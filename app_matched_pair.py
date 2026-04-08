"""
app.py  —  Oblique Aerial Survey Planner v3
============================================
Per-camera table, portrait/landscape, tilt-axis selector.
Three diagram views: footprint plan (all frames), cross-section, multi-strip.

Geometry verified against Oblique_setup9_working_2.xls — all values match.

Spreadsheet convention notes
─────────────────────────────
The Landscape sheet uses portrait-mounted L/R cameras (narrow axis across-track).
Our geometry.py orientation='portrait' reproduces those values exactly.
The Portrait sheet tab uses landscape-mounted oblique cameras (long axis across-track)
at a different flying height — it represents a different physical rig configuration.

Run:
    streamlit run app.py
"""

import json
import html
import math
from dataclasses import replace
from pathlib import Path
import io
from datetime import datetime
from itertools import product
from xml.etree import ElementTree as ET

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import matplotlib.colors as mcolors
import numpy as np
import streamlit as st
from openpyxl import Workbook
from openpyxl.drawing.image import Image as XLImage
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Pt
try:
    from shapely.geometry import Polygon as ShapelyPolygon, box as shapely_box, LineString as ShapelyLineString
    from shapely.ops import unary_union
    from shapely.affinity import translate as shapely_translate, rotate as shapely_rotate
    SHAPELY_AVAILABLE = True
except Exception:
    ShapelyPolygon = None
    shapely_box = None
    ShapelyLineString = None
    unary_union = None
    shapely_translate = None
    shapely_rotate = None
    SHAPELY_AVAILABLE = False

from geometry_matched_pair import (
    normalize_tilt_angle,
    half_fov_deg,
    diag_pp_to_long_edge_mm,
    flying_height_for_gsd,
    calculate_camera_solution,
    calculate_multicamera_solution,
    m_to_unit,
    unit_to_m,
)

# ─────────────────────────────────────────────────────────────────────────────
# Constants
# ─────────────────────────────────────────────────────────────────────────────

CAM_COLOURS = ["#58a6ff", "#f85149", "#3fb950", "#d29922", "#bc8cff", "#39c5cf"]

BODY_PRESETS = {
    "Sony A7R IV":       {"w_mm": 35.7,    "h_mm": 23.8,    "w_px": 9504,  "h_px": 6336},
  "Sony A7R V":        {"w_mm": 35.7, "h_mm": 23.8, "w_px": 9504, "h_px": 6336},
    "Sony A6500":        {"w_mm": 23.5,    "h_mm": 15.6,    "w_px": 6000,  "h_px": 4000},
    "Phase One iXM-100": {"w_mm": 53.4,    "h_mm": 40.0,    "w_px": 11664, "h_px": 8750},
    "Canon 5DS R":       {"w_mm": 36.0,    "h_mm": 24.0,    "w_px": 8688,  "h_px": 5792},
    "Nikon D850":        {"w_mm": 35.9,    "h_mm": 23.9,    "w_px": 8256,  "h_px": 5504},
    "Canon 760D":        {"w_mm": 22.3,    "h_mm": 14.9,    "w_px": 6000,  "h_px": 4000},
}

# Default 3-camera rig matching Oblique_setup9_working_2.xls Landscape sheet exactly
DEFAULT_CAMERAS = [
    {
        "enabled": True,
        "label": "Nadir",
        "body": "Sony A7R V",
        "focal_mm": 21.0,
        "tilt_deg": 0.0,
        "tilt_conv": "nadir",
        "orientation": "landscape",
        "tilt_axis": "across",
    },
    {
        "enabled": True,
        "label": "Right oblique",
        "body": "Sony A7R V",
        "focal_mm": 50.0,
        "tilt_deg": 50.0,
        "tilt_conv": "horiz",
        "orientation": "portrait",
        "tilt_axis": "across",
    },
    {
        "enabled": True,
        "label": "Left oblique",
        "body": "Sony A7R V",
        "focal_mm": 50.0,
        "tilt_deg": 50.0,
        "tilt_conv": "horiz",
        "orientation": "portrait",
        "tilt_axis": "across",
    },
    {
        "enabled": True,
        "label": "Fore oblique",
        "body": "Sony A7R V",
        "focal_mm": 50.0,
        "tilt_deg": 50.0,
        "tilt_conv": "horiz",
        "orientation": "portrait",
        "tilt_axis": "along",
    },
    {
        "enabled": True,
        "label": "Aft oblique",
        "body": "Sony A7R V",
        "focal_mm": 50.0,
        "tilt_deg": 50.0,
        "tilt_conv": "horiz",
        "orientation": "portrait",
        "tilt_axis": "along",
    },
]

PRESET_FILE = Path("presets.json")
SCENARIO_DIR = Path("saved_scenarios")
AOI_LIBRARY_DIR = Path("aoi_library")
DEFAULT_ALTITUDE_M = 600.0
DEFAULT_SPEED_MS = 62.0
DEFAULT_SPEED_KTS = DEFAULT_SPEED_MS * 1.94384
OVERLAP_PRESETS = {
    "Standard oblique": {"forward": 80, "side": 60},
    "Balanced": {"forward": 60, "side": 40},
    "Regional": {"forward": 40, "side": 20},
    "Open terrain": {"forward": 35, "side": 15},
    "Custom": None,
}
REPORT_LOGO_CANDIDATES = [
    Path("assets/aerial_surveys_logo.png"),
    Path("aerial_surveys_logo.png"),
    Path("/mnt/data/cdb78639-5338-4c0f-88eb-fa20a9521e12.png"),
]

# ─────────────────────────────────────────────────────────────────────────────
# Page config
# ─────────────────────────────────────────────────────────────────────────────

st.set_page_config(page_title="Oblique Survey Planner", page_icon="ASL_Imagery_Icon.png", layout="wide")
st.markdown(
    """
    <style>
    .stApp {
        background: #123b6d !important;
    }

    [data-testid="stAppViewContainer"] {
        background: #123b6d !important;
    }

    [data-testid="stHeader"] {
        background: #123b6d !important;
    }
    </style>
    """,
    unsafe_allow_html=True,
)
# ─────────────────────────────────────────────────────────────────────────────
# Persistence helpers
# ─────────────────────────────────────────────────────────────────────────────

def load_saved_bodies():
    if PRESET_FILE.exists():
        try:
            return json.loads(PRESET_FILE.read_text())
        except Exception:
            pass
    return {}

def save_body_preset(name, data):
    existing = load_saved_bodies()
    existing[name] = data
    PRESET_FILE.write_text(json.dumps(existing, indent=2))

def ensure_scenario_dir():
    SCENARIO_DIR.mkdir(parents=True, exist_ok=True)
    return SCENARIO_DIR


def normalise_scenario_name(name):
    cleaned = "".join(ch if ch.isalnum() or ch in ("-", "_", " ") else "_" for ch in str(name or "").strip())
    cleaned = " ".join(cleaned.split())
    return cleaned or "my_survey"


def scenario_path_from_name(name):
    safe_name = normalise_scenario_name(name)
    if not safe_name.lower().endswith(".json"):
        safe_name = f"{safe_name}.json"
    return ensure_scenario_dir() / safe_name


def is_scenario_payload(data):
    return isinstance(data, dict) and isinstance(data.get("cameras"), list)


def save_scenario(data, name):
    path = scenario_path_from_name(name)
    path.write_text(json.dumps(data, indent=2, default=str))
    return path


def load_scenario(path_or_name):
    candidate = Path(path_or_name)
    possible_paths = []
    if candidate.exists():
        possible_paths.append(candidate)
    else:
        possible_paths.append(scenario_path_from_name(path_or_name))
        if candidate.suffix:
            possible_paths.append(Path(candidate.name))
        else:
            possible_paths.append(Path(f"{candidate.name}.json"))

    for path in possible_paths:
        try:
            data = json.loads(Path(path).read_text())
            if is_scenario_payload(data):
                return data
        except Exception:
            continue
    return None


def list_saved_scenarios():
    ensure_scenario_dir()
    found = []
    seen = set()

    def _append(path, origin):
        try:
            data = json.loads(path.read_text())
        except Exception:
            return
        if not is_scenario_payload(data):
            return
        resolved = str(path.resolve())
        if resolved in seen:
            return
        seen.add(resolved)
        label = path.name if origin == "saved_scenarios" else f"{path.name} (project root)"
        found.append({"label": label, "path": str(path), "origin": origin})

    for path in sorted(SCENARIO_DIR.glob("*.json"), key=lambda p: p.name.lower()):
        _append(path, "saved_scenarios")

    for path in sorted(Path(".").glob("*.json"), key=lambda p: p.name.lower()):
        if path.name == PRESET_FILE.name:
            continue
        _append(path, "project_root")

    return found

def all_body_names():
    return list(BODY_PRESETS.keys()) + list(load_saved_bodies().keys())

def get_body(name):
    saved = load_saved_bodies()
    if name in BODY_PRESETS:
        return BODY_PRESETS[name]
    if name in saved:
        d = saved[name]
        return {"w_mm": d["w_mm"], "h_mm": d["h_mm"], "w_px": d["w_px"], "h_px": d["h_px"]}
    return BODY_PRESETS["Sony A7R V"]
def get_inner_outer_angles(sol):
    return sol.near_angle_deg, sol.far_angle_deg


def oblique_descriptor(angle_deg: float) -> str:
    if angle_deg < 5:
        return "Nadir"
    if angle_deg < 15:
        return "Mild"
    if angle_deg < 30:
        return "Moderate"
    if angle_deg < 45:
        return "Strong"
    return "Very strong"


def format_oblique(angle_deg: float) -> str:
    return f"{angle_deg:.1f}° ({oblique_descriptor(angle_deg)})"


def mirror_solution_for_label(sol):
    """Mirror canonical positive-tilt solutions to the correct side by label."""
    label = (sol.label or "").lower()

    if sol.tilt_axis == "across" and "left" in label:
        return replace(
            sol,
            near_edge_m=-sol.near_edge_m,
            centre_m=-sol.centre_m,
            far_edge_m=-sol.far_edge_m,
            corner_near_top=(-sol.corner_near_top[0], sol.corner_near_top[1]),
            corner_near_bot=(-sol.corner_near_bot[0], sol.corner_near_bot[1]),
            corner_far_top=(-sol.corner_far_top[0], sol.corner_far_top[1]),
            corner_far_bot=(-sol.corner_far_bot[0], sol.corner_far_bot[1]),
        )

    if sol.tilt_axis == "along" and any(k in label for k in ["aft", "rear", "back"]):
        return replace(
            sol,
            near_edge_m=-sol.near_edge_m,
            centre_m=-sol.centre_m,
            far_edge_m=-sol.far_edge_m,
            corner_near_top=(sol.corner_near_top[0], -sol.corner_near_top[1]),
            corner_near_bot=(sol.corner_near_bot[0], -sol.corner_near_bot[1]),
            corner_far_top=(sol.corner_far_top[0], -sol.corner_far_top[1]),
            corner_far_bot=(sol.corner_far_bot[0], -sol.corner_far_bot[1]),
        )

    return sol
def build_export_data(solutions, mc, altitude_m, speed_ms, fwd_frac, side_frac, dist_unit="m", reciprocal=True):
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
    speed_kts = speed_ms * 1.94384
    settings_rows = [
        ["Prepared", timestamp],
        ["Altitude AGL", f"{m_to_unit(altitude_m, dist_unit):.1f} {dist_unit}"],
        ["Aircraft speed", f"{speed_ms:.1f} m/s ({speed_kts:.1f} kts)"],
        ["Forward overlap target", f"{fwd_frac * 100:.1f}%"],
        ["Sidelap target", f"{side_frac * 100:.1f}%"],
        ["Reciprocal strips", "Yes" if reciprocal else "No"],
        ["Enabled cameras", ", ".join(sol.label for _, sol, _ in solutions)],
    ]

    system_rows = []
    if mc is not None:
        line_spacing = getattr(mc, "recommended_line_spacing_m", float("nan"))
        photo_spacing = getattr(mc, "recommended_photo_spacing_m", float("nan"))
        exposure_interval = getattr(mc, "photo_interval_s", float("nan"))
        combined_swath = getattr(mc, "combined_swath_m", float("nan"))
        sidelap_achieved = getattr(mc, "sidelap_achieved", float("nan"))
        system_rows.extend([
            ["Combined swath", f"{m_to_unit(combined_swath, dist_unit):.1f} {dist_unit}" if np.isfinite(combined_swath) else "—"],
            ["Recommended line spacing", f"{m_to_unit(line_spacing, dist_unit):.1f} {dist_unit}" if np.isfinite(line_spacing) else "—"],
            ["Recommended photo spacing", f"{m_to_unit(photo_spacing, dist_unit):.1f} {dist_unit}" if np.isfinite(photo_spacing) else "—"],
            ["Exposure interval", f"{exposure_interval:.2f} s" if np.isfinite(exposure_interval) else "—"],
            ["Achieved sidelap", f"{sidelap_achieved * 100:.1f}%" if np.isfinite(sidelap_achieved) else "—"],
        ])

    representative = next((s for _, s, _ in solutions if abs(s.tilt_from_nadir_deg) > 1), solutions[0][1] if solutions else None)
    if representative is not None:
        system_rows.extend([
            [f"Inner GSD ({representative.label})", fmt_gsd(min(representative.near_gsd_m, representative.far_gsd_m))],
            ["Centre GSD", fmt_gsd(representative.centre_gsd_m)],
            [f"Outer GSD ({representative.label})", fmt_gsd(max(representative.near_gsd_m, representative.far_gsd_m))],
            [f"Obliqueness ratio ({representative.label})", f"{obliqueness_ratio(representative):.2f}×"],
        ])

    camera_rows = []
    for cam, sol, _ in solutions:
        camera_rows.append([
            cam["label"],
            round(sol.near_angle_deg, 2),
            round(sol.far_angle_deg, 2),
            round(sol.near_angle_deg, 2),
            oblique_descriptor(sol.near_angle_deg),
            round(sol.far_angle_deg, 2),
            oblique_descriptor(sol.far_angle_deg),
            round(min(sol.near_gsd_m, sol.far_gsd_m) * 100, 2),
            round(sol.centre_gsd_m * 100, 2),
            round(max(sol.near_gsd_m, sol.far_gsd_m) * 100, 2),
            round(abs(sol.near_edge_m), 2),
            round(sol.centre_m, 2),
            round(abs(sol.far_edge_m), 2),
        ])

    return settings_rows, system_rows, camera_rows


def set_table_style(table, style_name):
    try:
        table.style = style_name
    except Exception:
        pass


def fig_to_png_bytes(fig, dpi=180):
    bio = io.BytesIO()
    fig.savefig(bio, format="png", dpi=dpi, bbox_inches="tight", facecolor=fig.get_facecolor())
    bio.seek(0)
    return bio.getvalue()


def find_report_logo():
    for candidate in REPORT_LOGO_CANDIDATES:
        candidate = Path(candidate)
        if candidate.exists():
            return candidate
    return None


def build_mission_export_rows(mission_outputs, dist_unit="m"):
    if mission_outputs is None:
        return []
    area_km2 = float(mission_outputs.get("area_m2", 0.0)) / 1_000_000.0
    total_line_km = float(mission_outputs.get("total_line_length_m", 0.0)) / 1000.0
    avg_line_km = float(mission_outputs.get("average_line_length_m", 0.0)) / 1000.0
    return [
        ["AOI name", mission_outputs.get("name", "AOI")],
        ["AOI source", mission_outputs.get("source", "unknown")],
        ["AOI area", f"{area_km2:.1f} km²"],
        ["Flight azimuth", f"{float(mission_outputs.get('flight_azimuth_deg', 0.0)):.1f}°"],
        ["Lead-in / out per line", f"{m_to_unit(float(mission_outputs.get('lead_in_out_m', 0.0)), dist_unit):.0f} {dist_unit}"],
        ["Flight lines", int(mission_outputs.get("line_count", 0))],
        ["Trigger events", int(mission_outputs.get("trigger_events", 0))],
        ["Frames per camera", int(mission_outputs.get("frames_per_camera", 0))],
        ["Total images", int(mission_outputs.get("total_images", 0))],
        ["Line spacing", fmt(float(mission_outputs.get("line_spacing_m", 0.0)), dist_unit)],
        ["Photo spacing", fmt(float(mission_outputs.get("photo_spacing_m", 0.0)), dist_unit)],
        ["Total line length", f"{total_line_km:,.2f} km"],
        ["Average line length", f"{avg_line_km:,.2f} km"],
        ["Estimated storage", f"{float(mission_outputs.get('total_storage_mb', 0.0)) / 1024.0:.1f} GB"],
        ["Estimated flying time", f"{float(mission_outputs.get('flight_time_s', 0.0)) / 60.0:.1f} min"],
    ]


def make_excel_export(settings_rows, system_rows, camera_rows, mission_rows=None, mission_figure_bytes=None):
    wb = Workbook()

    ws1 = wb.active
    ws1.title = "Settings"
    for row in settings_rows:
        ws1.append(row)

    ws2 = wb.create_sheet("System Summary")
    for row in system_rows:
        ws2.append(row)

    ws3 = wb.create_sheet("Camera Results")
    ws3.append([
        "Camera",
        "Inner edge oblique (deg)",
        "Outer edge oblique (deg)",
        "Near angle (deg)",
        "Near class",
        "Far angle (deg)",
        "Far class",
        "Inner GSD (cm/px)",
        "Centre GSD (cm/px)",
        "Outer GSD (cm/px)",
        "Inner edge (m)",
        "Centre (m)",
        "Outer edge (m)",
    ])

    for row in camera_rows:
        ws3.append(row)

    if mission_rows:
        ws4 = wb.create_sheet("Sample Area")
        ws4.append(["Metric", "Value"])
        for row in mission_rows:
            ws4.append(row)
        if mission_figure_bytes:
            try:
                img = XLImage(io.BytesIO(mission_figure_bytes))
                img.width = 520
                img.height = 390
                ws4.add_image(img, "D2")
            except Exception:
                pass

    bio = io.BytesIO()
    wb.save(bio)
    bio.seek(0)
    return bio.getvalue()


def make_word_export(settings_rows, system_rows, camera_rows, mission_rows=None, mission_figure_bytes=None, report_figures=None):
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    logo_path = find_report_logo()
    if logo_path is not None:
        doc.add_picture(str(logo_path), width=Inches(2.3))

    title = doc.add_paragraph()
    title_run = title.add_run("Oblique Aerial Survey Planner Report — Matched Pair")
    title_run.bold = True
    title_run.font.size = Pt(20)

    subtitle = doc.add_paragraph("Client summary of survey settings, coverage and camera geometry.")
    subtitle.runs[0].italic = True

    doc.add_heading("Summary of settings", level=2)
    table = doc.add_table(rows=1, cols=2)
    set_table_style(table, "Light List Accent 1")
    hdr = table.rows[0].cells
    hdr[0].text = "Item"
    hdr[1].text = "Value"

    for item, value in settings_rows:
        row = table.add_row().cells
        row[0].text = str(item)
        row[1].text = str(value)

    doc.add_heading("System summary", level=2)
    table_sys = doc.add_table(rows=1, cols=2)
    set_table_style(table_sys, "Light Grid Accent 1")
    hdr_sys = table_sys.rows[0].cells
    hdr_sys[0].text = "Metric"
    hdr_sys[1].text = "Value"
    for item, value in system_rows:
        row = table_sys.add_row().cells
        row[0].text = str(item)
        row[1].text = str(value)

    if mission_rows:
        doc.add_heading("Sample area and flight plan", level=2)
        mission_table = doc.add_table(rows=1, cols=2)
        set_table_style(mission_table, "Light Grid Accent 1")
        hdr_m = mission_table.rows[0].cells
        hdr_m[0].text = "Metric"
        hdr_m[1].text = "Value"
        for item, value in mission_rows:
            row = mission_table.add_row().cells
            row[0].text = str(item)
            row[1].text = str(value)
        if mission_figure_bytes:
            title_p = doc.add_paragraph("Sample area and generated flight lines")
            title_p.paragraph_format.keep_with_next = True
            doc.add_picture(io.BytesIO(mission_figure_bytes), width=Inches(5.8))

    doc.add_heading("Camera results", level=2)
    table2 = doc.add_table(rows=1, cols=13)
    set_table_style(table2, "Medium Grid 1 Accent 1")
    hdr2 = table2.rows[0].cells
    headers = [
        "Camera",
        "Inner edge oblique",
        "Outer edge oblique",
        "Near angle",
        "Near class",
        "Far angle",
        "Far class",
        "Inner GSD",
        "Centre GSD",
        "Outer GSD",
        "Inner edge",
        "Centre",
        "Outer edge",
    ]

    for i, h in enumerate(headers):
        hdr2[i].text = h

    for row_data in camera_rows:
        row = table2.add_row().cells
        for i, value in enumerate(row_data):
            row[i].text = str(value)

    if report_figures:
        doc.add_heading("Coverage diagrams", level=2)
        for figure_title, figure_bytes in report_figures:
            para = doc.add_paragraph(figure_title)
            para.paragraph_format.keep_with_next = True
            doc.add_picture(io.BytesIO(figure_bytes), width=Inches(6.7))

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# Session state
# ─────────────────────────────────────────────────────────────────────────────

if "cameras" not in st.session_state:
    st.session_state.cameras = [dict(c) for c in DEFAULT_CAMERAS]
if "planner_altitude_m" not in st.session_state:
    st.session_state.planner_altitude_m = DEFAULT_ALTITUDE_M
if "planner_dist_unit_last" not in st.session_state:
    st.session_state.planner_dist_unit_last = "m"
if "sidebar_alt_input" not in st.session_state:
    st.session_state.sidebar_alt_input = round(m_to_unit(st.session_state.planner_altitude_m, "m"), 1)
if "sidebar_speed_ms" not in st.session_state:
    st.session_state.sidebar_speed_ms = round(DEFAULT_SPEED_MS, 1)
if "sidebar_reciprocal" not in st.session_state:
    st.session_state.sidebar_reciprocal = True
if "pending_loaded_scenario" not in st.session_state:
    st.session_state.pending_loaded_scenario = None
if "saved_scenario_selector" not in st.session_state:
    st.session_state.saved_scenario_selector = None
if "saved_scenario_selector_prev" not in st.session_state:
    st.session_state.saved_scenario_selector_prev = None
if "camera_widget_nonce" not in st.session_state:
    st.session_state.camera_widget_nonce = 0
if "pending_optimizer_apply" not in st.session_state:
    st.session_state.pending_optimizer_apply = None
if "selected_aoi_name" not in st.session_state:
    st.session_state.selected_aoi_name = None
if "selected_aoi_payload" not in st.session_state:
    st.session_state.selected_aoi_payload = None


def bump_camera_widget_nonce():
    st.session_state.camera_widget_nonce = int(st.session_state.get("camera_widget_nonce", 0)) + 1

# ─────────────────────────────────────────────────────────────────────────────
# Helpers — geometry and display
# ─────────────────────────────────────────────────────────────────────────────

def fmt(v_m, unit, d=1):
    return f"{m_to_unit(v_m, unit):.{d}f} {unit}"

def fmt_gsd(v_m, d=2):
    return f"{v_m * 100:.{d}f} cm/px"

def obliqueness_ratio(sol):
    """
    Ratio of far-edge GSD to inner-edge GSD.  1.0 = nadir (no variation).
    Higher = more oblique — GSD varies more across the image.
    """
    inner_gsd = min(sol.near_gsd_m, sol.far_gsd_m)
    outer_gsd = max(sol.near_gsd_m, sol.far_gsd_m)
    if inner_gsd <= 0:
        return float("inf")
    return outer_gsd / inner_gsd

def dark_fig(w=12, h=6):
    fig, ax = plt.subplots(figsize=(w, h))
    fig.patch.set_facecolor("#0d1117")
    ax.set_facecolor("#161b22")
    ax.tick_params(colors="#8b949e", labelsize=8)
    for sp in ax.spines.values():
        sp.set_color("#30363d")
    return fig, ax

def corner_inner_outer(sol):
    """
    Return (inner_gx_or_gy, outer_gx_or_gy) where inner = edge closer to nadir.
    Works for both positive-tilt (right) and negative-tilt (left) cameras.
    """
    a, b = sol.near_edge_m, sol.far_edge_m
    if np.isfinite(a) and np.isfinite(b):
        return (a, b) if abs(a) <= abs(b) else (b, a)

    extent = polygon_extent(sol)
    if extent is not None:
        xmin, xmax, ymin, ymax = extent
        if sol.tilt_axis == "across":
            return (xmin, xmax) if abs(xmin) <= abs(xmax) else (xmax, xmin)
        return (ymin, ymax) if abs(ymin) <= abs(ymax) else (ymax, ymin)

    return 0.0, 0.0

def inner_outer_corners(sol):
    """
    Split the 4 footprint corners into (inner_top, inner_bot) and (outer_top, outer_bot).
    Inner = nadir side. Falls back to the rendered footprint polygon for near-nadir cases.
    """
    nt, nb = sol.corner_near_top, sol.corner_near_bot
    ft, fb = sol.corner_far_top, sol.corner_far_bot
    raw = [nt, nb, ft, fb]

    if all(np.isfinite(v) for xy in raw for v in xy):
        if abs(nt[0]) <= abs(ft[0]):
            return (nt, nb), (ft, fb)
        return (ft, fb), (nt, nb)

    poly = camera_polygon(sol)
    if len(poly) != 4:
        return ((0.0, 0.0), (0.0, 0.0)), ((0.0, 0.0), (0.0, 0.0))

    eps = 1e-6
    if sol.tilt_axis == "across":
        inner_abs = min(abs(p[0]) for p in poly)
        outer_abs = max(abs(p[0]) for p in poly)
        inner = [p for p in poly if abs(abs(p[0]) - inner_abs) < eps or np.isclose(abs(p[0]), inner_abs)]
        outer = [p for p in poly if abs(abs(p[0]) - outer_abs) < eps or np.isclose(abs(p[0]), outer_abs)]
        inner = sorted(inner, key=lambda p: p[1], reverse=True)[:2]
        outer = sorted(outer, key=lambda p: p[1], reverse=True)[:2]
    else:
        inner_abs = min(abs(p[1]) for p in poly)
        outer_abs = max(abs(p[1]) for p in poly)
        inner = [p for p in poly if abs(abs(p[1]) - inner_abs) < eps or np.isclose(abs(p[1]), inner_abs)]
        outer = [p for p in poly if abs(abs(p[1]) - outer_abs) < eps or np.isclose(abs(p[1]), outer_abs)]
        inner = sorted(inner, key=lambda p: p[0], reverse=True)[:2]
        outer = sorted(outer, key=lambda p: p[0], reverse=True)[:2]

    while len(inner) < 2:
        inner.append(inner[0])
    while len(outer) < 2:
        outer.append(outer[0])
    return (inner[0], inner[1]), (outer[0], outer[1])

def along_lengths_for_display(sol):
    """Inner and outer along-track (or across-track for along-tilt) footprint lengths."""
    (it, ib), (ot, ob) = inner_outer_corners(sol)
    return abs(it[1] - ib[1]), abs(ot[1] - ob[1])

def safe_corners(sol):
    """Return all 4 corner (Gx, Gy) tuples, filtered to finite values only."""
    return [c for c in camera_polygon(sol) if np.isfinite(c[0]) and np.isfinite(c[1])]

def axis_limits_from_solutions(solutions, pad=0.18):
    """
    Compute square axis limits that fit all camera footprints with a margin.
    Returns (lim,) where axes run from -lim to +lim.
    Falls back to 1.0 if no finite corners found.
    """
    all_x, all_y = [], []
    for _, sol, _ in solutions:
        for c in safe_corners(sol):
            all_x.append(c[0])
            all_y.append(c[1])
    if not all_x:
        return 1.0
    x_span = max(abs(min(all_x)), abs(max(all_x))) * (1 + pad)
    y_span = max(abs(min(all_y)), abs(max(all_y))) * (1 + pad)
    lim = max(x_span, y_span)
    return lim if np.isfinite(lim) and lim > 0 else 1.0


def camera_polygon(sol):
    corners = [
        (sol.corner_near_top[0], sol.corner_near_top[1]),
        (sol.corner_far_top[0],  sol.corner_far_top[1]),
        (sol.corner_far_bot[0],  sol.corner_far_bot[1]),
        (sol.corner_near_bot[0], sol.corner_near_bot[1]),
    ]
    if all(np.isfinite(v) for xy in corners for v in xy):
        return corners

    if abs(getattr(sol, "tilt_from_nadir_deg", 0.0)) < 0.25:
        half_across = sol.centre_slant_m * math.tan(math.radians(sol.half_fov_across_deg))
        half_along = sol.centre_slant_m * math.tan(math.radians(sol.half_fov_along_deg))
        if np.isfinite(half_across) and np.isfinite(half_along):
            return [
                (-half_across, half_along),
                (half_across, half_along),
                (half_across, -half_along),
                (-half_across, -half_along),
            ]

    return []


def polygon_extent(sol):
    poly = camera_polygon(sol)
    if not poly:
        return None
    xs = [p[0] for p in poly]
    ys = [p[1] for p in poly]
    return min(xs), max(xs), min(ys), max(ys)


def fallback_multistrip_spacing(multistrip_solutions, fwd_frac, side_frac, mc=None):
    line_spacing = getattr(mc, "recommended_line_spacing_m", float("nan")) if mc is not None else float("nan")
    photo_spacing = getattr(mc, "recommended_photo_spacing_m", float("nan")) if mc is not None else float("nan")
    used_fallback = False

    extents = [polygon_extent(sol) for _, sol, _ in multistrip_solutions]
    extents = [e for e in extents if e is not None]

    if not (np.isfinite(line_spacing) and line_spacing > 0) and extents:
        swath_width = max(e[1] for e in extents) - min(e[0] for e in extents)
        line_spacing = swath_width * max(0.05, (1.0 - side_frac))
        used_fallback = True

    if not (np.isfinite(photo_spacing) and photo_spacing > 0) and extents:
        lengths = [e[3] - e[2] for e in extents if np.isfinite(e[3] - e[2]) and (e[3] - e[2]) > 0]
        if lengths:
            photo_spacing = min(lengths) * max(0.05, (1.0 - fwd_frac))
            used_fallback = True

    if not (np.isfinite(line_spacing) and line_spacing > 0):
        line_spacing = 1.0
    if not (np.isfinite(photo_spacing) and photo_spacing > 0):
        photo_spacing = 1.0

    return line_spacing, photo_spacing, used_fallback


def matched_sidelap_band(multistrip_solutions, line_spacing, frame_y=0.0):
    """
    Return a simple cross-track overlap band for the matched reciprocal pair:
    Right oblique on the current strip versus Left oblique on the adjacent strip.
    This is intended to communicate sidelap as edge-to-edge cross-track overlap,
    not as shared plan-area intersection of skewed polygons.
    """
    right_entry = next(
        ((cam, sol, col) for cam, sol, col in multistrip_solutions
         if sol.tilt_axis == "across" and "right" in sol.label.lower()),
        None,
    )
    left_entry = next(
        ((cam, sol, col) for cam, sol, col in multistrip_solutions
         if sol.tilt_axis == "across" and "left" in sol.label.lower()),
        None,
    )
    if right_entry is None or left_entry is None:
        return None

    _, right_sol, right_col = right_entry
    _, left_sol, _ = left_entry
    right_extent = polygon_extent(right_sol)
    left_extent = polygon_extent(left_sol)
    if right_extent is None or left_extent is None:
        return None

    r_x0, r_x1, r_y0, r_y1 = right_extent
    l_x0, l_x1, l_y0, l_y1 = left_extent

    # Current strip at x offset 0; adjacent reciprocal strip at +line_spacing
    band_x0 = max(r_x0, l_x0 + line_spacing)
    band_x1 = min(r_x1, l_x1 + line_spacing)
    if not (np.isfinite(band_x0) and np.isfinite(band_x1) and band_x1 > band_x0):
        return None

    band_y0 = max(r_y0 + frame_y, l_y0 + frame_y)
    band_y1 = min(r_y1 + frame_y, l_y1 + frame_y)
    if not (np.isfinite(band_y0) and np.isfinite(band_y1) and band_y1 > band_y0):
        return None

    right_width = r_x1 - r_x0
    overlap_width = band_x1 - band_x0
    overlap_frac = overlap_width / right_width if right_width > 0 else float("nan")

    return {
        "x0": band_x0,
        "x1": band_x1,
        "y0": band_y0,
        "y1": band_y1,
        "overlap_width": overlap_width,
        "right_width": right_width,
        "overlap_frac": overlap_frac,
        "colour": right_col,
    }

def point_on_segment(px, py, ax, ay, bx, by, eps=1e-9):
    cross = (px - ax) * (by - ay) - (py - ay) * (bx - ax)
    if abs(cross) > eps:
        return False
    dot = (px - ax) * (bx - ax) + (py - ay) * (by - ay)
    if dot < -eps:
        return False
    sq_len = (bx - ax) ** 2 + (by - ay) ** 2
    if dot - sq_len > eps:
        return False
    return True


def point_in_polygon(px, py, polygon):
    inside = False
    n = len(polygon)
    for i in range(n):
        x1, y1 = polygon[i]
        x2, y2 = polygon[(i + 1) % n]
        if point_on_segment(px, py, x1, y1, x2, y2):
            return True
        intersects = ((y1 > py) != (y2 > py)) and (px < (x2 - x1) * (py - y1) / ((y2 - y1) or 1e-12) + x1)
        if intersects:
            inside = not inside
    return inside


def coverage_view_family(label):
    lower = (label or "").strip().lower()
    if "nadir" in lower:
        return "Nadir"
    if "left" in lower:
        return "Left oblique"
    if "right" in lower:
        return "Right oblique"
    if any(token in lower for token in ["fore", "forward", "front"]):
        return "Fore oblique"
    if any(token in lower for token in ["aft", "rear", "back"]):
        return "Aft oblique"
    return (label or "Camera").strip() or "Camera"


def build_coverage_sources(solutions, line_spacing, photo_spacing):
    base_polygons = []
    max_abs_x = 0.0
    max_abs_y = 0.0
    for cam, sol, _ in solutions:
        poly = camera_polygon(sol)
        if len(poly) != 4:
            continue
        max_abs_x = max(max_abs_x, max(abs(x) for x, _ in poly))
        max_abs_y = max(max_abs_y, max(abs(y) for _, y in poly))
        source = {
            "label": sol.label,
            "family": coverage_view_family(sol.label),
            "polygon": poly,
        }
        if SHAPELY_AVAILABLE:
            try:
                geom = ShapelyPolygon(poly)
                if not geom.is_valid:
                    geom = geom.buffer(0)
                source["shape"] = geom
            except Exception:
                source["shape"] = None
        base_polygons.append(source)

    if not base_polygons:
        return []

    strip_repeats = max(1, int(math.ceil(max_abs_x / max(line_spacing, 1e-6))) + 1)
    frame_repeats = max(1, int(math.ceil(max_abs_y / max(photo_spacing, 1e-6))) + 1)

    sources = []
    for base in base_polygons:
        for strip_idx in range(-strip_repeats, strip_repeats + 1):
            dx = strip_idx * line_spacing
            for frame_idx in range(-frame_repeats, frame_repeats + 1):
                dy = frame_idx * photo_spacing
                shifted = [(x + dx, y + dy) for x, y in base["polygon"]]
                xs = [x for x, _ in shifted]
                ys = [y for _, y in shifted]
                item = {
                    "label": base["label"],
                    "family": base["family"],
                    "polygon": shifted,
                    "bbox": (min(xs), max(xs), min(ys), max(ys)),
                }
                if SHAPELY_AVAILABLE and base.get("shape") is not None:
                    try:
                        item["shape"] = shapely_translate(base["shape"], xoff=dx, yoff=dy)
                    except Exception:
                        item["shape"] = None
                sources.append(item)
    return sources


def compute_exact_gap_stats(sources, line_spacing, photo_spacing):
    if not SHAPELY_AVAILABLE or not sources or line_spacing <= 0 or photo_spacing <= 0:
        return None

    try:
        repeat_cell = shapely_box(-0.5 * line_spacing, -0.5 * photo_spacing, 0.5 * line_spacing, 0.5 * photo_spacing)
        clipped_shapes = []
        for source in sources:
            geom = source.get("shape")
            if geom is None:
                geom = ShapelyPolygon(source["polygon"])
            if geom.is_empty:
                continue
            clipped = geom.intersection(repeat_cell)
            if not clipped.is_empty:
                clipped_shapes.append(clipped)
        covered = unary_union(clipped_shapes) if clipped_shapes else None
        uncovered = repeat_cell if covered is None else repeat_cell.difference(covered)
        total_area = float(repeat_cell.area)
        uncovered_area = float(uncovered.area) if uncovered is not None else total_area
        tol = max(total_area, 1.0) * 1e-9
        return {
            "available": True,
            "total_area": total_area,
            "zero_hit_area": uncovered_area,
            "zero_hit_pct": (100.0 * uncovered_area / total_area) if total_area else 0.0,
            "zero_angle_area": uncovered_area,
            "zero_angle_pct": (100.0 * uncovered_area / total_area) if total_area else 0.0,
            "has_gap": uncovered_area > tol,
        }
    except Exception:
        return None


def compute_point_coverage(solutions, line_spacing, photo_spacing, samples_x=45, samples_y=45):
    sources = build_coverage_sources(solutions, line_spacing, photo_spacing)
    if not sources:
        return None

    xs = np.linspace(-0.5 * line_spacing, 0.5 * line_spacing, samples_x)
    ys = np.linspace(-0.5 * photo_spacing, 0.5 * photo_spacing, samples_y)
    hits = np.zeros((len(ys), len(xs)), dtype=int)
    angle_counts = np.zeros((len(ys), len(xs)), dtype=int)

    for iy, y in enumerate(ys):
        for ix, x in enumerate(xs):
            labels_here = []
            families_here = set()
            for source in sources:
                xmin, xmax, ymin, ymax = source["bbox"]
                if x < xmin or x > xmax or y < ymin or y > ymax:
                    continue
                if point_in_polygon(x, y, source["polygon"]):
                    labels_here.append(source["label"])
                    families_here.add(source["family"])
            hits[iy, ix] = len(labels_here)
            angle_counts[iy, ix] = len(families_here)

    exact_gap_stats = compute_exact_gap_stats(sources, line_spacing, photo_spacing)

    return {
        "xs": xs,
        "ys": ys,
        "hits": hits,
        "angle_counts": angle_counts,
        "sources": sources,
        "line_spacing": line_spacing,
        "photo_spacing": photo_spacing,
        "exact_gap_stats": exact_gap_stats,
    }


def point_coverage_at(x, y, sources):
    labels_here = []
    families_here = []
    for source in sources:
        xmin, xmax, ymin, ymax = source["bbox"]
        if x < xmin or x > xmax or y < ymin or y > ymax:
            continue
        if point_in_polygon(x, y, source["polygon"]):
            labels_here.append(source["label"])
            families_here.append(source["family"])
    unique_families = sorted(set(families_here))
    return {
        "hits": len(labels_here),
        "unique_angles": len(unique_families),
        "labels": labels_here,
        "families": unique_families,
    }


def coverage_summary(coverage):
    hits = coverage["hits"]
    angle_counts = coverage["angle_counts"]
    return {
        "hits_min": int(np.min(hits)),
        "hits_avg": float(np.mean(hits)),
        "hits_max": int(np.max(hits)),
        "angles_min": int(np.min(angle_counts)),
        "angles_avg": float(np.mean(angle_counts)),
        "angles_max": int(np.max(angle_counts)),
    }


def coverage_gap_stats(coverage):
    hits = coverage["hits"]
    angle_counts = coverage["angle_counts"]
    sample_total = int(hits.size)
    sample_zero_hit_points = int(np.sum(hits == 0))
    sample_zero_angle_points = int(np.sum(angle_counts == 0))
    stats = {
        "sample_total_points": sample_total,
        "sample_zero_hit_points": sample_zero_hit_points,
        "sample_zero_hit_pct": (100.0 * sample_zero_hit_points / sample_total) if sample_total else 0.0,
        "sample_zero_angle_points": sample_zero_angle_points,
        "sample_zero_angle_pct": (100.0 * sample_zero_angle_points / sample_total) if sample_total else 0.0,
        "has_gap": sample_zero_hit_points > 0 or sample_zero_angle_points > 0,
        "exact_available": False,
    }

    exact_gap_stats = coverage.get("exact_gap_stats") if isinstance(coverage, dict) else None
    if exact_gap_stats is not None:
        stats.update({
            "exact_available": True,
            "zero_hit_area": exact_gap_stats.get("zero_hit_area", 0.0),
            "zero_hit_pct": exact_gap_stats.get("zero_hit_pct", 0.0),
            "zero_angle_area": exact_gap_stats.get("zero_angle_area", 0.0),
            "zero_angle_pct": exact_gap_stats.get("zero_angle_pct", 0.0),
            "has_gap": bool(exact_gap_stats.get("has_gap", False)),
        })
    else:
        stats.update({
            "zero_hit_pct": stats["sample_zero_hit_pct"],
            "zero_angle_pct": stats["sample_zero_angle_pct"],
        })
    return stats




def format_gap_pct(pct: float) -> str:
    if pct <= 0:
        return "0.00%"
    if pct < 0.01:
        return "<0.01%"
    return f"{pct:.2f}%"


def coverage_sampling_label(samples_x: int, samples_y: int) -> str:
    total = int(samples_x) * int(samples_y)
    return f"{samples_x} × {samples_y} sample grid ({total:,} points)"


def classify_gap_presentation(gap_stats: dict) -> str:
    if not gap_stats.get("has_gap"):
        return "none"
    exact_available = bool(gap_stats.get("exact_available", False))
    sample_zero_hit = float(gap_stats.get("sample_zero_hit_pct", 0.0))
    sample_zero_angle = float(gap_stats.get("sample_zero_angle_pct", 0.0))
    exact_zero_hit = float(gap_stats.get("zero_hit_pct", 0.0))
    if exact_available and exact_zero_hit < 0.01 and sample_zero_hit == 0.0 and sample_zero_angle == 0.0:
        return "micro"
    return "shortfall"

def current_common_oblique_value(camera_configs, field_name, fallback):
    values = []
    for cam in camera_configs:
        if not cam.get("enabled", True):
            continue
        if "nadir" in str(cam.get("label", "")).lower():
            continue
        try:
            values.append(float(cam.get(field_name, fallback)))
        except Exception:
            continue
    if not values:
        return float(fallback)
    return round(sum(values) / len(values), 1)


def apply_common_oblique_settings(camera_configs, tilt_deg=None, focal_mm=None):
    for cam in camera_configs:
        if not cam.get("enabled", True):
            continue
        if "nadir" in str(cam.get("label", "")).lower():
            continue
        if tilt_deg is not None:
            cam["tilt_deg"] = float(tilt_deg)
        if focal_mm is not None:
            cam["focal_mm"] = float(focal_mm)
    return camera_configs


def build_solutions_for_optimizer(camera_configs, altitude_m):
    active_local = [dict(c) for c in camera_configs if c.get("enabled", False)]
    solutions_local = []
    errors_local = []
    for i, cam in enumerate(active_local):
        try:
            bd = get_body(cam["body"])
            tilt_n = normalize_tilt_angle(cam["tilt_deg"], cam["tilt_conv"])
            sol = calculate_camera_solution(
                altitude_m=altitude_m,
                tilt_from_nadir_deg=tilt_n,
                sensor_w_native_mm=bd["w_mm"],
                sensor_h_native_mm=bd["h_mm"],
                image_w_native_px=bd["w_px"],
                image_h_native_px=bd["h_px"],
                focal_length_mm=cam["focal_mm"],
                orientation=cam["orientation"],
                tilt_axis=cam["tilt_axis"],
                label=cam["label"],
            )
            sol = mirror_solution_for_label(sol)
            solutions_local.append((cam, sol, CAM_COLOURS[i % len(CAM_COLOURS)]))
        except Exception as exc:
            errors_local.append(f"{cam.get('label', 'Camera')}: {exc}")
    return solutions_local, errors_local


def optimizer_penalty(candidate, requirements):
    penalty = 0.0
    penalty += max(0.0, float(requirements["min_viewing_angles"]) - float(candidate["angles_min"])) * 250.0
    penalty += max(0.0, float(requirements["min_image_hits"]) - float(candidate["hits_min"])) * 120.0
    if requirements.get("require_no_gaps", True) and candidate.get("has_gap", False):
        penalty += 1000.0 + float(candidate.get("zero_hit_pct", 0.0)) * 50.0
    if requirements.get("max_inner_gsd_cm") is not None:
        penalty += max(0.0, float(candidate["inner_gsd_cm"]) - float(requirements["max_inner_gsd_cm"])) * 15.0
    if requirements.get("max_outer_gsd_cm") is not None:
        penalty += max(0.0, float(candidate["outer_gsd_cm"]) - float(requirements["max_outer_gsd_cm"])) * 25.0
    if requirements.get("min_achieved_sidelap_pct") is not None:
        penalty += max(0.0, float(requirements["min_achieved_sidelap_pct"]) - float(candidate["achieved_sidelap_pct"])) * 6.0
    return penalty


def evaluate_optimizer_candidate(
    base_cameras,
    altitude_m,
    speed_ms,
    forward_overlap_pct,
    side_overlap_pct,
    reciprocal,
    samples_x,
    samples_y,
    requirements,
    common_oblique_tilt_deg=None,
    common_oblique_focal_mm=None,
):
    candidate_cameras = [dict(cam) for cam in base_cameras]
    if common_oblique_tilt_deg is not None or common_oblique_focal_mm is not None:
        candidate_cameras = apply_common_oblique_settings(
            candidate_cameras,
            tilt_deg=common_oblique_tilt_deg,
            focal_mm=common_oblique_focal_mm,
        )

    solutions_local, errors_local = build_solutions_for_optimizer(candidate_cameras, altitude_m)
    if errors_local:
        return {
            "valid": False,
            "reason": "; ".join(errors_local),
            "forward_pct": int(forward_overlap_pct),
            "side_pct": int(side_overlap_pct),
            "reciprocal": bool(reciprocal),
            "altitude_m": float(altitude_m),
            "speed_ms": float(speed_ms),
        }
    if not solutions_local:
        return {
            "valid": False,
            "reason": "No enabled cameras to optimise.",
            "forward_pct": int(forward_overlap_pct),
            "side_pct": int(side_overlap_pct),
            "reciprocal": bool(reciprocal),
            "altitude_m": float(altitude_m),
            "speed_ms": float(speed_ms),
        }

    try:
        mc_local = calculate_multicamera_solution(
            camera_solutions=[sol for _, sol, _ in solutions_local],
            arrangement="custom",
            altitude_m=altitude_m,
            aircraft_speed_ms=speed_ms,
            forward_overlap_fraction=float(forward_overlap_pct) / 100.0,
            sidelap_fraction=float(side_overlap_pct) / 100.0,
            reciprocal_flying=reciprocal,
        )
    except Exception as exc:
        return {
            "valid": False,
            "reason": str(exc),
            "forward_pct": int(forward_overlap_pct),
            "side_pct": int(side_overlap_pct),
            "reciprocal": bool(reciprocal),
            "altitude_m": float(altitude_m),
            "speed_ms": float(speed_ms),
        }

    line_spacing_local, photo_spacing_local, used_fallback = fallback_multistrip_spacing(
        solutions_local,
        fwd_frac=float(forward_overlap_pct) / 100.0,
        side_frac=float(side_overlap_pct) / 100.0,
        mc=mc_local,
    )
    coverage_local = compute_point_coverage(
        solutions_local,
        line_spacing=line_spacing_local,
        photo_spacing=photo_spacing_local,
        samples_x=samples_x,
        samples_y=samples_y,
    )
    if coverage_local is None:
        return {
            "valid": False,
            "reason": "Coverage could not be estimated for this candidate.",
            "forward_pct": int(forward_overlap_pct),
            "side_pct": int(side_overlap_pct),
            "reciprocal": bool(reciprocal),
            "altitude_m": float(altitude_m),
            "speed_ms": float(speed_ms),
        }

    summary_local = coverage_summary(coverage_local)
    gap_stats_local = coverage_gap_stats(coverage_local)

    reference_solutions = [sol for _, sol, _ in solutions_local if abs(sol.tilt_from_nadir_deg) > 1.0]
    if not reference_solutions:
        reference_solutions = [sol for _, sol, _ in solutions_local]

    inner_gsd_cm = max((min(sol.near_gsd_m, sol.far_gsd_m) * 100.0) for sol in reference_solutions)
    outer_gsd_cm = max((max(sol.near_gsd_m, sol.far_gsd_m) * 100.0) for sol in reference_solutions)
    min_oblique_angle = min(min(sol.near_angle_deg, sol.far_angle_deg) for sol in reference_solutions)
    max_oblique_angle = max(max(sol.near_angle_deg, sol.far_angle_deg) for sol in reference_solutions)
    achieved_sidelap_pct = mc_local.sidelap_achieved * 100.0 if np.isfinite(mc_local.sidelap_achieved) else float("nan")
    reciprocal_factor = 2.0 if reciprocal else 1.0
    efficiency_index_m2 = (line_spacing_local * photo_spacing_local) / reciprocal_factor if line_spacing_local > 0 and photo_spacing_local > 0 else 0.0

    result = {
        "valid": True,
        "reason": "Pass",
        "forward_pct": int(forward_overlap_pct),
        "side_pct": int(side_overlap_pct),
        "reciprocal": bool(reciprocal),
        "altitude_m": float(altitude_m),
        "speed_ms": float(speed_ms),
        "common_oblique_tilt_deg": current_common_oblique_value(candidate_cameras, "tilt_deg", 0.0),
        "common_oblique_focal_mm": current_common_oblique_value(candidate_cameras, "focal_mm", 0.0),
        "hits_min": int(summary_local["hits_min"]),
        "hits_avg": float(summary_local["hits_avg"]),
        "hits_max": int(summary_local["hits_max"]),
        "angles_min": int(summary_local["angles_min"]),
        "angles_avg": float(summary_local["angles_avg"]),
        "angles_max": int(summary_local["angles_max"]),
        "has_gap": bool(gap_stats_local.get("has_gap", False)),
        "zero_hit_pct": float(gap_stats_local.get("zero_hit_pct", 0.0)),
        "exact_gap_available": bool(gap_stats_local.get("exact_available", False)),
        "inner_gsd_cm": float(inner_gsd_cm),
        "outer_gsd_cm": float(outer_gsd_cm),
        "min_oblique_angle_deg": float(min_oblique_angle),
        "max_oblique_angle_deg": float(max_oblique_angle),
        "achieved_sidelap_pct": float(achieved_sidelap_pct),
        "line_spacing_m": float(line_spacing_local),
        "photo_spacing_m": float(photo_spacing_local),
        "efficiency_index_m2": float(efficiency_index_m2),
        "used_fallback": bool(used_fallback),
    }

    failure_reasons = []
    if result["angles_min"] < int(requirements["min_viewing_angles"]):
        failure_reasons.append(f"min viewing angles {result['angles_min']} < {int(requirements['min_viewing_angles'])}")
    if result["hits_min"] < int(requirements["min_image_hits"]):
        failure_reasons.append(f"min image hits {result['hits_min']} < {int(requirements['min_image_hits'])}")
    if requirements.get("require_no_gaps", True) and result["has_gap"]:
        failure_reasons.append(f"coverage gap {format_gap_pct(result['zero_hit_pct'])}")
    if requirements.get("max_inner_gsd_cm") is not None and result["inner_gsd_cm"] > float(requirements["max_inner_gsd_cm"]):
        failure_reasons.append(f"inner GSD {result['inner_gsd_cm']:.2f} cm > {float(requirements['max_inner_gsd_cm']):.2f} cm")
    if requirements.get("max_outer_gsd_cm") is not None and result["outer_gsd_cm"] > float(requirements["max_outer_gsd_cm"]):
        failure_reasons.append(f"outer GSD {result['outer_gsd_cm']:.2f} cm > {float(requirements['max_outer_gsd_cm']):.2f} cm")
    if requirements.get("min_achieved_sidelap_pct") is not None and np.isfinite(result["achieved_sidelap_pct"]) and result["achieved_sidelap_pct"] < float(requirements["min_achieved_sidelap_pct"]):
        failure_reasons.append(f"achieved sidelap {result['achieved_sidelap_pct']:.1f}% < {float(requirements['min_achieved_sidelap_pct']):.1f}%")

    result["valid"] = len(failure_reasons) == 0
    result["reason"] = "Pass" if result["valid"] else "; ".join(failure_reasons)
    result["penalty"] = optimizer_penalty(result, requirements)
    return result


def optimizer_valid_sort_key(candidate):
    return (
        -float(candidate.get("efficiency_index_m2", 0.0)),
        float(candidate.get("outer_gsd_cm", 0.0)),
        float(candidate.get("inner_gsd_cm", 0.0)),
        -float(candidate.get("angles_avg", 0.0)),
        -float(candidate.get("hits_avg", 0.0)),
        float(candidate.get("forward_pct", 0.0) + candidate.get("side_pct", 0.0)),
    )


def optimizer_near_miss_sort_key(candidate):
    return (
        float(candidate.get("penalty", 0.0)),
        -float(candidate.get("efficiency_index_m2", 0.0)),
        float(candidate.get("outer_gsd_cm", 0.0)),
        float(candidate.get("inner_gsd_cm", 0.0)),
    )


def optimizer_candidate_row(candidate, unit):
    return {
        "Status": "Pass" if candidate.get("valid") else "Near miss",
        "Forward %": int(candidate.get("forward_pct", 0)),
        "Side %": int(candidate.get("side_pct", 0)),
        "Reciprocal": "Yes" if candidate.get("reciprocal") else "No",
        f"Altitude ({unit})": round(m_to_unit(float(candidate.get("altitude_m", 0.0)), unit), 1),
        "Oblique tilt °": round(float(candidate.get("common_oblique_tilt_deg", 0.0)), 1),
        "Oblique FL mm": round(float(candidate.get("common_oblique_focal_mm", 0.0)), 1),
        "Inner edge oblique °": round(float(candidate.get("min_oblique_angle_deg", 0.0)), 2),
        "Outer edge oblique °": round(float(candidate.get("max_oblique_angle_deg", 0.0)), 2),
        "Min views": int(candidate.get("angles_min", 0)),
        "Min hits": int(candidate.get("hits_min", 0)),
        "Inner GSD cm": round(float(candidate.get("inner_gsd_cm", 0.0)), 2),
        "Outer GSD cm": round(float(candidate.get("outer_gsd_cm", 0.0)), 2),
        "Achieved sidelap %": round(float(candidate.get("achieved_sidelap_pct", 0.0)), 1),
        f"Line spacing ({unit})": round(m_to_unit(float(candidate.get("line_spacing_m", 0.0)), unit), 1),
        f"Photo spacing ({unit})": round(m_to_unit(float(candidate.get("photo_spacing_m", 0.0)), unit), 1),
        f"Efficiency ({unit}²/pass)": round((m_to_unit(1.0, unit) ** 2) * float(candidate.get("efficiency_index_m2", 0.0)), 1),
        "Notes": candidate.get("reason", ""),
    }


def coverage_heatmap_figure(xs, ys, grid, title, colorbar_label, probe_xy=None):
    fig, ax = dark_fig(8.8, 6.2)
    extent = [xs[0], xs[-1], ys[0], ys[-1]]

    grid_int = np.asarray(grid, dtype=int)
    max_value = int(np.max(grid_int)) if grid_int.size else 0
    display_max = max(max_value, 1)

    viridis = plt.get_cmap("viridis")
    colour_steps = ["#ff3b30"]
    for i in range(1, display_max + 1):
        frac = i / max(display_max, 1)
        colour_steps.append(viridis(frac))
    cmap = mcolors.ListedColormap(colour_steps)
    boundaries = np.arange(-0.5, display_max + 1.5, 1.0)
    norm = mcolors.BoundaryNorm(boundaries, cmap.N)

    img = ax.imshow(
        grid_int,
        origin="lower",
        extent=extent,
        aspect="auto",
        interpolation="nearest",
        cmap=cmap,
        norm=norm,
    )
    cbar = fig.colorbar(img, ax=ax, shrink=0.88, ticks=np.arange(0, display_max + 1, 1))
    cbar.set_label(colorbar_label, color="#c9d1d9", fontsize=8)
    cbar.ax.yaxis.set_tick_params(color="#8b949e")
    plt.setp(cbar.ax.get_yticklabels(), color="#8b949e", fontsize=8)
    if probe_xy is not None:
        px, py = probe_xy
        ax.scatter([px], [py], s=70, marker="x", color="#f0c040", linewidths=2.0, zorder=5)
    ax.axvline(0, color="#30363d", lw=0.9, ls="--", zorder=2)
    ax.axhline(0, color="#30363d", lw=0.9, ls="--", zorder=2)
    xt = ax.get_xticks()
    ax.set_xticklabels([f"{m_to_unit(t, dist_unit):.0f}" for t in xt], color="#8b949e")
    yt = ax.get_yticks()
    ax.set_yticklabels([f"{m_to_unit(t, dist_unit):.0f}" for t in yt], color="#8b949e")
    ax.set_xlabel(f"Across-track within repeat cell ({dist_unit})", color="#8b949e")
    ax.set_ylabel(f"Along-track within repeat cell ({dist_unit})", color="#8b949e")
    ax.set_title(title, color="#c9d1d9", fontsize=11)
    fig.tight_layout()
    return fig




def _tooltip_attr_text(text):
    return html.escape(str(text), quote=True).replace("\n", "&#10;")


def metric_with_help(container, label, value, help_text):
    escaped_label = html.escape(str(label))
    escaped_value = html.escape(str(value))
    tooltip = _tooltip_attr_text(help_text)
    container.markdown(
        f"""
        <div style="border:1px solid rgba(250,250,250,0.12); border-radius:0.75rem; padding:0.8rem 0.95rem; margin-bottom:0.35rem; background:rgba(255,255,255,0.02);">
            <div style="display:flex; align-items:center; gap:0.35rem; color:#9aa4b2; font-size:0.92rem; font-weight:600; margin-bottom:0.35rem;">
                <span>{escaped_label}</span>
                <span title="{tooltip}" style="display:inline-flex; align-items:center; justify-content:center; width:1.05rem; height:1.05rem; border-radius:999px; border:1px solid #6e7681; color:#c9d1d9; font-size:0.72rem; cursor:help; line-height:1;">?</span>
            </div>
            <div style="color:#f0f6fc; font-size:1.9rem; font-weight:700; line-height:1.1;">{escaped_value}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

def help_toggle(title, text, key, level="info"):
    col_title, col_btn = st.columns([18, 1])
    with col_title:
        st.markdown(f"##### {title}")
    with col_btn:
        toggle_key = f"show_help_{key}"
        if st.button("?", key=f"btn_{key}", help=f"Show help for {title}"):
            st.session_state[toggle_key] = not st.session_state.get(toggle_key, False)
    if st.session_state.get(f"show_help_{key}", False):
        getattr(st, level)(text)


def ensure_aoi_library_dir():
    AOI_LIBRARY_DIR.mkdir(parents=True, exist_ok=True)
    return AOI_LIBRARY_DIR


def list_library_kmls():
    ensure_aoi_library_dir()
    return sorted(AOI_LIBRARY_DIR.glob("*.kml"), key=lambda p: p.name.lower())


def kml_ring_to_lonlat(coords_text):
    pts = []
    for token in str(coords_text or "").replace("\n", " ").split():
        parts = token.split(",")
        if len(parts) < 2:
            continue
        try:
            lon = float(parts[0])
            lat = float(parts[1])
        except Exception:
            continue
        pts.append((lon, lat))
    if len(pts) >= 3 and pts[0] != pts[-1]:
        pts.append(pts[0])
    return pts


def lonlat_to_local_xy(points):
    if not points:
        return [], 0.0, 0.0
    lons = [p[0] for p in points]
    lats = [p[1] for p in points]
    lon0 = float(sum(lons) / len(lons))
    lat0 = float(sum(lats) / len(lats))
    r = 6378137.0
    cos_lat = math.cos(math.radians(lat0))
    local = []
    for lon, lat in points:
        x = r * math.radians(lon - lon0) * cos_lat
        y = r * math.radians(lat - lat0)
        local.append((x, y))
    return local, lon0, lat0


def parse_kml_aoi(path):
    if not SHAPELY_AVAILABLE:
        raise RuntimeError("Shapely is required for AOI / mission outputs.")

    raw_text = None
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            raw_text = Path(path).read_text(encoding=encoding)
            break
        except Exception:
            continue
    if raw_text is None:
        raise RuntimeError(f"Could not read {Path(path).name}.")

    try:
        root = ET.fromstring(raw_text)
    except Exception as exc:
        raise RuntimeError(f"KML parse error: {exc}") from exc

    coord_texts = []
    for elem in root.iter():
        if str(elem.tag).lower().endswith("coordinates") and (elem.text or "").strip():
            coord_texts.append(elem.text)

    polygons = []
    for coords_text in coord_texts:
        lonlat_ring = kml_ring_to_lonlat(coords_text)
        if len(lonlat_ring) < 4:
            continue
        local_ring, _, _ = lonlat_to_local_xy(lonlat_ring)
        try:
            poly = ShapelyPolygon(local_ring)
            if not poly.is_valid:
                poly = poly.buffer(0)
            if not poly.is_empty and poly.area > 0:
                polygons.append(poly)
        except Exception:
            continue

    if not polygons:
        raise RuntimeError("No valid polygon coordinates were found in that KML.")

    aoi_poly = unary_union(polygons)
    if aoi_poly.is_empty:
        raise RuntimeError("The KML geometry did not produce a usable AOI polygon.")

    return {
        "name": Path(path).stem,
        "source": "kml_library",
        "polygon": aoi_poly,
        "area_m2": float(aoi_poly.area),
    }


def build_standard_aoi(area_km2=100.0):
    if not SHAPELY_AVAILABLE:
        raise RuntimeError("Shapely is required for AOI / mission outputs.")
    area_m2 = max(0.01, float(area_km2)) * 1_000_000.0
    side_m = math.sqrt(area_m2)
    half = side_m / 2.0
    poly = ShapelyPolygon([(-half, -half), (half, -half), (half, half), (-half, half)])
    return {
        "name": f"Standard {float(area_km2):.1f} km² square",
        "source": "standard_example",
        "polygon": poly,
        "area_m2": float(poly.area),
    }


def iter_line_geometries(geom):
    if geom is None or getattr(geom, "is_empty", True):
        return
    gtype = getattr(geom, "geom_type", "")
    if gtype == "LineString":
        yield geom
    elif hasattr(geom, "geoms"):
        for sub_geom in geom.geoms:
            yield from iter_line_geometries(sub_geom)


def estimate_camera_file_size_mb(cam):
    bd = get_body(cam["body"])
    megapixels = (float(bd["w_px"]) * float(bd["h_px"])) / 1_000_000.0
    return megapixels * 1.0


def compute_aoi_mission_outputs(aoi_payload, line_spacing_m, photo_spacing_m, speed_ms, enabled_cameras, flight_azimuth_deg=0.0, lead_in_out_m=150.0):
    if not SHAPELY_AVAILABLE:
        return None
    if aoi_payload is None:
        return None
    polygon = aoi_payload.get("polygon")
    if polygon is None or getattr(polygon, "is_empty", True):
        return None
    if not (np.isfinite(line_spacing_m) and line_spacing_m > 0 and np.isfinite(photo_spacing_m) and photo_spacing_m > 0):
        return None

    rotated = shapely_rotate(polygon, float(flight_azimuth_deg), origin="centroid", use_radians=False)
    minx, miny, maxx, maxy = rotated.bounds
    width = maxx - minx
    if not np.isfinite(width) or width < 0:
        return None

    n_candidates = max(1, int(math.ceil(width / line_spacing_m)) + 1)
    offsets = [minx + i * line_spacing_m for i in range(n_candidates)]
    pad = max(maxy - miny, 1.0) + max(float(lead_in_out_m), 0.0) + 1000.0

    line_count = 0
    total_line_length_m = 0.0
    total_exposures = 0
    line_lengths_m = []

    for x in offsets:
        base_line = ShapelyLineString([(x, miny - pad), (x, maxy + pad)])
        intersection = rotated.intersection(base_line)
        for seg in iter_line_geometries(intersection):
            seg_length = float(seg.length)
            if seg_length <= 0:
                continue
            mission_length = seg_length + 2.0 * max(float(lead_in_out_m), 0.0)
            line_count += 1
            total_line_length_m += mission_length
            line_lengths_m.append(mission_length)
            total_exposures += max(1, int(math.ceil(mission_length / photo_spacing_m)) + 1)

    enabled_cam_count = max(1, len(enabled_cameras))
    total_images = total_exposures * enabled_cam_count
    per_trigger_storage_mb = sum(estimate_camera_file_size_mb(cam) for cam in enabled_cameras)
    total_storage_mb = total_exposures * per_trigger_storage_mb
    flight_time_s = total_line_length_m / speed_ms if speed_ms > 0 else float("nan")

    mission_line_geometries = []
    for x in offsets:
        base_line = ShapelyLineString([(x, miny - pad), (x, maxy + pad)])
        intersection = rotated.intersection(base_line)
        for seg in iter_line_geometries(intersection):
            seg_length = float(seg.length)
            if seg_length <= 0:
                continue
            coords = list(seg.coords)
            if len(coords) < 2:
                continue
            start = coords[0]
            end = coords[-1]
            dy = end[1] - start[1]
            seg_len = math.hypot(end[0] - start[0], dy)
            if seg_len <= 0:
                continue
            extend = max(float(lead_in_out_m), 0.0)
            ux = (end[0] - start[0]) / seg_len
            uy = dy / seg_len
            ext_start = (start[0] - ux * extend, start[1] - uy * extend)
            ext_end = (end[0] + ux * extend, end[1] + uy * extend)
            try:
                ext_line = ShapelyLineString([ext_start, ext_end])
                mission_line_geometries.append(
                    shapely_rotate(ext_line, -float(flight_azimuth_deg), origin=polygon.centroid, use_radians=False)
                )
            except Exception:
                continue

    return {
        "name": aoi_payload.get("name", "AOI"),
        "source": aoi_payload.get("source", "unknown"),
        "area_m2": float(aoi_payload.get("area_m2", polygon.area)),
        "line_count": int(line_count),
        "total_line_length_m": float(total_line_length_m),
        "average_line_length_m": float(np.mean(line_lengths_m)) if line_lengths_m else 0.0,
        "photo_spacing_m": float(photo_spacing_m),
        "line_spacing_m": float(line_spacing_m),
        "trigger_events": int(total_exposures),
        "frames_per_camera": int(total_exposures),
        "total_images": int(total_images),
        "per_trigger_storage_mb": float(per_trigger_storage_mb),
        "total_storage_mb": float(total_storage_mb),
        "flight_time_s": float(flight_time_s),
        "flight_azimuth_deg": float(flight_azimuth_deg),
        "lead_in_out_m": float(lead_in_out_m),
        "mission_line_geometries": mission_line_geometries,
        "aoi_polygon": polygon,
    }


def make_aoi_mission_figure(mission_outputs, dist_unit="m"):
    if mission_outputs is None:
        return None
    polygon = mission_outputs.get("aoi_polygon")
    if polygon is None or getattr(polygon, "is_empty", True):
        return None

    fig, ax = dark_fig(4.2, 3.35)
    ax.set_aspect("equal")

    all_x = []
    all_y = []
    try:
        minx, miny, maxx, maxy = polygon.bounds
        all_x.extend([minx, maxx])
        all_y.extend([miny, maxy])
    except Exception:
        pass

    for line in mission_outputs.get("mission_line_geometries", []):
        try:
            x, y = line.xy
        except Exception:
            continue
        all_x.extend(list(x))
        all_y.extend(list(y))

    if all_x and all_y:
        x_min, x_max = min(all_x), max(all_x)
        y_min, y_max = min(all_y), max(all_y)
        x_span = max(x_max - x_min, 1.0)
        y_span = max(y_max - y_min, 1.0)
        x_mid = 0.5 * (x_min + x_max)
        y_mid = 0.5 * (y_min + y_max)
        half_span = 0.5 * max(x_span, y_span)
        pad = max(half_span * 0.12, 1.0)
        x_lo, x_hi = x_mid - half_span - pad, x_mid + half_span + pad
        y_lo, y_hi = y_mid - half_span - pad, y_mid + half_span + pad
    else:
        x_lo, x_hi, y_lo, y_hi = -1.0, 1.0, -1.0, 1.0

    display_in_km = dist_unit == "m"
    scale = 1000.0 if display_in_km else 1.0
    axis_unit = "km" if display_in_km else dist_unit

    def _sx(vals):
        return [v / scale for v in vals]

    def _plot_poly(poly, edgecolor="#58a6ff", facealpha=0.10, lw=2.0):
        geoms = [poly] if getattr(poly, "geom_type", "") == "Polygon" else list(getattr(poly, "geoms", []))
        for geom in geoms:
            if getattr(geom, "is_empty", True):
                continue
            x, y = geom.exterior.xy
            ax.fill(_sx(x), _sx(y), color=edgecolor, alpha=facealpha, zorder=2)
            ax.plot(_sx(x), _sx(y), color=edgecolor, lw=lw, zorder=3)
            for interior in getattr(geom, "interiors", []):
                ix, iy = interior.xy
                ax.plot(_sx(ix), _sx(iy), color="#30363d", lw=1.0, zorder=2)

    _plot_poly(polygon)

    for line in mission_outputs.get("mission_line_geometries", []):
        try:
            x, y = line.xy
        except Exception:
            continue
        ax.plot(_sx(x), _sx(y), color="#f0c040", lw=1.5, alpha=0.95, zorder=4)

    ax.set_xlim(x_lo / scale, x_hi / scale)
    ax.set_ylim(y_lo / scale, y_hi / scale)

    try:
        from matplotlib.ticker import MaxNLocator
        ax.xaxis.set_major_locator(MaxNLocator(nbins=5, prune=None))
        ax.yaxis.set_major_locator(MaxNLocator(nbins=6, prune=None))
    except Exception:
        pass

    xt = ax.get_xticks()
    ax.set_xticklabels([f"{t:.1f}" if display_in_km else f"{t:.0f}" for t in xt], color="#8b949e")
    yt = ax.get_yticks()
    ax.set_yticklabels([f"{t:.1f}" if display_in_km else f"{t:.0f}" for t in yt], color="#8b949e")
    ax.set_xlabel(f"Local easting ({axis_unit})", color="#8b949e")
    ax.set_ylabel(f"Local northing ({axis_unit})", color="#8b949e")
    ax.set_title("AOI and generated flight lines", color="#c9d1d9", fontsize=11)
    fig.tight_layout()
    return fig


# Apply any scenario requested on the previous run before widgets are created.
pending_loaded_scenario = st.session_state.get("pending_loaded_scenario")
if pending_loaded_scenario:
    sc = pending_loaded_scenario.get("data") or {}
    scenario_label = pending_loaded_scenario.get("label", "scenario")
    if sc and "cameras" in sc:
        loaded_altitude_m = float(sc.get("altitude_m", DEFAULT_ALTITUDE_M))
        current_dist_unit = st.session_state.get("dist_unit", st.session_state.get("planner_dist_unit_last", "m"))
        st.session_state.cameras = sc["cameras"]
        bump_camera_widget_nonce()
        st.session_state.planner_altitude_m = loaded_altitude_m
        st.session_state.sidebar_alt_input = round(m_to_unit(loaded_altitude_m, current_dist_unit), 1)
        st.session_state.sidebar_speed_ms = float(sc.get("speed_ms", DEFAULT_SPEED_MS))
        st.session_state.sidebar_reciprocal = bool(sc.get("reciprocal", True))
        st.session_state.fwd_pct = int(sc.get("fwd_overlap_pct", sc.get("forward_overlap_pct", OVERLAP_PRESETS["Balanced"]["forward"])))
        st.session_state.side_pct = int(sc.get("sidelap_pct", OVERLAP_PRESETS["Balanced"]["side"]))
        st.session_state.overlap_preset = "Custom"
        st.session_state.overlap_preset_last = "Custom"
        st.session_state.selected_scenario_label = scenario_label
        st.session_state.saved_scenario_selector = scenario_label
        st.session_state.saved_scenario_selector_prev = scenario_label
        st.session_state.scenario_flash_message = {
            "level": "success",
            "text": f"Loaded {scenario_label}",
        }
        st.session_state.pop("optimizer_results", None)
    st.session_state.pending_loaded_scenario = None

pending_optimizer_apply = st.session_state.get("pending_optimizer_apply")
if pending_optimizer_apply:
    candidate = pending_optimizer_apply
    current_dist_unit = st.session_state.get("dist_unit", st.session_state.get("planner_dist_unit_last", "m"))
    updated_cameras = [dict(c) for c in st.session_state.cameras]
    apply_common_oblique_settings(
        updated_cameras,
        tilt_deg=candidate.get("common_oblique_tilt_deg"),
        focal_mm=candidate.get("common_oblique_focal_mm"),
    )
    st.session_state.cameras = updated_cameras
    bump_camera_widget_nonce()
    st.session_state.planner_altitude_m = float(candidate.get("altitude_m", st.session_state.planner_altitude_m))
    st.session_state.sidebar_alt_input = round(m_to_unit(float(candidate.get("altitude_m", st.session_state.planner_altitude_m)), current_dist_unit), 1)
    st.session_state.sidebar_speed_ms = float(candidate.get("speed_ms", st.session_state.sidebar_speed_ms))
    st.session_state.sidebar_reciprocal = bool(candidate.get("reciprocal", st.session_state.sidebar_reciprocal))
    st.session_state.fwd_pct = int(candidate.get("forward_pct", st.session_state.get("fwd_pct", OVERLAP_PRESETS["Balanced"]["forward"])))
    st.session_state.side_pct = int(candidate.get("side_pct", st.session_state.get("side_pct", OVERLAP_PRESETS["Balanced"]["side"])))
    st.session_state.overlap_preset = "Custom"
    st.session_state.overlap_preset_last = "Custom"
    st.session_state.scenario_flash_message = {
        "level": "success",
        "text": "Applied optimiser recommendation to planner",
    }
    st.session_state.pending_optimizer_apply = None

# ─────────────────────────────────────────────────────────────────────────────
# Sidebar
# ─────────────────────────────────────────────────────────────────────────────

st.sidebar.image("ASL_Imagery_Icon.png", width=85)
st.sidebar.title("Oblique Survey Planner\nMatched Pair")
st.sidebar.markdown("---")
st.sidebar.subheader("Flight Parameters")

dist_unit = st.sidebar.selectbox("Distance unit", ["m", "ft"], index=0, key="dist_unit")

if st.session_state.planner_dist_unit_last != dist_unit:
    st.session_state.sidebar_alt_input = round(m_to_unit(st.session_state.planner_altitude_m, dist_unit), 1)
    st.session_state.planner_dist_unit_last = dist_unit

alt_input = st.sidebar.number_input(
    f"Altitude AGL ({dist_unit})",
    value=float(st.session_state.sidebar_alt_input),
    min_value=1.0, max_value=m_to_unit(10000.0, dist_unit),
    step=max(1.0, m_to_unit(10.0, dist_unit)),
    key="sidebar_alt_input",
)
altitude_m = unit_to_m(alt_input, dist_unit)
st.session_state.planner_altitude_m = altitude_m

speed_ms = st.sidebar.number_input(
    "Aircraft speed (m/s)",
    value=float(st.session_state.sidebar_speed_ms),
    min_value=1.0,
    max_value=300.0,
    step=0.5,
    key="sidebar_speed_ms",
)
st.sidebar.caption(f"≈ {speed_ms * 1.94384:.1f} kts  |  {speed_ms * 3.6:.1f} km/h")
reciprocal = st.sidebar.checkbox("Reciprocal (bidirectional) strips", value=bool(st.session_state.sidebar_reciprocal), key="sidebar_reciprocal")

st.sidebar.markdown("---")
st.sidebar.subheader("Overlap Targets")

if "overlap_preset" not in st.session_state:
    st.session_state.overlap_preset = "Balanced"
if "overlap_preset_last" not in st.session_state:
    st.session_state.overlap_preset_last = st.session_state.overlap_preset
if "fwd_pct" not in st.session_state:
    st.session_state.fwd_pct = OVERLAP_PRESETS["Balanced"]["forward"]
if "side_pct" not in st.session_state:
    st.session_state.side_pct = OVERLAP_PRESETS["Balanced"]["side"]

selected_overlap_preset = st.sidebar.selectbox(
    "Overlap preset",
    list(OVERLAP_PRESETS.keys()),
    key="overlap_preset",
)

if selected_overlap_preset != st.session_state.overlap_preset_last:
    preset_values = OVERLAP_PRESETS[selected_overlap_preset]
    if preset_values is not None:
        st.session_state.fwd_pct = preset_values["forward"]
        st.session_state.side_pct = preset_values["side"]
    st.session_state.overlap_preset_last = selected_overlap_preset

fwd_pct  = st.sidebar.slider("Forward overlap (%)", 0, 95, key="fwd_pct")
side_pct = st.sidebar.slider("Sidelap (%)", 0, 95, key="side_pct")
fwd_frac  = fwd_pct  / 100.0
side_frac = side_pct / 100.0

st.sidebar.markdown("---")
st.sidebar.subheader("💾 Save / Load")

if "scenario_flash_message" in st.session_state:
    flash = st.session_state.pop("scenario_flash_message")
    flash_level = flash.get("level", "success")
    getattr(st.sidebar, flash_level, st.sidebar.success)(flash.get("text", "Scenario updated."))

sc_name = st.sidebar.text_input("Scenario name", "my_survey")
if st.sidebar.button("Save scenario"):
    saved_path = save_scenario({
        "cameras": st.session_state.cameras,
        "altitude_m": altitude_m,
        "speed_ms": speed_ms,
        "fwd_overlap_pct": fwd_pct,
        "sidelap_pct": side_pct,
        "reciprocal": reciprocal,
    }, sc_name)
    st.session_state.scenario_flash_message = {
        "level": "success",
        "text": f"Saved {saved_path.name} to {saved_path.parent.as_posix()}",
    }
    st.session_state.selected_scenario_label = saved_path.name
    st.session_state.saved_scenario_selector = saved_path.name
    st.session_state.saved_scenario_selector_prev = saved_path.name
    st.rerun()

scenario_records = list_saved_scenarios()
scenario_labels = [record["label"] for record in scenario_records]
scenario_lookup = {record["label"]: record["path"] for record in scenario_records}

if scenario_records:
    default_index = 0
    selected_label = st.session_state.get("selected_scenario_label")
    if selected_label in scenario_labels:
        default_index = scenario_labels.index(selected_label)

    if st.session_state.saved_scenario_selector not in scenario_labels:
        st.session_state.saved_scenario_selector = scenario_labels[default_index]

    load_name = st.sidebar.selectbox(
        "Saved scenarios",
        scenario_labels,
        index=scenario_labels.index(st.session_state.saved_scenario_selector),
        key="saved_scenario_selector",
    )
    selected_record = next((record for record in scenario_records if record["label"] == load_name), None)
    if selected_record and selected_record["origin"] == "project_root":
        st.sidebar.caption("Legacy scenario found in the project root. Saving again will place it into saved_scenarios.")

    if st.session_state.saved_scenario_selector_prev is None:
        st.session_state.saved_scenario_selector_prev = load_name
    elif load_name != st.session_state.saved_scenario_selector_prev:
        sc = load_scenario(scenario_lookup.get(load_name, load_name))
        st.session_state.saved_scenario_selector_prev = load_name
        if sc and "cameras" in sc:
            st.session_state.pending_loaded_scenario = {
                "label": load_name,
                "data": sc,
            }
            st.rerun()
        else:
            st.sidebar.error(f"Could not load '{load_name}'")

    st.sidebar.caption("Selecting a saved scenario loads it automatically.")
else:
    load_name = None
    st.sidebar.info("No saved scenarios found in saved_scenarios yet.")

# ─────────────────────────────────────────────────────────────────────────────
# Main header
# ─────────────────────────────────────────────────────────────────────────────
col_left_logo, col_left, col_spacer, col_right = st.columns([1.2, 5.5, 1.3, 2.0])

with col_left_logo:
    st.markdown("<div style='padding-top: 22px;'></div>", unsafe_allow_html=True)
    st.image("ASL_Imagery_Icon.png", width=110)

with col_left:
    st.title("Aerial Surveys Oblique Planner")
    st.caption(
        f"Altitude: **{fmt(altitude_m, dist_unit)}**  |  "
        f"Speed: **{speed_ms:.0f} m/s ({speed_ms * 1.94384:.0f} kts)**  |  "
        f"Fwd overlap: **{fwd_pct}%**  |  Sidelap: **{side_pct}%**"
    )

with col_right:
    st.markdown("<div style='padding-top: 18px;'></div>", unsafe_allow_html=True)
    st.image("ASL_Logo_White.png", width=200)
# ─────────────────────────────────────────────────────────────────────────────
# Camera configuration table
# ─────────────────────────────────────────────────────────────────────────────

st.subheader("Camera Configuration")

with st.expander("Orientation & tilt-axis reference", expanded=False):
    st.markdown("""
| Setting | Option | Sensor across-track | Best for |
|---|---|---|---|
| **Orientation** | Portrait | **Narrow** (short) axis | L/R oblique — limits far-edge GSD stretch |
| | Landscape | **Long** axis | Nadir — maximises swath width |
| **Tilt axis** | Across (L/R) | Tilts left/right about along-track axis | Left & Right oblique |
| | Along (F/A) | Tilts fore/aft about across-track axis | Fore & Aft oblique |

**Spreadsheet match:** The Landscape sheet in the reference spreadsheet uses portrait-mounted
L/R cameras (narrow axis across-track). The default settings here reproduce those values exactly.
For the Left camera, enter the same tilt angle as Right — the mirroring is handled automatically.
    """)

cameras      = st.session_state.cameras
bodies_avail = all_body_names()

# Column headers
hdr = st.columns([0.35, 1.5, 1.7, 0.7, 0.65, 0.8, 0.85, 1.0, 0.3])
for col, lbl in zip(hdr, ["✓", "Label", "Body", "FL mm", "Tilt °", "From", "Orient.", "Tilt axis", ""]):
    col.markdown(f"<span style='color:#8b949e;font-size:0.78em;font-weight:600'>{lbl}</span>",
                 unsafe_allow_html=True)

to_delete = None
camera_widget_nonce = st.session_state.get("camera_widget_nonce", 0)
for i, cam in enumerate(cameras):
    c_en, c_lbl, c_body, c_fl, c_tilt, c_conv, c_orient, c_axis, c_del = \
        st.columns([0.35, 1.5, 1.7, 0.7, 0.65, 0.8, 0.85, 1.0, 0.3])
    colour = CAM_COLOURS[i % len(CAM_COLOURS)]

    cam["enabled"] = c_en.checkbox("##e", value=cam["enabled"], key=f"en_{camera_widget_nonce}_{i}",
                                    label_visibility="collapsed")
    c_lbl.markdown(f"<span style='color:{colour}'>●</span>", unsafe_allow_html=True)
    cam["label"] = c_lbl.text_input("##l", value=cam["label"], key=f"lbl_{camera_widget_nonce}_{i}",
                                     label_visibility="collapsed")
    body_idx = bodies_avail.index(cam["body"]) if cam["body"] in bodies_avail else 0
    cam["body"] = c_body.selectbox("##b", bodies_avail, index=body_idx, key=f"body_{camera_widget_nonce}_{i}",
                                    label_visibility="collapsed")
    cam["focal_mm"] = float(c_fl.number_input("##f", value=float(cam["focal_mm"]),
                                               min_value=1.0, max_value=2000.0, step=1.0,
                                               key=f"fl_{camera_widget_nonce}_{i}", label_visibility="collapsed"))
    cam["tilt_deg"] = float(c_tilt.number_input("##t", value=float(cam["tilt_deg"]),
                                                  min_value=0.0, max_value=85.0, step=0.5,
                                                  key=f"tilt_{camera_widget_nonce}_{i}", label_visibility="collapsed"))
    cam["tilt_conv"] = c_conv.selectbox("##c", ["horiz", "nadir"], key=f"conv_{camera_widget_nonce}_{i}",
                                         index=0 if cam["tilt_conv"] == "horiz" else 1,
                                         label_visibility="collapsed",
                                         format_func=lambda x: "Horizontal" if x == "horiz" else "Nadir")
    cam["orientation"] = c_orient.selectbox("##o", ["portrait", "landscape"], key=f"orient_{camera_widget_nonce}_{i}",
                                             index=0 if cam["orientation"] == "portrait" else 1,
                                             label_visibility="collapsed",
                                             format_func=lambda x: "Portrait" if x == "portrait" else "Landscape")
    cam["tilt_axis"] = c_axis.selectbox("##a", ["across", "along"], key=f"axis_{camera_widget_nonce}_{i}",
                                         index=0 if cam["tilt_axis"] == "across" else 1,
                                         label_visibility="collapsed",
                                         format_func=lambda x: "Across (L/R)" if x == "across" else "Along (F/A)")
    if c_del.button("✕", key=f"del_{camera_widget_nonce}_{i}", help="Remove camera"):
        to_delete = i

if to_delete is not None:
    cameras.pop(to_delete)
    st.session_state.pop("optimizer_results", None)
    bump_camera_widget_nonce()
    st.rerun()

b1, b2, b3 = st.columns([1, 1, 6])
if b1.button("➕ Add camera"):
    cameras.append({"enabled": True, "label": f"Camera {len(cameras)+1}",
                    "body": "Custom body", "focal_mm": 50.0, "tilt_deg": 50.0,
                    "tilt_conv": "horiz", "orientation": "portrait", "tilt_axis": "across"})
    st.session_state.pop("optimizer_results", None)
    bump_camera_widget_nonce()
    st.rerun()
if b2.button("↺ Reset defaults"):
    st.session_state.cameras = [dict(c) for c in DEFAULT_CAMERAS]
    st.session_state.pop("optimizer_results", None)
    bump_camera_widget_nonce()
    st.rerun()

with st.expander("Save a camera body as a custom preset"):
    pc1, pc2 = st.columns([2, 1])
    p_base = pc1.selectbox("Base body", list(BODY_PRESETS.keys()), key="p_base")
    p_name = pc2.text_input("Preset name", value=p_base, key="p_name")
    p_wmm  = pc1.number_input("Sensor long axis mm", value=BODY_PRESETS[p_base]["w_mm"], key="p_wmm")
    p_hmm  = pc1.number_input("Sensor short axis mm", value=BODY_PRESETS[p_base]["h_mm"], key="p_hmm")
    p_wpx  = pc1.number_input("Pixel count long axis", value=BODY_PRESETS[p_base]["w_px"], step=100, key="p_wpx")
    p_hpx  = pc1.number_input("Pixel count short axis", value=BODY_PRESETS[p_base]["h_px"], step=100, key="p_hpx")
    if pc2.button("Save preset"):
        save_body_preset(p_name, {"w_mm": p_wmm, "h_mm": p_hmm,
                                   "w_px": int(p_wpx), "h_px": int(p_hpx)})
        st.success(f"Saved '{p_name}'. Reload page to use it.")

st.markdown("---")

# ─────────────────────────────────────────────────────────────────────────────
# Build camera solutions
# ─────────────────────────────────────────────────────────────────────────────

active = [c for c in cameras if c["enabled"]]

solutions = []   # list of (cam_dict, CameraSolution, colour_str)
errors    = []

for i, cam in enumerate(active):
    try:
        bd     = get_body(cam["body"])
        tilt_n = normalize_tilt_angle(cam["tilt_deg"], cam["tilt_conv"])
        sol    = calculate_camera_solution(
            altitude_m          = altitude_m,
            tilt_from_nadir_deg = tilt_n,
            sensor_w_native_mm  = bd["w_mm"],
            sensor_h_native_mm  = bd["h_mm"],
            image_w_native_px   = bd["w_px"],
            image_h_native_px   = bd["h_px"],
            focal_length_mm     = cam["focal_mm"],
            orientation         = cam["orientation"],
            tilt_axis           = cam["tilt_axis"],
            label               = cam["label"],
        )
        sol = mirror_solution_for_label(sol)
        solutions.append((cam, sol, CAM_COLOURS[i % len(CAM_COLOURS)]))
    except Exception as e:
        errors.append(f"**{cam['label']}**: {e}")

for e in errors:
    st.error(e)

if not solutions:
    st.warning("No cameras enabled. Enable at least one camera in the table above.")
    st.stop()

sol_list = [s for _, s, _ in solutions]

mc = None
try:
    mc = calculate_multicamera_solution(
        camera_solutions      = sol_list,
        arrangement           = "custom",
        altitude_m            = altitude_m,
        aircraft_speed_ms     = speed_ms,
        forward_overlap_fraction = fwd_frac,
        sidelap_fraction      = side_frac,
        reciprocal_flying     = reciprocal,
    )
except Exception as e:
    st.error(f"System calculation error: {e}")

if mc and mc.warnings:
    with st.expander("⚠️ Warnings", expanded=True):
        for w in mc.warnings:
            st.warning(w)

# ─────────────────────────────────────────────────────────────────────────────
# Summary cards
# ─────────────────────────────────────────────────────────────────────────────

st.subheader("System Summary")
if mc:
    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("Combined swath",    fmt(mc.combined_swath_m,             dist_unit))
    c2.metric("Line spacing",      fmt(mc.recommended_line_spacing_m,  dist_unit))
    c3.metric("Photo spacing",     fmt(mc.recommended_photo_spacing_m, dist_unit))
    c4.metric("Exposure interval", f"{mc.photo_interval_s:.2f} s")
    c5.metric("Sidelap achieved",  f"{mc.sidelap_achieved * 100:.1f}%")

    rep = next((s for _, s, _ in solutions if abs(s.tilt_from_nadir_deg) > 1), sol_list[0])
    c6, c7, c8, c9 = st.columns(4)
    c6.metric(f"GSD inner ({rep.label})", fmt_gsd(min(rep.near_gsd_m, rep.far_gsd_m)))
    c7.metric("GSD centre",               fmt_gsd(rep.centre_gsd_m))
    c8.metric("GSD outer",                fmt_gsd(max(rep.near_gsd_m, rep.far_gsd_m)))
    c9.metric("Fwd overlap near/ctr/far",
              f"{mc.forward_overlap_near*100:.0f}% / "
              f"{mc.forward_overlap_centre*100:.0f}% / "
              f"{mc.forward_overlap_far*100:.0f}%")

mission_outputs = None
aoi_fig = None
st.markdown("---")
st.subheader("AOI / Mission Outputs")
help_toggle(
    "Area-based outputs",
    "This section turns the current line spacing and photo spacing into job-level outputs. You can use the built-in 100 km² example for like-for-like comparisons, or load a KML from the aoi_library folder on the server. Outputs are illustrative planning numbers based on clipped parallel lines across the AOI, the current photo spacing, enabled camera count and a simple RAW storage estimate of about 1 MB per megapixel per image.",
    key="aoi_outputs_intro",
)

if not SHAPELY_AVAILABLE:
    st.info("AOI / mission outputs require Shapely to be installed in the app environment.")
elif not mc:
    st.info("AOI / mission outputs are unavailable until the system spacing has been calculated.")
else:
    aoi_mode = st.radio(
        "AOI source",
        ["Standard example", "Load KML from library"],
        horizontal=True,
        key="aoi_source_mode",
    )

    aoi_payload = None
    controls_left, controls_right = st.columns([2.2, 1.4])

    with controls_left:
        if aoi_mode == "Standard example":
            example_area_km2 = st.number_input(
                "Example AOI area (km²)",
                min_value=1.0,
                max_value=5000.0,
                value=100.0,
                step=10.0,
                key="example_aoi_area_km2",
            )
            aoi_payload = build_standard_aoi(example_area_km2)
            st.session_state.selected_aoi_payload = aoi_payload
            st.session_state.selected_aoi_name = aoi_payload["name"]
            st.caption("Standard example uses a square AOI so you can compare option changes on a common 100 km²-style block.")
        else:
            kml_files = list_library_kmls()
            if not kml_files:
                st.info("No KML files found yet. Add .kml files to the aoi_library folder on the server to use this option.")
            else:
                selected_kml = st.selectbox(
                    "KML in aoi_library",
                    [p.name for p in kml_files],
                    key="selected_kml_name",
                )
                if st.button("Load KML", key="load_kml_button"):
                    selected_path = next((p for p in kml_files if p.name == selected_kml), None)
                    try:
                        st.session_state.selected_aoi_payload = parse_kml_aoi(selected_path)
                        st.session_state.selected_aoi_name = selected_kml
                        st.success(f"Loaded {selected_kml}")
                    except Exception as exc:
                        st.session_state.selected_aoi_payload = None
                        st.error(str(exc))
                if st.session_state.get("selected_aoi_payload") is not None:
                    aoi_payload = st.session_state.selected_aoi_payload
                    st.caption(f"Loaded AOI: {st.session_state.get('selected_aoi_name', aoi_payload.get('name', 'KML'))}")

    with controls_right:
        flight_azimuth_deg = st.number_input(
            "Flight azimuth (deg)",
            min_value=0.0,
            max_value=359.9,
            value=0.0,
            step=1.0,
            key="aoi_flight_azimuth_deg",
            help="Flight-line direction measured clockwise from north. 0° = north-south lines, 90° = east-west lines.",
        )
        lead_in_out_display = st.number_input(
            f"Lead-in / out per line ({dist_unit})",
            min_value=0.0,
            max_value=m_to_unit(5000.0, dist_unit),
            value=m_to_unit(150.0, dist_unit),
            step=max(1.0, m_to_unit(25.0, dist_unit)),
            key="aoi_lead_in_out_display",
        )
        lead_in_out_m = unit_to_m(lead_in_out_display, dist_unit)

    if aoi_payload is not None:
        mission_outputs = compute_aoi_mission_outputs(
            aoi_payload=aoi_payload,
            line_spacing_m=mc.recommended_line_spacing_m,
            photo_spacing_m=mc.recommended_photo_spacing_m,
            speed_ms=speed_ms,
            enabled_cameras=active,
            flight_azimuth_deg=flight_azimuth_deg,
            lead_in_out_m=lead_in_out_m,
        )

    if mission_outputs is not None:
        m1, m2, m3, m4 = st.columns(4)
        m1.metric("AOI area", f"{mission_outputs['area_m2'] / 1_000_000.0:.1f} km²")
        m2.metric("Flight lines", f"{mission_outputs['line_count']}")
        m3.metric("Trigger events", f"{mission_outputs['trigger_events']}")
        m4.metric("Total images", f"{mission_outputs['total_images']:,}")

        m5, m6, m7, m8 = st.columns(4)
        m5.metric("Total line length (km)", f"{mission_outputs['total_line_length_m'] / 1000.0:,.2f}")
        m6.metric("Avg line length (km)", f"{mission_outputs['average_line_length_m'] / 1000.0:,.2f}")
        m7.metric("Estimated storage", f"{mission_outputs['total_storage_mb'] / 1024.0:.1f} GB")
        m8.metric("Estimated flying time", f"{mission_outputs['flight_time_s'] / 60.0:.1f} min")

        st.caption(
            f"Frames per camera = {mission_outputs['frames_per_camera']}. Total images assumes {len(active)} enabled camera(s) fire at each trigger event. "
            f"Storage is a rough estimate using about 1 MB per megapixel per image. Current mission spacing is {fmt(mission_outputs['line_spacing_m'], dist_unit)} by {fmt(mission_outputs['photo_spacing_m'], dist_unit)} with {m_to_unit(mission_outputs['lead_in_out_m'], dist_unit):.0f} {dist_unit} lead-in / out per line."
        )

        aoi_fig = make_aoi_mission_figure(mission_outputs, dist_unit=dist_unit)
        if aoi_fig is not None:
            aoi_left, aoi_mid, aoi_right = st.columns([1.2, 2.6, 1.2])
            with aoi_mid:
                st.pyplot(aoi_fig)
    else:
        st.info("Select the standard example or load a KML from the library to see lines, frames, images and storage estimates.")

st.markdown("---")

# ─────────────────────────────────────────────────────────────────────────────
# Per-camera results table  (includes obliqueness ratio)
# ─────────────────────────────────────────────────────────────────────────────

st.subheader("Per-Camera Results")

rows = []
for cam, sol, _ in solutions:
    inner_gx, outer_gx = corner_inner_outer(sol)
    inner_len, outer_len = along_lengths_for_display(sol)
    inner_gsd = min(sol.near_gsd_m, sol.far_gsd_m)
    outer_gsd = max(sol.near_gsd_m, sol.far_gsd_m)
    tilt_h = 90.0 - sol.tilt_from_nadir_deg

    rows.append({
    "Camera":                    sol.label,
    "Body / FL":                 f"{cam['body']}  {cam['focal_mm']:.0f} mm",
    "Orient.":                   sol.orientation,
    "Tilt axis":                 sol.tilt_axis,
    "Tilt nadir °":              round(sol.tilt_from_nadir_deg, 1),
    "Tilt horiz °":              round(tilt_h, 1),

    # ── Oblique angles ──
    "Inner edge oblique °":      round(sol.near_angle_deg, 2),
    "Outer edge oblique °":      round(sol.far_angle_deg, 2),
    "Near oblique angle":        format_oblique(sol.near_angle_deg),
    "Far oblique angle":         format_oblique(sol.far_angle_deg),
    # ─────────────────────

    "Pixel µm":                  round(sol.pixel_size_mm * 1000, 2),
    "FOV across °":              round(sol.full_fov_across_deg, 2),
    "FOV along °":               round(sol.full_fov_along_deg,  2),
    f"Inner edge ({dist_unit})": round(m_to_unit(abs(inner_gx), dist_unit), 1),
    f"Outer edge ({dist_unit})": round(m_to_unit(abs(outer_gx), dist_unit), 1),
    f"Inner length ({dist_unit})": round(m_to_unit(inner_len, dist_unit), 1),
    f"Outer length ({dist_unit})": round(m_to_unit(outer_len, dist_unit), 1),
    "Inner GSD cm":              round(inner_gsd * 100, 3),
    "Centre GSD cm":             round(sol.centre_gsd_m * 100, 3),
    "Outer GSD cm":              round(outer_gsd * 100, 3),
    "Inner slant m":             round(min(sol.near_slant_m, sol.far_slant_m), 1),
    "Outer slant m":             round(max(sol.near_slant_m, sol.far_slant_m), 1),
    "Diag image mm":             round(sol.diag_image_mm, 4),
})

st.dataframe(rows, width="stretch")
st.caption(
    "**Near/Far oblique angle** = angle from vertical (nadir) at the inner and outer image edges. "
    "Higher angles mean a stronger oblique view."
)
st.markdown("---")

settings_rows, system_rows, camera_rows = build_export_data(
    solutions=solutions,
    mc=mc,
    altitude_m=altitude_m,
    speed_ms=speed_ms,
    fwd_frac=fwd_frac,
    side_frac=side_frac,
    dist_unit=dist_unit,
    reciprocal=reciprocal,
)

st.markdown("---")

# then your first diagram section starts below here

# ─────────────────────────────────────────────────────────────────────────────
# DIAGRAM 1 — Footprint plan view — ALL cameras, ALL frames shown
# ─────────────────────────────────────────────────────────────────────────────

st.subheader("Footprint Plan View — All Cameras")
st.caption(
    "**Near/Far oblique angle** = angle from vertical (nadir) at the inner and outer image edges. "
    "Higher angles mean a stronger oblique view."
)

lim = axis_limits_from_solutions(solutions)

fig_fp, ax_fp = dark_fig(w=10, h=10)
ax_fp.set_aspect("equal")
ax_fp.set_xlim(-lim, lim)
ax_fp.set_ylim(-lim, lim)

# Grid / nadir
ax_fp.axhline(0, color="#21262d", lw=0.8, zorder=1)
ax_fp.axvline(0, color="#21262d", lw=0.8, zorder=1)
ax_fp.scatter([0], [0], s=200, color="#f0c040", zorder=8, marker="x", linewidths=2.5)
ax_fp.annotate("Nadir", (0, 0), xytext=(8, -14), textcoords="offset points",
               color="#f0c040", fontsize=8.5, fontweight="bold")

for cam, sol, colour in solutions:
    corners_xy = camera_polygon(sol)
    if len(corners_xy) != 4:
        continue

    poly = plt.Polygon(
        corners_xy,
        closed=True,
        facecolor=colour,
        alpha=0.18,
        edgecolor=colour,
        linewidth=2.2,
        zorder=3,
    )
    ax_fp.add_patch(poly)
    ax_fp.scatter(*zip(*corners_xy), color=colour, s=30, zorder=5, edgecolors="none")

    cx = sum(c[0] for c in corners_xy) / 4
    cy = sum(c[1] for c in corners_xy) / 4
    ax_fp.text(
        cx,
        cy,
        sol.label,
        color=colour,
        fontsize=8,
        fontweight="bold",
        ha="center",
        va="center",
        zorder=7,
        bbox=dict(facecolor="#0d1117", alpha=0.70, pad=2.5, edgecolor=colour, linewidth=0.8, boxstyle="round,pad=0.3"),
    )

    inner_gx, outer_gx = corner_inner_outer(sol)
    inner_len, outer_len = along_lengths_for_display(sol)
    (it, ib), (ot, ob) = inner_outer_corners(sol)
    label_box = dict(facecolor="#0d1117", alpha=0.82, edgecolor=colour, linewidth=0.6, boxstyle="round,pad=0.18")

    if sol.tilt_axis == "across":
        ap = dict(arrowstyle="-|>", lw=1.3, mutation_scale=11)
        ax_fp.annotate("", xy=(inner_gx, 0), xytext=(0, 0), arrowprops={**ap, "color": colour}, zorder=4)
        ax_fp.annotate("", xy=(outer_gx, 0), xytext=(inner_gx, 0), arrowprops={**ap, "color": colour}, zorder=4)

        outer_tip_dx = -16 if outer_gx >= 0 else 16
        ax_fp.annotate(
            f"Inner offset\n{m_to_unit(abs(inner_gx), dist_unit):.0f} {dist_unit}",
            xy=(inner_gx / 2, 0),
            xytext=(0, 12),
            textcoords="offset points",
            color=colour,
            fontsize=6.8,
            ha="center",
            va="bottom",
            bbox=label_box,
        )
        ax_fp.annotate(
            f"Outer offset\n{m_to_unit(abs(outer_gx), dist_unit):.0f} {dist_unit}",
            xy=(outer_gx, 0),
            xytext=(outer_tip_dx, 12),
            textcoords="offset points",
            color=colour,
            fontsize=6.8,
            ha="right" if outer_gx >= 0 else "left",
            va="bottom",
            bbox=label_box,
        )

        ixm = (it[0] + ib[0]) / 2
        iym = (it[1] + ib[1]) / 2
        inner_top_y = max(it[1], ib[1])
        inner_label_y = inner_top_y * 0.5
        ax_fp.text(
            ixm,
            inner_label_y,
            f"Inner edge length\n{m_to_unit(inner_len, dist_unit):.0f} {dist_unit}",
            color=colour,
            fontsize=6.2,
            ha="center",
            va="center",
            bbox=label_box,
            zorder=7,
        )

        oxm = (ot[0] + ob[0]) / 2
        oym = (ot[1] + ob[1]) / 2
        outer_top_y = max(ot[1], ob[1])
        outer_label_y = outer_top_y * 0.5
        ax_fp.text(
            oxm,
            outer_label_y,
            f"Outer edge length\n{m_to_unit(outer_len, dist_unit):.0f} {dist_unit}",
            color=colour,
            fontsize=6.2,
            ha="center",
            va="center",
            bbox=label_box,
            zorder=7,
        )

        inner_gsd = min(sol.near_gsd_m, sol.far_gsd_m)
        outer_gsd = max(sol.near_gsd_m, sol.far_gsd_m)
        ax_fp.annotate(
            f"{inner_gsd*100:.2f} cm/px",
            xy=(ixm, iym),
            xytext=(0, 12),
            textcoords="offset points",
            color=colour,
            fontsize=5.5,
            alpha=0.9,
            ha="center",
            va="bottom",
        )
        ax_fp.annotate(
            f"{outer_gsd*100:.2f} cm/px",
            xy=(oxm, oym),
            xytext=(0, 12),
            textcoords="offset points",
            color=colour,
            fontsize=5.5,
            alpha=0.9,
            ha="center",
            va="bottom",
        )

    else:
        ap = dict(arrowstyle="-|>", lw=1.3, mutation_scale=11)
        inner_gy = sol.near_edge_m if abs(sol.near_edge_m) < abs(sol.far_edge_m) else sol.far_edge_m
        outer_gy = sol.far_edge_m if abs(sol.far_edge_m) > abs(sol.near_edge_m) else sol.near_edge_m
        if not (np.isfinite(inner_gy) and np.isfinite(outer_gy)):
            extent = polygon_extent(sol)
            if extent is not None:
                _, _, ymin, ymax = extent
                inner_gy, outer_gy = (ymin, ymax) if abs(ymin) <= abs(ymax) else (ymax, ymin)
            else:
                inner_gy, outer_gy = 0.0, 0.0
        ax_fp.annotate("", xy=(0, inner_gy), xytext=(0, 0), arrowprops={**ap, "color": colour}, zorder=4)
        ax_fp.annotate("", xy=(0, outer_gy), xytext=(0, inner_gy), arrowprops={**ap, "color": colour}, zorder=4)
        ax_fp.annotate(
            f"Inner offset\n{m_to_unit(abs(inner_gy), dist_unit):.0f} {dist_unit}",
            xy=(0, inner_gy / 2),
            xytext=(12, 0),
            textcoords="offset points",
            color=colour,
            fontsize=6.8,
            ha="left",
            va="center",
            bbox=label_box,
        )
        ax_fp.annotate(
            f"Outer offset\n{m_to_unit(abs(outer_gy), dist_unit):.0f} {dist_unit}",
            xy=(0, outer_gy),
            xytext=(12, -12 if outer_gy >= 0 else 12),
            textcoords="offset points",
            color=colour,
            fontsize=6.8,
            ha="left",
            va="top" if outer_gy >= 0 else "bottom",
            bbox=label_box,
        )
# Tick labels in display units
xt = ax_fp.get_xticks()
ax_fp.set_xticklabels([f"{m_to_unit(t, dist_unit):.0f}" for t in xt], color="#8b949e")
yt = ax_fp.get_yticks()
ax_fp.set_yticklabels([f"{m_to_unit(t, dist_unit):.0f}" for t in yt], color="#8b949e")
ax_fp.set_xlabel(f"← Across-track ({dist_unit}) →", color="#8b949e")
ax_fp.set_ylabel(f"↑ Along-track / forward ({dist_unit})", color="#8b949e")
ax_fp.set_title("Single-Frame Footprint Plan View — All Cameras", color="#c9d1d9",
                fontsize=11, pad=10)

legend_fp = [mpatches.Patch(color=col, label=cam["label"], alpha=0.85)
             for cam, _, col in solutions]
ax_fp.legend(handles=legend_fp, loc="upper right", fontsize=8, framealpha=0.4,
             labelcolor="#c9d1d9", facecolor="#161b22")

# Forward arrow
ax_fp.annotate("", xy=(0, lim * 0.90), xytext=(0, lim * 0.74),
               arrowprops=dict(arrowstyle="->", color="#c9d1d9", lw=1.8))
ax_fp.text(lim * 0.03, lim * 0.82, "Fwd", color="#c9d1d9", fontsize=7.5, va="center")

fig_fp.tight_layout()
st.pyplot(fig_fp)
st.markdown("---")

# ─────────────────────────────────────────────────────────────────────────────
# DIAGRAM 2 — Cross-section (across-track cameras)
# ─────────────────────────────────────────────────────────────────────────────

across_sols = [(c, s, col) for c, s, col in solutions if s.tilt_axis == "across"]

st.subheader("Cross-Section View — Across-Track")
st.caption(
    "Side view looking forward. "
    "Dotted line = inner edge ray, solid = centre ray, dashed = outer edge ray. "
    "GSD values shown at ground intercepts."
)

if across_sols:
    fig_xs, ax_xs = dark_fig(12, 5)
    H = altitude_m

    ax_xs.axvline(0, color="#30363d", lw=1.0, ls="--", zorder=1)
    ax_xs.axhline(0, color="#444", lw=1.5, zorder=1)
    ax_xs.scatter([0], [H], marker="^", s=220, color="#f0c040", zorder=7)
    ax_xs.annotate(f"{fmt(H, dist_unit)} AGL", (0, H), xytext=(10, -2),
                   textcoords="offset points", color="#f0c040", fontsize=8, fontweight="bold")

    candidate_x = [0.0]

    for cam, sol, colour in across_sols:
        inner_gx, outer_gx = corner_inner_outer(sol)
        centre_gx = sol.centre_m if np.isfinite(sol.centre_m) else 0.0
        if not (np.isfinite(inner_gx) and np.isfinite(outer_gx)):
            extent = polygon_extent(sol)
            if extent is not None:
                xmin, xmax, _, _ = extent
                inner_gx, outer_gx = (xmin, xmax) if abs(xmin) <= abs(xmax) else (xmax, xmin)
                centre_gx = 0.5 * (xmin + xmax)
        inner_gsd = min(sol.near_gsd_m, sol.far_gsd_m)
        outer_gsd = max(sol.near_gsd_m, sol.far_gsd_m)
        tilt_h = 90.0 - sol.tilt_from_nadir_deg
        near_a, far_a = get_inner_outer_angles(sol)

        for gx, ls, lw in [(inner_gx, ":", 1.4), (centre_gx, "-", 2.0), (outer_gx, "--", 1.4)]:
            if np.isfinite(gx):
                ax_xs.plot([0, gx], [H, 0], color=colour, ls=ls, lw=lw, zorder=3)
                candidate_x.append(gx)

        xs = [gx for gx in [inner_gx, centre_gx, outer_gx] if np.isfinite(gx)]
        if xs:
            ax_xs.scatter(xs, [0] * len(xs), color=colour, s=50, zorder=5, edgecolors="none")

        if np.isfinite(inner_gx):
            ax_xs.text(inner_gx, -H * 0.04,
                       f"{sol.label}\n{m_to_unit(abs(inner_gx), dist_unit):.0f} {dist_unit}\nnear {near_a:.1f}°",
                       color=colour, fontsize=6, ha="center", va="top")
            ax_xs.text(inner_gx * 0.55, H * 0.18,
                       f"{inner_gsd*100:.1f} cm/px",
                       color=colour, fontsize=6, ha="center", alpha=0.9)

        if np.isfinite(outer_gx):
            ax_xs.text(outer_gx, -H * 0.09,
                       f"{m_to_unit(abs(outer_gx), dist_unit):.0f} {dist_unit}\nfar {far_a:.1f}°",
                       color=colour, fontsize=6, ha="center", va="top")
            ax_xs.text(outer_gx * 0.55, H * 0.18,
                       f"{outer_gsd*100:.1f} cm/px",
                       color=colour, fontsize=6, ha="center", alpha=0.9)

        if np.isfinite(centre_gx):
            ax_xs.text(centre_gx * 0.5, H * 0.55,
                       f"tilt {tilt_h:.1f}° from horiz",
                       color=colour, fontsize=6, ha="center", alpha=0.75, style="italic")

    valid_x = [abs(v) for v in candidate_x if np.isfinite(v)]
    max_x = max(valid_x) * 1.25 if valid_x else 1.0
    if not np.isfinite(max_x) or max_x <= 0:
        max_x = 1.0
    ax_xs.set_xlim(-max_x, max_x)
    ax_xs.set_ylim(-H * 0.18, H * 1.15)

    xt = ax_xs.get_xticks()
    ax_xs.set_xticklabels([f"{m_to_unit(t, dist_unit):.0f}" for t in xt], color="#8b949e")
    yt = ax_xs.get_yticks()
    ax_xs.set_yticklabels([f"{m_to_unit(t, dist_unit):.0f}" for t in yt], color="#8b949e")
    ax_xs.set_xlabel(f"← Across-track ({dist_unit}) →", color="#8b949e")
    ax_xs.set_ylabel(f"Altitude ({dist_unit})", color="#8b949e")
    ax_xs.set_title("Cross-Section — Across-Track Cameras", color="#c9d1d9", fontsize=11)

    legend_xs = [mpatches.Patch(color=col, label=cam["label"], alpha=0.8)
                 for cam, _, col in across_sols]
    legend_xs += [
        mpatches.Patch(color="white", alpha=0, label=""),
        plt.Line2D([0],[0], color="white", ls=":", lw=1.4, label="Inner edge ray"),
        plt.Line2D([0],[0], color="white", ls="-", lw=2.0, label="Centre ray"),
        plt.Line2D([0],[0], color="white", ls="--", lw=1.4, label="Outer edge ray"),
    ]
    ax_xs.legend(handles=legend_xs, fontsize=7, framealpha=0.35,
                 labelcolor="#c9d1d9", facecolor="#161b22", loc="upper right", ncol=2)

    fig_xs.tight_layout()
    st.pyplot(fig_xs)
else:
    st.info("No across-track cameras enabled.")
st.markdown("---")

# ─────────────────────────────────────────────────────────────────────────────
# DIAGRAM 3 — Multi-strip plan view
# ─────────────────────────────────────────────────────────────────────────────

st.subheader("Multi-Strip Plan View")
show_nadir_multistrip = st.checkbox(
    "Include nadir footprint in multi-strip view",
    value=True,
    key="show_nadir_multistrip",
)
st.caption(
    "Three adjacent flight strips offset by the recommended line spacing. "
    "Each strip shows two successive frame instances so both sidelap and forward overlap are visible. "
    "The white hatched band shows the matched-frame cross-track sidelap between the right oblique on one strip "
    "and the left oblique on the adjacent reciprocal strip."
)

multistrip_solutions = [
    (cam, sol, col)
    for cam, sol, col in solutions
    if show_nadir_multistrip or "nadir" not in sol.label.lower()
]

if mc and multistrip_solutions:
    fig_ms, ax_ms = dark_fig(14, 8)
    ax_ms.set_aspect("equal")

    line_spacing, photo_spacing, used_spacing_fallback = fallback_multistrip_spacing(
        multistrip_solutions,
        fwd_frac=fwd_frac,
        side_frac=side_frac,
        mc=mc,
    )
    if used_spacing_fallback:
        st.info("Using geometric fallback spacing in the multi-strip view because one or more system spacing values were non-finite.")

    n_strips = 3
    strip_alphas = [0.55, 0.38, 0.22]
    strip_lws = [1.6, 0.9, 0.5]
    strip_cols = ["#c9d1d9", "#6e7681", "#444"]
    strip_positions = [((i - (n_strips // 2)) * line_spacing) for i in range(n_strips)]
    frame_offsets_y = [0.0, photo_spacing]

    all_x_ms = []
    all_y_ms = []

    for si, x_off in enumerate(strip_positions):
        for frame_idx, y_off in enumerate(frame_offsets_y):
            for cam, sol, colour in multistrip_solutions:
                base_corners = camera_polygon(sol)
                if len(base_corners) != 4:
                    continue
                corners = [(x + x_off, y + y_off) for x, y in base_corners]
                alpha = strip_alphas[si] if frame_idx == 0 else strip_alphas[si] * 0.45
                poly = plt.Polygon(
                    corners,
                    closed=True,
                    facecolor=colour,
                    alpha=alpha,
                    edgecolor=colour,
                    linewidth=strip_lws[si],
                    zorder=3,
                )
                ax_ms.add_patch(poly)
                for x, y in corners:
                    all_x_ms.append(x)
                    all_y_ms.append(y)

        ax_ms.axvline(x_off, color=strip_cols[si], lw=0.9, ls="-.", alpha=0.7, zorder=2)

    if all_y_ms:
        top_y = max(all_y_ms)
        for si, x_off in enumerate(strip_positions):
            ax_ms.text(
                x_off,
                top_y * 1.02,
                f"Strip {si+1}",
                color=strip_cols[si],
                fontsize=8,
                ha="center",
                va="bottom",
            )

    sidelap_band = matched_sidelap_band(multistrip_solutions, line_spacing=line_spacing, frame_y=0.0)
    if sidelap_band is not None:
        hatch_rect = mpatches.Rectangle(
            (sidelap_band["x0"], sidelap_band["y0"]),
            sidelap_band["x1"] - sidelap_band["x0"],
            sidelap_band["y1"] - sidelap_band["y0"],
            facecolor=(1, 1, 1, 0.06),
            edgecolor="white",
            hatch="////",
            linewidth=1.1,
            zorder=4,
        )
        ax_ms.add_patch(hatch_rect)
        label_x = 0.5 * (sidelap_band["x0"] + sidelap_band["x1"])
        label_y = sidelap_band["y1"] + 0.03 * max((max(all_y_ms) - min(all_y_ms)), 1.0)
        band_pct = sidelap_band["overlap_frac"] * 100 if np.isfinite(sidelap_band["overlap_frac"]) else mc.sidelap_achieved * 100
        ax_ms.text(
            label_x,
            label_y,
            f"Matched R/L sidelap\n{band_pct:.0f}%",
            color="white",
            fontsize=7,
            ha="center",
            va="bottom",
            bbox=dict(facecolor="#21262d", alpha=0.80, pad=3, edgecolor="#c9d1d9", lw=0.6),
            zorder=6,
        )

    if all_x_ms and all_y_ms:
        y_range = max(all_y_ms) - min(all_y_ms)
        x_range = max(all_x_ms) - min(all_x_ms)
        line_arrow_y = min(all_y_ms) - 0.10 * y_range
        ax_ms.annotate(
            "",
            xy=(line_spacing, line_arrow_y),
            xytext=(0, line_arrow_y),
            arrowprops=dict(arrowstyle="<->", color="white", lw=2.0, mutation_scale=13),
            zorder=5,
        )
        ax_ms.text(
            line_spacing / 2,
            line_arrow_y - 0.03 * y_range,
            f"Line spacing\n{m_to_unit(line_spacing, dist_unit):.0f} {dist_unit}",
            color="white",
            fontsize=8,
            ha="center",
            va="top",
            bbox=dict(facecolor="#21262d", alpha=0.75, pad=3, edgecolor="#555", lw=0.5),
        )

        x_for_photo = min(all_x_ms) - 0.06 * x_range
        ax_ms.annotate(
            "",
            xy=(x_for_photo, photo_spacing),
            xytext=(x_for_photo, 0),
            arrowprops=dict(arrowstyle="<->", color="white", lw=2.0, mutation_scale=13),
            zorder=5,
        )
        ax_ms.text(
            x_for_photo + 0.03 * x_range,
            photo_spacing / 2,
            f"Photo spacing\n{m_to_unit(photo_spacing, dist_unit):.0f} {dist_unit}",
            color="white",
            fontsize=8,
            ha="left",
            va="center",
            bbox=dict(facecolor="#21262d", alpha=0.75, pad=3, edgecolor="#555", lw=0.5),
        )

        x_min = min(all_x_ms) - 0.12 * x_range
        x_max = max(all_x_ms) + 0.12 * x_range
        y_min = min(all_y_ms) - 0.16 * y_range
        y_max = max(all_y_ms) + 0.10 * y_range
    else:
        x_min, x_max, y_min, y_max = -1, 1, -1, 1

    ax_ms.set_xlim(x_min, x_max)
    ax_ms.set_ylim(y_min, y_max)

    xt = ax_ms.get_xticks()
    ax_ms.set_xticklabels([f"{m_to_unit(t, dist_unit):.0f}" for t in xt], color="#8b949e")
    yt = ax_ms.get_yticks()
    ax_ms.set_yticklabels([f"{m_to_unit(t, dist_unit):.0f}" for t in yt], color="#8b949e")
    ax_ms.set_xlabel(f"← Across-track ({dist_unit}) →", color="#8b949e")
    ax_ms.set_ylabel(f"↑ Along-track ({dist_unit})", color="#8b949e")
    ax_ms.set_title("Multi-Strip Plan View (3 strips, 2 frame instances)", color="#c9d1d9", fontsize=11)

    legend_ms = [mpatches.Patch(color=col, label=cam["label"], alpha=0.8)
                 for cam, _, col in multistrip_solutions]
    ax_ms.legend(handles=legend_ms, fontsize=8, framealpha=0.35,
                 labelcolor="#c9d1d9", facecolor="#161b22", loc="upper right")

    fig_ms.tight_layout()
    st.pyplot(fig_ms)

elif not mc:
    st.info("System solution unavailable.")
else:
    st.info("No cameras enabled for the multi-strip view.")

st.markdown("---")

# ─────────────────────────────────────────────────────────────────────────────
# DIAGRAM 4 — Point coverage / views per point
# ─────────────────────────────────────────────────────────────────────────────

coverage_fig_hits = None
coverage_fig_angles = None
coverage_result = None
coverage_probe = None

help_toggle(
    "Point Coverage / Views Per Point",
    "This section samples one repeating survey cell to estimate how many times a ground point is captured within the pattern. Image hits count every photo that sees the point. Distinct viewing angles count unique camera directions such as nadir, left oblique, right oblique, fore oblique and aft oblique. Use the min, average and max together because coverage is not uniform across the block. The min/average/max figures and heatmaps are sampled across one repeat cell, while the red gap warning below now also uses an exact geometric zero-coverage check when Shapely is available. Standard precision means a 61 × 61 sample grid, which is 3,721 tested points across one repeat cell.",
    key="point_coverage_intro",
)

samples_x, samples_y = 61, 61
sample_grid_label = coverage_sampling_label(samples_x, samples_y)
st.caption(f"Standard: {sample_grid_label} tested across one repeat cell.")

line_spacing_pc, photo_spacing_pc, used_pc_fallback = fallback_multistrip_spacing(
    solutions,
    fwd_frac=fwd_frac,
    side_frac=side_frac,
    mc=mc,
)

coverage_result = compute_point_coverage(
    solutions,
    line_spacing=line_spacing_pc,
    photo_spacing=photo_spacing_pc,
    samples_x=samples_x,
    samples_y=samples_y,
)

if coverage_result is not None:
    if used_pc_fallback:
        st.info("Using geometric fallback spacing for the point-coverage estimate because one or more system spacing values were non-finite.")

    summary = coverage_summary(coverage_result)
    gap_stats = coverage_gap_stats(coverage_result)

    help_toggle(
        "How to read the summary",
        "Minimum = the weakest sampled point inside the repeating cell. Average = a typical sampled point. Maximum = the strongest sampled overlap zone. For client conversations, the safest wording is usually to quote the range and then mention the typical value, for example: each point gets between X and Y image hits, with an average of Z. The heatmaps and min/avg/max remain sampled, but the zero-coverage warning below now also uses an exact geometric gap test when available. Standard precision uses a 61 × 61 sample grid, not a 61 m by 61 m area.",
        key="point_coverage_summary",
    )

    gap_mode = classify_gap_presentation(gap_stats)

    if gap_mode == "shortfall":
        warning_title = "⚠ COVERAGE SHORTFALL DETECTED"
        if gap_stats["exact_available"]:
            warning_summary = (
                f"Exact zero-coverage gap present. {format_gap_pct(gap_stats['zero_hit_pct'])} of the repeat-cell area has no image coverage."
            )
            warning_detail = (
                f"The heatmaps and min/avg/max values are still sampled at Standard "
                f"using a {sample_grid_label}, but the gap warning itself is based on an exact geometric check. "
                f"For reference, {gap_stats['sample_zero_hit_pct']:.1f}% of sampled points had zero image hits and "
                f"{gap_stats['sample_zero_angle_pct']:.1f}% had zero viewing angles."
            )
        else:
            warning_summary = (
                f"Sampled coverage gaps detected at Standard using a {sample_grid_label}."
            )
            warning_detail = (
                f"{gap_stats['sample_zero_hit_pct']:.1f}% of sampled points had zero image hits, and "
                f"{gap_stats['sample_zero_angle_pct']:.1f}% had zero viewing angles. "
                "This warning is currently based on the sampled repeat-cell check."
            )
        st.markdown(
            f"""
            <div style="background:rgba(248,81,73,0.18); border:2px solid rgba(248,81,73,0.95); border-left:10px solid #f85149; border-radius:0.95rem; padding:1rem 1.15rem; margin:0.45rem 0 1rem 0; box-shadow:0 0 0 1px rgba(248,81,73,0.18) inset;">
                <div style="font-size:1.42rem; font-weight:900; letter-spacing:0.02em; color:#ffb3ad; margin-bottom:0.28rem;">{warning_title}</div>
                <div style="font-size:1.06rem; font-weight:800; color:#f0f6fc; margin-bottom:0.28rem;">{warning_summary}</div>
                <div style="font-size:0.98rem; font-weight:600; color:#f0f6fc; line-height:1.45;">{warning_detail}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    elif gap_mode == "micro":
        st.warning(
            f"Possible microscopic exact-gap sliver detected: {format_gap_pct(gap_stats['zero_hit_pct'])} of the repeat-cell area. "
            f"The exact geometric check found a tiny uncovered sliver, but the {sample_grid_label} found 0.0% zero-hit and 0.0% zero-angle sampled points. "
            "Treat this as a vanishingly small seam or numerical precision edge unless it persists at higher precision."
        )
    else:
        if gap_stats["exact_available"]:
            st.success(
                f"No exact zero-coverage gaps detected inside the repeat cell. "
               f"The heatmaps and min/avg/max below are still sampled at Standard using a {sample_grid_label}."
            )
        else:
            st.success(
                f"No sampled gaps detected at Standard using a {sample_grid_label}. "
                "This is a sampled check of the repeat cell, not a formal geometric proof."
            )

    m1, m2, m3 = st.columns(3)
    metric_with_help(
        m1,
        "Min image hits",
        f"{summary['hits_min']}",
        "The weakest sampled point inside the repeating survey cell. If this is zero, the sampled check found a coverage gap somewhere under the current settings. For client conversations, this is the safest number to describe as the minimum image count per point.",
    )
    metric_with_help(
        m2,
        "Average image hits",
        f"{summary['hits_avg']:.1f}",
        "A typical sampled point across the repeating survey cell. This is useful for explaining the general level of redundancy, but it is not a guarantee because some points will sit in weaker or stronger overlap zones.",
    )
    metric_with_help(
        m3,
        "Max image hits",
        f"{summary['hits_max']}",
        "The strongest sampled overlap zone inside the repeating survey cell. This often occurs where strip overlap and forward overlap stack together. It is helpful for understanding peak redundancy, but it should not be presented as typical coverage.",
    )

    m4, m5, m6 = st.columns(3)
    metric_with_help(
        m4,
        "Min viewing angles",
        f"{summary['angles_min']}",
        "The weakest directional diversity at any sampled point. A value of zero means there is no sampled coverage at that location. A value of one means the point is seen from only one camera family or direction. This is the conservative number to quote when clients ask about minimum angular diversity.",
    )
    metric_with_help(
        m5,
        "Average viewing angles",
        f"{summary['angles_avg']:.1f}",
        "A typical number of distinct camera directions seeing a sampled point across the repeat cell. This is often the easiest client-facing indicator of how much 3D viewing diversity the survey provides overall.",
    )
    metric_with_help(
        m6,
        "Max viewing angles",
        f"{summary['angles_max']}",
        "The strongest directional diversity inside the sampled repeat cell. These are usually the best-covered overlap zones where several camera families see the same point. Use this as an upper bound, not the expected value everywhere.",
    )

    st.caption(
        "Image hits = total photos containing a point. Viewing angles = unique camera families seeing that point, not repeated photos from the same direction."
    )

    help_toggle(
        "Probe a single point",
        "Use this to answer a very specific question such as: how many photos and how many distinct angles would a point near the strip centre, overlap zone or edge receive? \
The probe point sits inside one repeating line-spacing by photo-spacing cell.",
        key="point_coverage_probe",
    )

    probe_col1, probe_col2 = st.columns(2)
    with probe_col1:
        probe_x_display = st.slider(
            f"Probe across-track position ({dist_unit})",
            min_value=float(m_to_unit(-0.5 * line_spacing_pc, dist_unit)),
            max_value=float(m_to_unit(0.5 * line_spacing_pc, dist_unit)),
            value=0.0,
            step=max(0.5, round(m_to_unit(line_spacing_pc / 40.0, dist_unit), 1)),
        )
    with probe_col2:
        probe_y_display = st.slider(
            f"Probe along-track position ({dist_unit})",
            min_value=float(m_to_unit(-0.5 * photo_spacing_pc, dist_unit)),
            max_value=float(m_to_unit(0.5 * photo_spacing_pc, dist_unit)),
            value=0.0,
            step=max(0.5, round(m_to_unit(photo_spacing_pc / 40.0, dist_unit), 1)),
        )

    probe_x = unit_to_m(probe_x_display, dist_unit)
    probe_y = unit_to_m(probe_y_display, dist_unit)
    coverage_probe = point_coverage_at(probe_x, probe_y, coverage_result["sources"])

    p1, p2 = st.columns(2)
    with p1:
        metric_with_help(
            p1,
            "Probe image hits",
            str(coverage_probe["hits"]),
            "This is the exact number of photos covering the chosen probe point. Use it when a client asks about a specific position, such as the strip centre, overlap zone, or edge, rather than the whole coverage pattern.",
        )
        metric_with_help(
            p1,
            "Probe viewing angles",
            str(coverage_probe["unique_angles"]),
            "This is the exact number of distinct camera directions covering the chosen probe point. It helps explain whether that location is being seen from just one direction, or from several different look angles that improve 3D interpretation.",
        )
    with p2:
        st.markdown("**Viewing angle families at probe point**")
        if coverage_probe["families"]:
            st.write(", ".join(coverage_probe["families"]))
        else:
            st.write("No enabled camera footprints intersect this point.")

    help_toggle(
        "Heatmaps",
        "The heatmaps show how sampled coverage changes across one repeating cell of the survey pattern. Red cells mark sampled zero-coverage holes and the colour scale starts at zero, so it is easier to spot likely gaps. The separate warning above now also uses an exact geometric zero-coverage check when available. The current precision setting controls how many test points are laid across that repeat cell, for example standard is a 61 × 61 sample grid. Central overlap zones will often have more hits and sometimes more unique viewing angles than edge zones.",
        key="point_coverage_heatmaps",
    )

    coverage_fig_hits = coverage_heatmap_figure(
        coverage_result["xs"],
        coverage_result["ys"],
        coverage_result["hits"],
        title="Point Coverage Heatmap — Total Image Hits",
        colorbar_label="Image hits per point",
        probe_xy=(probe_x, probe_y),
    )
    coverage_fig_angles = coverage_heatmap_figure(
        coverage_result["xs"],
        coverage_result["ys"],
        coverage_result["angle_counts"],
        title="Point Coverage Heatmap — Distinct Viewing Angles",
        colorbar_label="Distinct viewing angles",
        probe_xy=(probe_x, probe_y),
    )

    h1, h2 = st.columns(2)
    with h1:
        st.pyplot(coverage_fig_hits)
    with h2:
        st.pyplot(coverage_fig_angles)

    st.caption(
        f"Repeat cell used for the estimate: {m_to_unit(line_spacing_pc, dist_unit):.0f} {dist_unit} line spacing by {m_to_unit(photo_spacing_pc, dist_unit):.0f} {dist_unit} photo spacing."
    )
else:
    st.info("Point coverage could not be estimated from the current camera footprints.")

st.markdown("---")



st.subheader("Optimiser")
help_toggle(
    "Client requirements → best config",
    "This optimiser is a sandbox for finding the lightest survey pattern that still satisfies a client brief. It does not replace any of the existing sections above. Enter the minimum coverage or quality rules you need to hit, define the ranges you want searched, then run it. The optimiser ranks passing candidates by efficiency, which here means the largest line-spacing by photo-spacing area per pass, adjusted for whether reciprocal flying is required.",
    key="optimiser_intro",
)

current_oblique_tilt = current_common_oblique_value(st.session_state.cameras, "tilt_deg", 50.0)
current_oblique_fl = current_common_oblique_value(st.session_state.cameras, "focal_mm", 50.0)
enabled_obliques_present = any(
    cam.get("enabled", False) and "nadir" not in str(cam.get("label", "")).lower()
    for cam in st.session_state.cameras
)
current_hits_min_default = 1
current_angles_min_default = 2
current_inner_gsd_default = 5.0
current_outer_gsd_default = 8.5
if solutions:
    optimizer_reference_solutions = [sol for _, sol, _ in solutions if abs(sol.tilt_from_nadir_deg) > 1.0]
    if not optimizer_reference_solutions:
        optimizer_reference_solutions = [sol for _, sol, _ in solutions]
    current_inner_gsd_default = max((min(sol.near_gsd_m, sol.far_gsd_m) * 100.0) for sol in optimizer_reference_solutions)
    current_outer_gsd_default = max((max(sol.near_gsd_m, sol.far_gsd_m) * 100.0) for sol in optimizer_reference_solutions)
if coverage_result is not None:
    current_summary_for_optimizer = coverage_summary(coverage_result)
    current_hits_min_default = max(1, int(current_summary_for_optimizer["hits_min"]))
    current_angles_min_default = max(1, int(current_summary_for_optimizer["angles_min"]))

req_col1, req_col2 = st.columns(2)
with req_col1:
    opt_min_viewing_angles = st.number_input(
        "Minimum viewing angles anywhere in the survey",
        min_value=0,
        max_value=10,
        value=max(2, current_angles_min_default),
        step=1,
        key="opt_min_viewing_angles",
    )
    opt_min_image_hits = st.number_input(
        "Minimum image hits anywhere in the survey",
        min_value=0,
        max_value=50,
        value=current_hits_min_default,
        step=1,
        key="opt_min_image_hits",
    )
    opt_require_no_gaps = st.checkbox(
        "Require full repeat-cell coverage (no gaps)",
        value=True,
        key="opt_require_no_gaps",
    )
    opt_enforce_inner_gsd = st.checkbox(
        "Set a maximum inner-edge GSD",
        value=False,
        key="opt_enforce_inner_gsd",
    )
    opt_max_inner_gsd_cm = None
    if opt_enforce_inner_gsd:
        opt_max_inner_gsd_cm = st.number_input(
            "Maximum inner-edge GSD (cm/px)",
            min_value=0.5,
            max_value=100.0,
            value=round(current_inner_gsd_default, 2),
            step=0.1,
            key="opt_max_inner_gsd_cm",
        )
    opt_enforce_outer_gsd = st.checkbox(
        "Set a maximum outer-edge GSD",
        value=False,
        key="opt_enforce_outer_gsd",
    )
    opt_max_outer_gsd_cm = None
    if opt_enforce_outer_gsd:
        opt_max_outer_gsd_cm = st.number_input(
            "Maximum outer-edge GSD (cm/px)",
            min_value=0.5,
            max_value=100.0,
            value=round(current_outer_gsd_default, 2),
            step=0.1,
            key="opt_max_outer_gsd_cm",
        )

with req_col2:
    opt_min_achieved_sidelap = st.number_input(
        "Minimum achieved sidelap (%)",
        min_value=0.0,
        max_value=95.0,
        value=float(side_pct),
        step=1.0,
        key="opt_min_achieved_sidelap",
    )
    opt_reciprocal_mode = st.selectbox(
        "Reciprocal strips during optimisation",
        ["Keep current", "Force on", "Force off", "Try both"],
        index=3,
        key="opt_reciprocal_mode",
    )
    opt_precision = st.selectbox(
        "Optimiser coverage precision",
        ["Fast preview", "Standard"],
        index=0,
        key="opt_precision",
        help="Fast preview is better for scanning lots of combinations. Standard is more defensible if you are checking a narrow final range.",
    )
    optimiser_precision_map = {"Fast preview": (31, 31), "Standard": (61, 61)}
    opt_samples_x, opt_samples_y = optimiser_precision_map[opt_precision]

range_col1, range_col2, range_col3 = st.columns(3)
with range_col1:
    st.markdown("**Overlap search range**")
    opt_fwd_min = st.number_input("Forward overlap min (%)", min_value=10, max_value=95, value=max(10, int(fwd_pct) - 15), step=1, key="opt_fwd_min")
    opt_fwd_max = st.number_input("Forward overlap max (%)", min_value=10, max_value=95, value=min(95, int(fwd_pct) + 15), step=1, key="opt_fwd_max")
    opt_fwd_step = st.number_input("Forward overlap step (%)", min_value=1, max_value=20, value=5, step=1, key="opt_fwd_step")

with range_col2:
    st.markdown("**Sidelap search range**")
    opt_side_min = st.number_input("Sidelap min (%)", min_value=10, max_value=95, value=max(10, int(side_pct) - 20), step=1, key="opt_side_min")
    opt_side_max = st.number_input("Sidelap max (%)", min_value=10, max_value=95, value=min(95, int(side_pct) + 20), step=1, key="opt_side_max")
    opt_side_step = st.number_input("Sidelap step (%)", min_value=1, max_value=20, value=5, step=1, key="opt_side_step")

with range_col3:
    st.markdown("**Optional camera / height search**")
    opt_search_tilt = st.checkbox("Search common oblique tilt", value=False, key="opt_search_tilt", disabled=not enabled_obliques_present)
    opt_search_focal = st.checkbox("Search common oblique focal length", value=False, key="opt_search_focal", disabled=not enabled_obliques_present)
    opt_search_altitude = st.checkbox("Search altitude", value=True, key="opt_search_altitude")
    st.caption("Altitude search is useful when the client brief is driven by inner-edge and outer-edge GSD limits.")

opt_extra_col1, opt_extra_col2, opt_extra_col3 = st.columns(3)
with opt_extra_col1:
    if opt_search_tilt and enabled_obliques_present:
        opt_tilt_min = st.number_input("Oblique tilt min (°)", min_value=0.0, max_value=85.0, value=max(0.0, current_oblique_tilt - 10.0), step=0.5, key="opt_tilt_min")
        opt_tilt_max = st.number_input("Oblique tilt max (°)", min_value=0.0, max_value=85.0, value=min(85.0, current_oblique_tilt + 10.0), step=0.5, key="opt_tilt_max")
        opt_tilt_step = st.number_input("Oblique tilt step (°)", min_value=0.5, max_value=20.0, value=2.5, step=0.5, key="opt_tilt_step")
    else:
        opt_tilt_min = opt_tilt_max = current_oblique_tilt
        opt_tilt_step = 1.0

with opt_extra_col2:
    if opt_search_focal and enabled_obliques_present:
        opt_fl_min = st.number_input("Oblique FL min (mm)", min_value=1.0, max_value=300.0, value=max(1.0, current_oblique_fl - 15.0), step=1.0, key="opt_fl_min")
        opt_fl_max = st.number_input("Oblique FL max (mm)", min_value=1.0, max_value=300.0, value=min(300.0, current_oblique_fl + 15.0), step=1.0, key="opt_fl_max")
        opt_fl_step = st.number_input("Oblique FL step (mm)", min_value=1.0, max_value=50.0, value=5.0, step=1.0, key="opt_fl_step")
    else:
        opt_fl_min = opt_fl_max = current_oblique_fl
        opt_fl_step = 1.0

with opt_extra_col3:
    current_alt_display = float(m_to_unit(altitude_m, dist_unit))
    if opt_search_altitude:
        opt_alt_min_display = st.number_input(f"Altitude min ({dist_unit})", min_value=1.0, max_value=m_to_unit(10000.0, dist_unit), value=max(1.0, current_alt_display - m_to_unit(150.0, dist_unit)), step=max(1.0, m_to_unit(10.0, dist_unit)), key="opt_alt_min")
        opt_alt_max_display = st.number_input(f"Altitude max ({dist_unit})", min_value=1.0, max_value=m_to_unit(10000.0, dist_unit), value=min(m_to_unit(10000.0, dist_unit), current_alt_display + m_to_unit(150.0, dist_unit)), step=max(1.0, m_to_unit(10.0, dist_unit)), key="opt_alt_max")
        opt_alt_step_display = st.number_input(f"Altitude step ({dist_unit})", min_value=max(0.5, m_to_unit(5.0, dist_unit)), max_value=m_to_unit(500.0, dist_unit), value=max(1.0, m_to_unit(25.0, dist_unit)), step=max(0.5, m_to_unit(5.0, dist_unit)), key="opt_alt_step")
    else:
        opt_alt_min_display = opt_alt_max_display = current_alt_display
        opt_alt_step_display = max(1.0, m_to_unit(10.0, dist_unit))

def build_search_values(start_value, end_value, step_value, decimals=0):
    if step_value <= 0:
        return [round(float(start_value), decimals)]
    lo = min(float(start_value), float(end_value))
    hi = max(float(start_value), float(end_value))
    values = []
    current = lo
    while current <= hi + (step_value * 0.25):
        values.append(round(current, decimals))
        current += float(step_value)
    if not values:
        values = [round(lo, decimals)]
    if round(hi, decimals) not in values:
        values.append(round(hi, decimals))
    return sorted(set(values))

forward_values = build_search_values(opt_fwd_min, opt_fwd_max, opt_fwd_step, decimals=0)
side_values = build_search_values(opt_side_min, opt_side_max, opt_side_step, decimals=0)
if opt_reciprocal_mode == "Keep current":
    reciprocal_values = [bool(reciprocal)]
elif opt_reciprocal_mode == "Force on":
    reciprocal_values = [True]
elif opt_reciprocal_mode == "Force off":
    reciprocal_values = [False]
else:
    reciprocal_values = [False, True]

tilt_values = build_search_values(opt_tilt_min, opt_tilt_max, opt_tilt_step, decimals=1) if opt_search_tilt and enabled_obliques_present else [round(current_oblique_tilt, 1)]
focal_values = build_search_values(opt_fl_min, opt_fl_max, opt_fl_step, decimals=1) if opt_search_focal and enabled_obliques_present else [round(current_oblique_fl, 1)]
altitude_values_m = [unit_to_m(v, dist_unit) for v in build_search_values(opt_alt_min_display, opt_alt_max_display, opt_alt_step_display, decimals=1)] if opt_search_altitude else [altitude_m]

combination_count = len(forward_values) * len(side_values) * len(reciprocal_values) * len(tilt_values) * len(focal_values) * len(altitude_values_m)
st.caption(
    f"Search size: {combination_count:,} candidate combinations using {coverage_sampling_label(opt_samples_x, opt_samples_y)}. "
    f"Oblique tilt search is {'on' if opt_search_tilt and enabled_obliques_present else 'off'}, focal search is {'on' if opt_search_focal and enabled_obliques_present else 'off'}, altitude search is {'on' if opt_search_altitude else 'off'}."
)

requirements = {
    "min_viewing_angles": int(opt_min_viewing_angles),
    "min_image_hits": int(opt_min_image_hits),
    "require_no_gaps": bool(opt_require_no_gaps),
    "max_inner_gsd_cm": float(opt_max_inner_gsd_cm) if opt_enforce_inner_gsd and opt_max_inner_gsd_cm is not None else None,
    "max_outer_gsd_cm": float(opt_max_outer_gsd_cm) if opt_enforce_outer_gsd and opt_max_outer_gsd_cm is not None else None,
    "min_achieved_sidelap_pct": float(opt_min_achieved_sidelap),
}

run_optimiser = st.button("Run optimiser", key="run_optimiser")
if run_optimiser:
    if combination_count > 2500:
        st.error("The current optimiser range is very large. Reduce one or more ranges or increase the step size so the search stays under about 2,500 combinations.")
    else:
        with st.spinner("Testing candidate survey configurations..."):
            candidates = []
            for forward_candidate, side_candidate, reciprocal_candidate, tilt_candidate, focal_candidate, altitude_candidate in product(
                forward_values,
                side_values,
                reciprocal_values,
                tilt_values,
                focal_values,
                altitude_values_m,
            ):
                candidate = evaluate_optimizer_candidate(
                    base_cameras=st.session_state.cameras,
                    altitude_m=float(altitude_candidate),
                    speed_ms=float(speed_ms),
                    forward_overlap_pct=int(forward_candidate),
                    side_overlap_pct=int(side_candidate),
                    reciprocal=bool(reciprocal_candidate),
                    samples_x=opt_samples_x,
                    samples_y=opt_samples_y,
                    requirements=requirements,
                    common_oblique_tilt_deg=float(tilt_candidate) if opt_search_tilt and enabled_obliques_present else None,
                    common_oblique_focal_mm=float(focal_candidate) if opt_search_focal and enabled_obliques_present else None,
                )
                candidates.append(candidate)

        valid_candidates = sorted([c for c in candidates if c.get("valid")], key=optimizer_valid_sort_key)
        near_miss_candidates = sorted([c for c in candidates if not c.get("valid")], key=optimizer_near_miss_sort_key)
        st.session_state.optimizer_results = {
            "best_valid": valid_candidates[0] if valid_candidates else None,
            "best_near_miss": near_miss_candidates[0] if near_miss_candidates else None,
            "top_rows": [optimizer_candidate_row(c, dist_unit) for c in (valid_candidates[:10] if valid_candidates else near_miss_candidates[:10])],
            "searched": combination_count,
            "requirements": requirements,
        }

optimizer_results = st.session_state.get("optimizer_results")
if optimizer_results:
    best_valid = optimizer_results.get("best_valid")
    best_near_miss = optimizer_results.get("best_near_miss")

    if best_valid:
        st.success("Best passing configuration found. This is the most efficient candidate that still satisfies the current client rules.")
        b1, b2, b3, b4 = st.columns(4)
        b1.metric("Recommended overlaps", f"{best_valid['forward_pct']}% / {best_valid['side_pct']}%")
        b2.metric("Min views / hits", f"{best_valid['angles_min']} / {best_valid['hits_min']}")
        b3.metric("Inner / outer GSD", f"{best_valid['inner_gsd_cm']:.2f} / {best_valid['outer_gsd_cm']:.2f} cm")
        b4.metric("Reciprocal", "Yes" if best_valid['reciprocal'] else "No")

        c1, c2, c3, c4 = st.columns(4)
        c1.metric(f"Altitude ({dist_unit})", f"{m_to_unit(best_valid['altitude_m'], dist_unit):.1f}")
        c2.metric(f"Line spacing ({dist_unit})", f"{m_to_unit(best_valid['line_spacing_m'], dist_unit):.1f}")
        c3.metric(f"Photo spacing ({dist_unit})", f"{m_to_unit(best_valid['photo_spacing_m'], dist_unit):.1f}")
        c4.metric(f"Efficiency ({dist_unit}²/pass)", f"{(m_to_unit(1.0, dist_unit) ** 2) * best_valid['efficiency_index_m2']:.1f}")

        st.caption(
            f"Common oblique values used for this recommendation: tilt {best_valid['common_oblique_tilt_deg']:.1f}° and focal length {best_valid['common_oblique_focal_mm']:.1f} mm. "
            f"Minimum oblique angle at the image edges stays between {best_valid['min_oblique_angle_deg']:.1f}° and {best_valid['max_oblique_angle_deg']:.1f}° from nadir."
        )

        if st.button("Apply best config to planner", key="apply_best_optimiser_config"):
            st.session_state.pending_optimizer_apply = dict(best_valid)
            st.rerun()
    elif best_near_miss:
        st.warning("No candidate met every rule in the current search range. Here is the closest near-miss so you can see what is failing and widen or shift the search.")
        n1, n2, n3, n4 = st.columns(4)
        n1.metric("Tried overlaps", f"{best_near_miss['forward_pct']}% / {best_near_miss['side_pct']}%")
        n2.metric("Min views / hits", f"{best_near_miss['angles_min']} / {best_near_miss['hits_min']}")
        n3.metric("Inner / outer GSD", f"{best_near_miss['inner_gsd_cm']:.2f} / {best_near_miss['outer_gsd_cm']:.2f} cm")
        n4.metric("Gap area", format_gap_pct(best_near_miss['zero_hit_pct']) if best_near_miss.get('has_gap') else "0.00%")
        st.caption(best_near_miss.get("reason", ""))

    top_rows = optimizer_results.get("top_rows") or []
    if top_rows:
        st.markdown("**Top optimiser results**")
        st.dataframe(top_rows, width="stretch")

st.markdown("---")

st.subheader("Export Results")
report_figures = [("Footprint plan view", fig_to_png_bytes(fig_fp))]
if across_sols:
    report_figures.append(("Cross-section view", fig_to_png_bytes(fig_xs)))
if mc and multistrip_solutions:
    report_figures.append(("Multi-strip plan view", fig_to_png_bytes(fig_ms)))
if coverage_fig_hits is not None:
    report_figures.append(("Point coverage heatmap - image hits", fig_to_png_bytes(coverage_fig_hits)))
if coverage_fig_angles is not None:
    report_figures.append(("Point coverage heatmap - viewing angles", fig_to_png_bytes(coverage_fig_angles)))

mission_rows = build_mission_export_rows(mission_outputs, dist_unit=dist_unit) if mission_outputs is not None else []
mission_figure_bytes = fig_to_png_bytes(aoi_fig) if aoi_fig is not None else None

col_exp1, col_exp2 = st.columns(2)
with col_exp1:
    excel_bytes = make_excel_export(
        settings_rows,
        system_rows,
        camera_rows,
        mission_rows=mission_rows,
        mission_figure_bytes=mission_figure_bytes,
    )
    st.download_button(
        label="Download Excel data report",
        data=excel_bytes,
        file_name="oblique_planner_report.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
with col_exp2:
    word_bytes = make_word_export(
        settings_rows,
        system_rows,
        camera_rows,
        mission_rows=mission_rows,
        mission_figure_bytes=mission_figure_bytes,
        report_figures=report_figures,
    )
    st.download_button(
        label="Download Word client report",
        data=word_bytes,
        file_name="oblique_planner_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

st.markdown("---")
# ─────────────────────────────────────────────────────────────────────────────
# Formula trace
# ─────────────────────────────────────────────────────────────────────────────

st.subheader("Formula Trace")
with st.expander("Show intermediate calculations per camera", expanded=False):
    for cam, sol, colour in solutions:
        bd = get_body(cam["body"])
        inner_gx, outer_gx = corner_inner_outer(sol)
        inner_len, outer_len = along_lengths_for_display(sol)
        inner_gsd = min(sol.near_gsd_m, sol.far_gsd_m)
        outer_gsd = max(sol.near_gsd_m, sol.far_gsd_m)
        tilt_h = 90.0 - sol.tilt_from_nadir_deg
        ob = obliqueness_ratio(sol)

        st.markdown(
            f"<span style='color:{colour}'>●</span> **{sol.label}**",
            unsafe_allow_html=True
        )
        st.markdown(f"""
| Quantity | Formula | Value |
|---|---|---|
| Sensor native | — | {bd['w_mm']} × {bd['h_mm']} mm, {bd['w_px']} × {bd['h_px']} px |
| Orientation | **{sol.orientation}** | across: **{sol.sensor_across_mm:.4f} mm**, along: **{sol.sensor_along_mm:.4f} mm** |
| Focal length | — | **{cam['focal_mm']} mm** |
| Pixel size | `sensor_across / image_across_px` | **{sol.pixel_size_mm*1000:.3f} µm** |
| Tilt | — | **{sol.tilt_from_nadir_deg:.2f}° from nadir** / **{tilt_h:.2f}° from horizontal** |
| Tilt axis | — | **{sol.tilt_axis}** |
| Half FOV across | `atan(sensor_across / (2×fl))` | {sol.half_fov_across_deg:.4f}° → full {sol.full_fov_across_deg:.4f}° |
| Half FOV along | `atan(sensor_along / (2×fl))` | {sol.half_fov_along_deg:.4f}° → full {sol.full_fov_along_deg:.4f}° |
| Diag PP→edge | `sqrt((sensor_across/2)² + fl²)` | **{sol.diag_image_mm:.4f} mm** |
| Inner edge | `H × tan(θ − φ_w)` | **{m_to_unit(abs(inner_gx), dist_unit):.2f} {dist_unit}** from nadir |
| Outer edge | `H × tan(θ + φ_w)` | **{m_to_unit(abs(outer_gx), dist_unit):.2f} {dist_unit}** from nadir |
| Inner slant | `sqrt(H² + Gx²)` | **{min(sol.near_slant_m, sol.far_slant_m):.2f} m** |
| Outer slant | `sqrt(H² + Gx²)` | **{max(sol.near_slant_m, sol.far_slant_m):.2f} m** |
| Inner length | exact 4-corner | **{m_to_unit(inner_len, dist_unit):.2f} {dist_unit}** |
| Outer length | exact 4-corner | **{m_to_unit(outer_len, dist_unit):.2f} {dist_unit}** |
| Inner GSD | `px_mm × slant_mm / diag_mm` | **{fmt_gsd(inner_gsd)}** |
| Centre GSD | — | **{fmt_gsd(sol.centre_gsd_m)}** |
| Outer GSD | `px_mm × slant_mm / diag_mm` | **{fmt_gsd(outer_gsd)}** |
| **Obliqueness ratio** | `outer GSD / inner GSD` | **{ob:.3f}×** |
        """)

st.markdown("---")

# ─────────────────────────────────────────────────────────────────────────────
# Assumptions
# ─────────────────────────────────────────────────────────────────────────────

with st.expander("ℹ️ Assumptions, conventions and spreadsheet verification"):
    st.markdown("""
### Spreadsheet verification (Oblique_setup9_working_2.xls)

**Landscape sheet** — Sony A7R V, nadir fl=21 mm, oblique fl=50 mm, 50° from horizontal, GSD 8.5 cm:

| Value | Our app | Spreadsheet |
|---|---|---|
| Flying height | 469.737 m | 469.737 m |
| Oblique inner edge | 233.820 m | 233.820 m |
| Oblique outer edge | 635.679 m | 635.679 m |
| Inner length | 368.473 m | 368.473 m |
| Outer length | 555.051 m | 555.051 m |
| Inner GSD | 3.877 cm/px | 3.877 cm/px |
| Outer GSD | 5.840 cm/px | 5.840 cm/px |
| Slope to inner | 524.714 m | 524.714 m |
| Slope to outer | 790.405 m | 790.405 m |

All values match to 3+ decimal places.

**Portrait sheet** — uses landscape-mounted oblique cameras (long axis across-track) at a different
flying height (H=631.58 m, nadir fl=48 mm, GSD 5 cm). To reproduce that sheet, set the L/R oblique
cameras to **Landscape** orientation. Both sheets match our model when the correct orientation is used.

### Obliqueness ratio
Defined as `outer GSD / inner GSD`. Values:
- **1.0** = nadir camera (uniform GSD across the image)
- **1.5** = mild oblique (50% more GSD at the far edge than inner edge)
- **>3** = highly oblique (consider whether far-edge GSD meets mission requirements)

### Sensor orientation convention

| Setting | Sensor dimension across-track | Along-track | Typical use |
|---|---|---|---|
| **Portrait** | Short (narrow) axis | Long axis | L/R oblique |
| **Landscape** | Long axis | Short axis | Nadir camera |

### GSD formula (slant-plane, matching spreadsheet)
```
GSD = pixel_size_mm × slant_2d_mm / diag_image_mm
diag_image_mm = sqrt((sensor_across/2)² + focal_length²)
```

### Photo spacing
Uses the **inner (minimum) footprint length** so the target forward overlap is
met at the most constrained position. The far edge will have higher overlap than requested.

### Assumptions (v1)
- Flat terrain
- Pinhole camera, square pixels
- Pure single-axis tilt per camera
- No lens distortion, wind drift, or lever-arm effects
    """)

st.markdown("---")
st.caption(
    "Oblique Survey Planner v3  ·  Flat terrain  ·  Pinhole model  ·  "
    "Verified against Oblique_setup9_working_2.xls"
)
