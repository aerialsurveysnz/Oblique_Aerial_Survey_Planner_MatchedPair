"""
app.py  —  Oblique Aerial Survey Planner v3
============================================
Per-camera table, portrait/landscape, tilt-axis selector.
Three diagram views: footprint plan (all frames), cross-section, multi-strip.

Geometry verified against Oblique_setup9_working_2.xls — all values match.

Spreadsheet convention notes
─────────────────────────────
The Landscape sheet uses portrait-mounted L/R cameras (narrow axis across-track).
Our geometry.py orientation='portrait' reproduces those values exactly.
The Portrait sheet tab uses landscape-mounted oblique cameras (long axis across-track)
at a different flying height — it represents a different physical rig configuration.

Run:
    streamlit run app.py
"""

import json
import html
import math
from dataclasses import replace
from pathlib import Path
import io
from datetime import datetime

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import matplotlib.colors as mcolors
import numpy as np
import streamlit as st
from openpyxl import Workbook
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Pt
try:
    from shapely.geometry import Polygon as ShapelyPolygon, box as shapely_box
    from shapely.ops import unary_union
    from shapely.affinity import translate as shapely_translate
    SHAPELY_AVAILABLE = True
except Exception:
    ShapelyPolygon = None
    shapely_box = None
    unary_union = None
    shapely_translate = None
    SHAPELY_AVAILABLE = False

from geometry_matched_pair import (
    normalize_tilt_angle,
    half_fov_deg,
    diag_pp_to_long_edge_mm,
    flying_height_for_gsd,
    calculate_camera_solution,
    calculate_multicamera_solution,
    m_to_unit,
    unit_to_m,
)

# ─────────────────────────────────────────────────────────────────────────────
# Constants
# ─────────────────────────────────────────────────────────────────────────────

CAM_COLOURS = ["#58a6ff", "#f85149", "#3fb950", "#d29922", "#bc8cff", "#39c5cf"]

BODY_PRESETS = {
    "Sony A7R IV":       {"w_mm": 35.7,    "h_mm": 23.8,    "w_px": 9504,  "h_px": 6336},
  "Sony A7R V":        {"w_mm": 35.7, "h_mm": 23.8, "w_px": 9504, "h_px": 6336},
    "Sony A6500":        {"w_mm": 23.5,    "h_mm": 15.6,    "w_px": 6000,  "h_px": 4000},
    "Phase One iXM-100": {"w_mm": 53.4,    "h_mm": 40.0,    "w_px": 11664, "h_px": 8750},
    "Canon 5DS R":       {"w_mm": 36.0,    "h_mm": 24.0,    "w_px": 8688,  "h_px": 5792},
    "Nikon D850":        {"w_mm": 35.9,    "h_mm": 23.9,    "w_px": 8256,  "h_px": 5504},
    "Canon 760D":        {"w_mm": 22.3,    "h_mm": 14.9,    "w_px": 6000,  "h_px": 4000},
}

# Default 3-camera rig matching Oblique_setup9_working_2.xls Landscape sheet exactly
DEFAULT_CAMERAS = [
    {
        "enabled": True,
        "label": "Nadir",
        "body": "Sony A7R V",
        "focal_mm": 21.0,
        "tilt_deg": 0.0,
        "tilt_conv": "nadir",
        "orientation": "landscape",
        "tilt_axis": "across",
    },
    {
        "enabled": True,
        "label": "Right oblique",
        "body": "Sony A7R V",
        "focal_mm": 50.0,
        "tilt_deg": 50.0,
        "tilt_conv": "horiz",
        "orientation": "portrait",
        "tilt_axis": "across",
    },
    {
        "enabled": True,
        "label": "Left oblique",
        "body": "Sony A7R V",
        "focal_mm": 50.0,
        "tilt_deg": 50.0,
        "tilt_conv": "horiz",
        "orientation": "portrait",
        "tilt_axis": "across",
    },
    {
        "enabled": True,
        "label": "Fore oblique",
        "body": "Sony A7R V",
        "focal_mm": 50.0,
        "tilt_deg": 50.0,
        "tilt_conv": "horiz",
        "orientation": "portrait",
        "tilt_axis": "along",
    },
    {
        "enabled": True,
        "label": "Aft oblique",
        "body": "Sony A7R V",
        "focal_mm": 50.0,
        "tilt_deg": 50.0,
        "tilt_conv": "horiz",
        "orientation": "portrait",
        "tilt_axis": "along",
    },
]

PRESET_FILE = Path("presets.json")
SCENARIO_DIR = Path("saved_scenarios")
DEFAULT_ALTITUDE_M = 600.0
DEFAULT_SPEED_MS = 62.0
DEFAULT_SPEED_KTS = DEFAULT_SPEED_MS * 1.94384
OVERLAP_PRESETS = {
    "Standard oblique": {"forward": 80, "side": 60},
    "Balanced": {"forward": 60, "side": 40},
    "Regional": {"forward": 40, "side": 20},
    "Open terrain": {"forward": 35, "side": 15},
    "Custom": None,
}
REPORT_LOGO_CANDIDATES = [
    Path("assets/aerial_surveys_logo.png"),
    Path("aerial_surveys_logo.png"),
    Path("/mnt/data/cdb78639-5338-4c0f-88eb-fa20a9521e12.png"),
]

# ─────────────────────────────────────────────────────────────────────────────
# Page config
# ─────────────────────────────────────────────────────────────────────────────

st.set_page_config(page_title="Oblique Survey Planner", page_icon="ASL_Imagery_Icon.png", layout="wide")
st.markdown(
    """
    <style>
    .stApp {
        background: #123b6d !important;
    }

    [data-testid="stAppViewContainer"] {
        background: #123b6d !important;
    }

    [data-testid="stHeader"] {
        background: #123b6d !important;
    }
    </style>
    """,
    unsafe_allow_html=True,
)
# ─────────────────────────────────────────────────────────────────────────────
# Persistence helpers
# ─────────────────────────────────────────────────────────────────────────────

def load_saved_bodies():
    if PRESET_FILE.exists():
        try:
            return json.loads(PRESET_FILE.read_text())
        except Exception:
            pass
    return {}

def save_body_preset(name, data):
    existing = load_saved_bodies()
    existing[name] = data
    PRESET_FILE.write_text(json.dumps(existing, indent=2))

def ensure_scenario_dir():
    SCENARIO_DIR.mkdir(parents=True, exist_ok=True)
    return SCENARIO_DIR


def normalise_scenario_name(name):
    cleaned = "".join(ch if ch.isalnum() or ch in ("-", "_", " ") else "_" for ch in str(name or "").strip())
    cleaned = " ".join(cleaned.split())
    return cleaned or "my_survey"


def scenario_path_from_name(name):
    safe_name = normalise_scenario_name(name)
    if not safe_name.lower().endswith(".json"):
        safe_name = f"{safe_name}.json"
    return ensure_scenario_dir() / safe_name


def is_scenario_payload(data):
    return isinstance(data, dict) and isinstance(data.get("cameras"), list)


def save_scenario(data, name):
    path = scenario_path_from_name(name)
    path.write_text(json.dumps(data, indent=2, default=str))
    return path


def load_scenario(path_or_name):
    candidate = Path(path_or_name)
    possible_paths = []
    if candidate.exists():
        possible_paths.append(candidate)
    else:
        possible_paths.append(scenario_path_from_name(path_or_name))
        if candidate.suffix:
            possible_paths.append(Path(candidate.name))
        else:
            possible_paths.append(Path(f"{candidate.name}.json"))

    for path in possible_paths:
        try:
            data = json.loads(Path(path).read_text())
            if is_scenario_payload(data):
                return data
        except Exception:
            continue
    return None


def list_saved_scenarios():
    ensure_scenario_dir()
    found = []
    seen = set()

    def _append(path, origin):
        try:
            data = json.loads(path.read_text())
        except Exception:
            return
        if not is_scenario_payload(data):
            return
        resolved = str(path.resolve())
        if resolved in seen:
            return
        seen.add(resolved)
        label = path.name if origin == "saved_scenarios" else f"{path.name} (project root)"
        found.append({"label": label, "path": str(path), "origin": origin})

    for path in sorted(SCENARIO_DIR.glob("*.json"), key=lambda p: p.name.lower()):
        _append(path, "saved_scenarios")

    for path in sorted(Path(".").glob("*.json"), key=lambda p: p.name.lower()):
        if path.name == PRESET_FILE.name:
            continue
        _append(path, "project_root")

    return found

def all_body_names():
    return list(BODY_PRESETS.keys()) + list(load_saved_bodies().keys())

def get_body(name):
    saved = load_saved_bodies()
    if name in BODY_PRESETS:
        return BODY_PRESETS[name]
    if name in saved:
        d = saved[name]
        return {"w_mm": d["w_mm"], "h_mm": d["h_mm"], "w_px": d["w_px"], "h_px": d["h_px"]}
    return BODY_PRESETS["Sony A7R V"]
def get_inner_outer_angles(sol):
    return sol.near_angle_deg, sol.far_angle_deg


def oblique_descriptor(angle_deg: float) -> str:
    if angle_deg < 5:
        return "Nadir"
    if angle_deg < 15:
        return "Mild"
    if angle_deg < 30:
        return "Moderate"
    if angle_deg < 45:
        return "Strong"
    return "Very strong"


def format_oblique(angle_deg: float) -> str:
    return f"{angle_deg:.1f}° ({oblique_descriptor(angle_deg)})"


def mirror_solution_for_label(sol):
    """Mirror canonical positive-tilt solutions to the correct side by label."""
    label = (sol.label or "").lower()

    if sol.tilt_axis == "across" and "left" in label:
        return replace(
            sol,
            near_edge_m=-sol.near_edge_m,
            centre_m=-sol.centre_m,
            far_edge_m=-sol.far_edge_m,
            corner_near_top=(-sol.corner_near_top[0], sol.corner_near_top[1]),
            corner_near_bot=(-sol.corner_near_bot[0], sol.corner_near_bot[1]),
            corner_far_top=(-sol.corner_far_top[0], sol.corner_far_top[1]),
            corner_far_bot=(-sol.corner_far_bot[0], sol.corner_far_bot[1]),
        )

    if sol.tilt_axis == "along" and any(k in label for k in ["aft", "rear", "back"]):
        return replace(
            sol,
            near_edge_m=-sol.near_edge_m,
            centre_m=-sol.centre_m,
            far_edge_m=-sol.far_edge_m,
            corner_near_top=(sol.corner_near_top[0], -sol.corner_near_top[1]),
            corner_near_bot=(sol.corner_near_bot[0], -sol.corner_near_bot[1]),
            corner_far_top=(sol.corner_far_top[0], -sol.corner_far_top[1]),
            corner_far_bot=(sol.corner_far_bot[0], -sol.corner_far_bot[1]),
        )

    return sol
def build_export_data(solutions, mc, altitude_m, speed_ms, fwd_frac, side_frac, dist_unit="m", reciprocal=True):
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
    speed_kts = speed_ms * 1.94384
    settings_rows = [
        ["Prepared", timestamp],
        ["Altitude AGL", f"{m_to_unit(altitude_m, dist_unit):.1f} {dist_unit}"],
        ["Aircraft speed", f"{speed_ms:.1f} m/s ({speed_kts:.1f} kts)"],
        ["Forward overlap target", f"{fwd_frac * 100:.1f}%"],
        ["Sidelap target", f"{side_frac * 100:.1f}%"],
        ["Reciprocal strips", "Yes" if reciprocal else "No"],
        ["Enabled cameras", ", ".join(sol.label for _, sol, _ in solutions)],
    ]

    system_rows = []
    if mc is not None:
        line_spacing = getattr(mc, "recommended_line_spacing_m", float("nan"))
        photo_spacing = getattr(mc, "recommended_photo_spacing_m", float("nan"))
        exposure_interval = getattr(mc, "photo_interval_s", float("nan"))
        combined_swath = getattr(mc, "combined_swath_m", float("nan"))
        sidelap_achieved = getattr(mc, "sidelap_achieved", float("nan"))
        system_rows.extend([
            ["Combined swath", f"{m_to_unit(combined_swath, dist_unit):.1f} {dist_unit}" if np.isfinite(combined_swath) else "—"],
            ["Recommended line spacing", f"{m_to_unit(line_spacing, dist_unit):.1f} {dist_unit}" if np.isfinite(line_spacing) else "—"],
            ["Recommended photo spacing", f"{m_to_unit(photo_spacing, dist_unit):.1f} {dist_unit}" if np.isfinite(photo_spacing) else "—"],
            ["Exposure interval", f"{exposure_interval:.2f} s" if np.isfinite(exposure_interval) else "—"],
            ["Achieved sidelap", f"{sidelap_achieved * 100:.1f}%" if np.isfinite(sidelap_achieved) else "—"],
        ])

    representative = next((s for _, s, _ in solutions if abs(s.tilt_from_nadir_deg) > 1), solutions[0][1] if solutions else None)
    if representative is not None:
        system_rows.extend([
            [f"Inner GSD ({representative.label})", fmt_gsd(min(representative.near_gsd_m, representative.far_gsd_m))],
            ["Centre GSD", fmt_gsd(representative.centre_gsd_m)],
            [f"Outer GSD ({representative.label})", fmt_gsd(max(representative.near_gsd_m, representative.far_gsd_m))],
            [f"Obliqueness ratio ({representative.label})", f"{obliqueness_ratio(representative):.2f}×"],
        ])

    camera_rows = []
    for cam, sol, _ in solutions:
        camera_rows.append([
            cam["label"],
            round(sol.near_angle_deg, 2),
            oblique_descriptor(sol.near_angle_deg),
            round(sol.far_angle_deg, 2),
            oblique_descriptor(sol.far_angle_deg),
            round(min(sol.near_gsd_m, sol.far_gsd_m) * 100, 2),
            round(sol.centre_gsd_m * 100, 2),
            round(max(sol.near_gsd_m, sol.far_gsd_m) * 100, 2),
            round(abs(sol.near_edge_m), 2),
            round(sol.centre_m, 2),
            round(abs(sol.far_edge_m), 2),
        ])

    return settings_rows, system_rows, camera_rows


def set_table_style(table, style_name):
    try:
        table.style = style_name
    except Exception:
        pass


def fig_to_png_bytes(fig, dpi=180):
    bio = io.BytesIO()
    fig.savefig(bio, format="png", dpi=dpi, bbox_inches="tight", facecolor=fig.get_facecolor())
    bio.seek(0)
    return bio.getvalue()


def find_report_logo():
    for candidate in REPORT_LOGO_CANDIDATES:
        candidate = Path(candidate)
        if candidate.exists():
            return candidate
    return None


def make_excel_export(settings_rows, system_rows, camera_rows):
    wb = Workbook()

    ws1 = wb.active
    ws1.title = "Settings"
    for row in settings_rows:
        ws1.append(row)

    ws2 = wb.create_sheet("System Summary")
    for row in system_rows:
        ws2.append(row)

    ws3 = wb.create_sheet("Camera Results")
    ws3.append([
        "Camera",
        "Near angle (deg)",
        "Near class",
        "Far angle (deg)",
        "Far class",
        "Inner GSD (cm/px)",
        "Centre GSD (cm/px)",
        "Outer GSD (cm/px)",
        "Inner edge (m)",
        "Centre (m)",
        "Outer edge (m)",
    ])

    for row in camera_rows:
        ws3.append(row)

    bio = io.BytesIO()
    wb.save(bio)
    bio.seek(0)
    return bio.getvalue()


def make_word_export(settings_rows, system_rows, camera_rows, report_figures=None):
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    logo_path = find_report_logo()
    if logo_path is not None:
        doc.add_picture(str(logo_path), width=Inches(2.3))

    title = doc.add_paragraph()
    title_run = title.add_run("Oblique Aerial Survey Planner Report — Matched Pair")
    title_run.bold = True
    title_run.font.size = Pt(20)

    subtitle = doc.add_paragraph("Client summary of survey settings, coverage and camera geometry.")
    subtitle.runs[0].italic = True

    doc.add_heading("Summary of settings", level=2)
    table = doc.add_table(rows=1, cols=2)
    set_table_style(table, "Light List Accent 1")
    hdr = table.rows[0].cells
    hdr[0].text = "Item"
    hdr[1].text = "Value"

    for item, value in settings_rows:
        row = table.add_row().cells
        row[0].text = str(item)
        row[1].text = str(value)

    doc.add_heading("System summary", level=2)
    table_sys = doc.add_table(rows=1, cols=2)
    set_table_style(table_sys, "Light Grid Accent 1")
    hdr_sys = table_sys.rows[0].cells
    hdr_sys[0].text = "Metric"
    hdr_sys[1].text = "Value"
    for item, value in system_rows:
        row = table_sys.add_row().cells
        row[0].text = str(item)
        row[1].text = str(value)

    doc.add_heading("Camera results", level=2)
    table2 = doc.add_table(rows=1, cols=11)
    set_table_style(table2, "Medium Grid 1 Accent 1")
    hdr2 = table2.rows[0].cells
    headers = [
        "Camera",
        "Near angle",
        "Near class",
        "Far angle",
        "Far class",
        "Inner GSD",
        "Centre GSD",
        "Outer GSD",
        "Inner edge",
        "Centre",
        "Outer edge",
    ]

    for i, h in enumerate(headers):
        hdr2[i].text = h

    for row_data in camera_rows:
        row = table2.add_row().cells
        for i, value in enumerate(row_data):
            row[i].text = str(value)

    if report_figures:
        doc.add_heading("Coverage diagrams", level=2)
        for figure_title, figure_bytes in report_figures:
            doc.add_paragraph(figure_title)
            doc.add_picture(io.BytesIO(figure_bytes), width=Inches(9.2))

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

# ─────────────────────────────────────────────────────────────────────────────
# Session state
# ─────────────────────────────────────────────────────────────────────────────

if "cameras" not in st.session_state:
    st.session_state.cameras = [dict(c) for c in DEFAULT_CAMERAS]
if "planner_altitude_m" not in st.session_state:
    st.session_state.planner_altitude_m = DEFAULT_ALTITUDE_M
if "planner_dist_unit_last" not in st.session_state:
    st.session_state.planner_dist_unit_last = "m"
if "sidebar_alt_input" not in st.session_state:
    st.session_state.sidebar_alt_input = round(m_to_unit(st.session_state.planner_altitude_m, "m"), 1)
if "sidebar_speed_ms" not in st.session_state:
    st.session_state.sidebar_speed_ms = round(DEFAULT_SPEED_MS, 1)
if "sidebar_reciprocal" not in st.session_state:
    st.session_state.sidebar_reciprocal = True
if "pending_loaded_scenario" not in st.session_state:
    st.session_state.pending_loaded_scenario = None
if "saved_scenario_selector" not in st.session_state:
    st.session_state.saved_scenario_selector = None
if "saved_scenario_selector_prev" not in st.session_state:
    st.session_state.saved_scenario_selector_prev = None
if "camera_widget_nonce" not in st.session_state:
    st.session_state.camera_widget_nonce = 0


def bump_camera_widget_nonce():
    st.session_state.camera_widget_nonce = int(st.session_state.get("camera_widget_nonce", 0)) + 1

# ─────────────────────────────────────────────────────────────────────────────
# Helpers — geometry and display
# ─────────────────────────────────────────────────────────────────────────────

def fmt(v_m, unit, d=1):
    return f"{m_to_unit(v_m, unit):.{d}f} {unit}"

def fmt_gsd(v_m, d=2):
    return f"{v_m * 100:.{d}f} cm/px"

def obliqueness_ratio(sol):
    """
    Ratio of far-edge GSD to inner-edge GSD.  1.0 = nadir (no variation).
    Higher = more oblique — GSD varies more across the image.
    """
    inner_gsd = min(sol.near_gsd_m, sol.far_gsd_m)
    outer_gsd = max(sol.near_gsd_m, sol.far_gsd_m)
    if inner_gsd <= 0:
        return float("inf")
    return outer_gsd / inner_gsd

def dark_fig(w=12, h=6):
    fig, ax = plt.subplots(figsize=(w, h))
    fig.patch.set_facecolor("#0d1117")
    ax.set_facecolor("#161b22")
    ax.tick_params(colors="#8b949e", labelsize=8)
    for sp in ax.spines.values():
        sp.set_color("#30363d")
    return fig, ax

def corner_inner_outer(sol):
    """
    Return (inner_gx_or_gy, outer_gx_or_gy) where inner = edge closer to nadir.
    Works for both positive-tilt (right) and negative-tilt (left) cameras.
    """
    a, b = sol.near_edge_m, sol.far_edge_m
    if np.isfinite(a) and np.isfinite(b):
        return (a, b) if abs(a) <= abs(b) else (b, a)

    extent = polygon_extent(sol)
    if extent is not None:
        xmin, xmax, ymin, ymax = extent
        if sol.tilt_axis == "across":
            return (xmin, xmax) if abs(xmin) <= abs(xmax) else (xmax, xmin)
        return (ymin, ymax) if abs(ymin) <= abs(ymax) else (ymax, ymin)

    return 0.0, 0.0

def inner_outer_corners(sol):
    """
    Split the 4 footprint corners into (inner_top, inner_bot) and (outer_top, outer_bot).
    Inner = nadir side. Falls back to the rendered footprint polygon for near-nadir cases.
    """
    nt, nb = sol.corner_near_top, sol.corner_near_bot
    ft, fb = sol.corner_far_top, sol.corner_far_bot
    raw = [nt, nb, ft, fb]

    if all(np.isfinite(v) for xy in raw for v in xy):
        if abs(nt[0]) <= abs(ft[0]):
            return (nt, nb), (ft, fb)
        return (ft, fb), (nt, nb)

    poly = camera_polygon(sol)
    if len(poly) != 4:
        return ((0.0, 0.0), (0.0, 0.0)), ((0.0, 0.0), (0.0, 0.0))

    eps = 1e-6
    if sol.tilt_axis == "across":
        inner_abs = min(abs(p[0]) for p in poly)
        outer_abs = max(abs(p[0]) for p in poly)
        inner = [p for p in poly if abs(abs(p[0]) - inner_abs) < eps or np.isclose(abs(p[0]), inner_abs)]
        outer = [p for p in poly if abs(abs(p[0]) - outer_abs) < eps or np.isclose(abs(p[0]), outer_abs)]
        inner = sorted(inner, key=lambda p: p[1], reverse=True)[:2]
        outer = sorted(outer, key=lambda p: p[1], reverse=True)[:2]
    else:
        inner_abs = min(abs(p[1]) for p in poly)
        outer_abs = max(abs(p[1]) for p in poly)
        inner = [p for p in poly if abs(abs(p[1]) - inner_abs) < eps or np.isclose(abs(p[1]), inner_abs)]
        outer = [p for p in poly if abs(abs(p[1]) - outer_abs) < eps or np.isclose(abs(p[1]), outer_abs)]
        inner = sorted(inner, key=lambda p: p[0], reverse=True)[:2]
        outer = sorted(outer, key=lambda p: p[0], reverse=True)[:2]

    while len(inner) < 2:
        inner.append(inner[0])
    while len(outer) < 2:
        outer.append(outer[0])
    return (inner[0], inner[1]), (outer[0], outer[1])

def along_lengths_for_display(sol):
    """Inner and outer along-track (or across-track for along-tilt) footprint lengths."""
    (it, ib), (ot, ob) = inner_outer_corners(sol)
    return abs(it[1] - ib[1]), abs(ot[1] - ob[1])

def safe_corners(sol):
    """Return all 4 corner (Gx, Gy) tuples, filtered to finite values only."""
    return [c for c in camera_polygon(sol) if np.isfinite(c[0]) and np.isfinite(c[1])]

def axis_limits_from_solutions(solutions, pad=0.18):
    """
    Compute square axis limits that fit all camera footprints with a margin.
    Returns (lim,) where axes run from -lim to +lim.
    Falls back to 1.0 if no finite corners found.
    """
    all_x, all_y = [], []
    for _, sol, _ in solutions:
        for c in safe_corners(sol):
            all_x.append(c[0])
            all_y.append(c[1])
    if not all_x:
        return 1.0
    x_span = max(abs(min(all_x)), abs(max(all_x))) * (1 + pad)
    y_span = max(abs(min(all_y)), abs(max(all_y))) * (1 + pad)
    lim = max(x_span, y_span)
    return lim if np.isfinite(lim) and lim > 0 else 1.0


def camera_polygon(sol):
    corners = [
        (sol.corner_near_top[0], sol.corner_near_top[1]),
        (sol.corner_far_top[0],  sol.corner_far_top[1]),
        (sol.corner_far_bot[0],  sol.corner_far_bot[1]),
        (sol.corner_near_bot[0], sol.corner_near_bot[1]),
    ]
    if all(np.isfinite(v) for xy in corners for v in xy):
        return corners

    if abs(getattr(sol, "tilt_from_nadir_deg", 0.0)) < 0.25:
        half_across = sol.centre_slant_m * math.tan(math.radians(sol.half_fov_across_deg))
        half_along = sol.centre_slant_m * math.tan(math.radians(sol.half_fov_along_deg))
        if np.isfinite(half_across) and np.isfinite(half_along):
            return [
                (-half_across, half_along),
                (half_across, half_along),
                (half_across, -half_along),
                (-half_across, -half_along),
            ]

    return []


def polygon_extent(sol):
    poly = camera_polygon(sol)
    if not poly:
        return None
    xs = [p[0] for p in poly]
    ys = [p[1] for p in poly]
    return min(xs), max(xs), min(ys), max(ys)


def fallback_multistrip_spacing(multistrip_solutions, fwd_frac, side_frac, mc=None):
    line_spacing = getattr(mc, "recommended_line_spacing_m", float("nan")) if mc is not None else float("nan")
    photo_spacing = getattr(mc, "recommended_photo_spacing_m", float("nan")) if mc is not None else float("nan")
    used_fallback = False

    extents = [polygon_extent(sol) for _, sol, _ in multistrip_solutions]
    extents = [e for e in extents if e is not None]

    if not (np.isfinite(line_spacing) and line_spacing > 0) and extents:
        swath_width = max(e[1] for e in extents) - min(e[0] for e in extents)
        line_spacing = swath_width * max(0.05, (1.0 - side_frac))
        used_fallback = True

    if not (np.isfinite(photo_spacing) and photo_spacing > 0) and extents:
        lengths = [e[3] - e[2] for e in extents if np.isfinite(e[3] - e[2]) and (e[3] - e[2]) > 0]
        if lengths:
            photo_spacing = min(lengths) * max(0.05, (1.0 - fwd_frac))
            used_fallback = True

    if not (np.isfinite(line_spacing) and line_spacing > 0):
        line_spacing = 1.0
    if not (np.isfinite(photo_spacing) and photo_spacing > 0):
        photo_spacing = 1.0

    return line_spacing, photo_spacing, used_fallback


def matched_sidelap_band(multistrip_solutions, line_spacing, frame_y=0.0):
    """
    Return a simple cross-track overlap band for the matched reciprocal pair:
    Right oblique on the current strip versus Left oblique on the adjacent strip.
    This is intended to communicate sidelap as edge-to-edge cross-track overlap,
    not as shared plan-area intersection of skewed polygons.
    """
    right_entry = next(
        ((cam, sol, col) for cam, sol, col in multistrip_solutions
         if sol.tilt_axis == "across" and "right" in sol.label.lower()),
        None,
    )
    left_entry = next(
        ((cam, sol, col) for cam, sol, col in multistrip_solutions
         if sol.tilt_axis == "across" and "left" in sol.label.lower()),
        None,
    )
    if right_entry is None or left_entry is None:
        return None

    _, right_sol, right_col = right_entry
    _, left_sol, _ = left_entry
    right_extent = polygon_extent(right_sol)
    left_extent = polygon_extent(left_sol)
    if right_extent is None or left_extent is None:
        return None

    r_x0, r_x1, r_y0, r_y1 = right_extent
    l_x0, l_x1, l_y0, l_y1 = left_extent

    # Current strip at x offset 0; adjacent reciprocal strip at +line_spacing
    band_x0 = max(r_x0, l_x0 + line_spacing)
    band_x1 = min(r_x1, l_x1 + line_spacing)
    if not (np.isfinite(band_x0) and np.isfinite(band_x1) and band_x1 > band_x0):
        return None

    band_y0 = max(r_y0 + frame_y, l_y0 + frame_y)
    band_y1 = min(r_y1 + frame_y, l_y1 + frame_y)
    if not (np.isfinite(band_y0) and np.isfinite(band_y1) and band_y1 > band_y0):
        return None

    right_width = r_x1 - r_x0
    overlap_width = band_x1 - band_x0
    overlap_frac = overlap_width / right_width if right_width > 0 else float("nan")

    return {
        "x0": band_x0,
        "x1": band_x1,
        "y0": band_y0,
        "y1": band_y1,
        "overlap_width": overlap_width,
        "right_width": right_width,
        "overlap_frac": overlap_frac,
        "colour": right_col,
    }

def point_on_segment(px, py, ax, ay, bx, by, eps=1e-9):
    cross = (px - ax) * (by - ay) - (py - ay) * (bx - ax)
    if abs(cross) > eps:
        return False
    dot = (px - ax) * (bx - ax) + (py - ay) * (by - ay)
    if dot < -eps:
        return False
    sq_len = (bx - ax) ** 2 + (by - ay) ** 2
    if dot - sq_len > eps:
        return False
    return True


def point_in_polygon(px, py, polygon):
    inside = False
    n = len(polygon)
    for i in range(n):
        x1, y1 = polygon[i]
        x2, y2 = polygon[(i + 1) % n]
        if point_on_segment(px, py, x1, y1, x2, y2):
            return True
        intersects = ((y1 > py) != (y2 > py)) and (px < (x2 - x1) * (py - y1) / ((y2 - y1) or 1e-12) + x1)
        if intersects:
            inside = not inside
    return inside


def coverage_view_family(label):
    lower = (label or "").strip().lower()
    if "nadir" in lower:
        return "Nadir"
    if "left" in lower:
        return "Left oblique"
    if "right" in lower:
        return "Right oblique"
    if any(token in lower for token in ["fore", "forward", "front"]):
        return "Fore oblique"
    if any(token in lower for token in ["aft", "rear", "back"]):
        return "Aft oblique"
    return (label or "Camera").strip() or "Camera"


def build_coverage_sources(solutions, line_spacing, photo_spacing):
    base_polygons = []
    max_abs_x = 0.0
    max_abs_y = 0.0
    for cam, sol, _ in solutions:
        poly = camera_polygon(sol)
        if len(poly) != 4:
            continue
        max_abs_x = max(max_abs_x, max(abs(x) for x, _ in poly))
        max_abs_y = max(max_abs_y, max(abs(y) for _, y in poly))
        source = {
            "label": sol.label,
            "family": coverage_view_family(sol.label),
            "polygon": poly,
        }
        if SHAPELY_AVAILABLE:
            try:
                geom = ShapelyPolygon(poly)
                if not geom.is_valid:
                    geom = geom.buffer(0)
                source["shape"] = geom
            except Exception:
                source["shape"] = None
        base_polygons.append(source)

    if not base_polygons:
        return []

    strip_repeats = max(1, int(math.ceil(max_abs_x / max(line_spacing, 1e-6))) + 1)
    frame_repeats = max(1, int(math.ceil(max_abs_y / max(photo_spacing, 1e-6))) + 1)

    sources = []
    for base in base_polygons:
        for strip_idx in range(-strip_repeats, strip_repeats + 1):
            dx = strip_idx * line_spacing
            for frame_idx in range(-frame_repeats, frame_repeats + 1):
                dy = frame_idx * photo_spacing
                shifted = [(x + dx, y + dy) for x, y in base["polygon"]]
                xs = [x for x, _ in shifted]
                ys = [y for _, y in shifted]
                item = {
                    "label": base["label"],
                    "family": base["family"],
                    "polygon": shifted,
                    "bbox": (min(xs), max(xs), min(ys), max(ys)),
                }
                if SHAPELY_AVAILABLE and base.get("shape") is not None:
                    try:
                        item["shape"] = shapely_translate(base["shape"], xoff=dx, yoff=dy)
                    except Exception:
                        item["shape"] = None
                sources.append(item)
    return sources


def compute_exact_gap_stats(sources, line_spacing, photo_spacing):
    if not SHAPELY_AVAILABLE or not sources or line_spacing <= 0 or photo_spacing <= 0:
        return None

    try:
        repeat_cell = shapely_box(-0.5 * line_spacing, -0.5 * photo_spacing, 0.5 * line_spacing, 0.5 * photo_spacing)
        clipped_shapes = []
        for source in sources:
            geom = source.get("shape")
            if geom is None:
                geom = ShapelyPolygon(source["polygon"])
            if geom.is_empty:
                continue
            clipped = geom.intersection(repeat_cell)
            if not clipped.is_empty:
                clipped_shapes.append(clipped)
        covered = unary_union(clipped_shapes) if clipped_shapes else None
        uncovered = repeat_cell if covered is None else repeat_cell.difference(covered)
        total_area = float(repeat_cell.area)
        uncovered_area = float(uncovered.area) if uncovered is not None else total_area
        tol = max(total_area, 1.0) * 1e-9
        return {
            "available": True,
            "total_area": total_area,
            "zero_hit_area": uncovered_area,
            "zero_hit_pct": (100.0 * uncovered_area / total_area) if total_area else 0.0,
            "zero_angle_area": uncovered_area,
            "zero_angle_pct": (100.0 * uncovered_area / total_area) if total_area else 0.0,
            "has_gap": uncovered_area > tol,
        }
    except Exception:
        return None


def compute_point_coverage(solutions, line_spacing, photo_spacing, samples_x=45, samples_y=45):
    sources = build_coverage_sources(solutions, line_spacing, photo_spacing)
    if not sources:
        return None

    xs = np.linspace(-0.5 * line_spacing, 0.5 * line_spacing, samples_x)
    ys = np.linspace(-0.5 * photo_spacing, 0.5 * photo_spacing, samples_y)
    hits = np.zeros((len(ys), len(xs)), dtype=int)
    angle_counts = np.zeros((len(ys), len(xs)), dtype=int)

    for iy, y in enumerate(ys):
        for ix, x in enumerate(xs):
            labels_here = []
            families_here = set()
            for source in sources:
                xmin, xmax, ymin, ymax = source["bbox"]
                if x < xmin or x > xmax or y < ymin or y > ymax:
                    continue
                if point_in_polygon(x, y, source["polygon"]):
                    labels_here.append(source["label"])
                    families_here.add(source["family"])
            hits[iy, ix] = len(labels_here)
            angle_counts[iy, ix] = len(families_here)

    exact_gap_stats = compute_exact_gap_stats(sources, line_spacing, photo_spacing)

    return {
        "xs": xs,
        "ys": ys,
        "hits": hits,
        "angle_counts": angle_counts,
        "sources": sources,
        "line_spacing": line_spacing,
        "photo_spacing": photo_spacing,
        "exact_gap_stats": exact_gap_stats,
    }


def point_coverage_at(x, y, sources):
    labels_here = []
    families_here = []
    for source in sources:
        xmin, xmax, ymin, ymax = source["bbox"]
        if x < xmin or x > xmax or y < ymin or y > ymax:
            continue
        if point_in_polygon(x, y, source["polygon"]):
            labels_here.append(source["label"])
            families_here.append(source["family"])
    unique_families = sorted(set(families_here))
    return {
        "hits": len(labels_here),
        "unique_angles": len(unique_families),
        "labels": labels_here,
        "families": unique_families,
    }


def coverage_summary(coverage):
    hits = coverage["hits"]
    angle_counts = coverage["angle_counts"]
    return {
        "hits_min": int(np.min(hits)),
        "hits_avg": float(np.mean(hits)),
        "hits_max": int(np.max(hits)),
        "angles_min": int(np.min(angle_counts)),
        "angles_avg": float(np.mean(angle_counts)),
        "angles_max": int(np.max(angle_counts)),
    }


def coverage_gap_stats(coverage):
    hits = coverage["hits"]
    angle_counts = coverage["angle_counts"]
    sample_total = int(hits.size)
    sample_zero_hit_points = int(np.sum(hits == 0))
    sample_zero_angle_points = int(np.sum(angle_counts == 0))
    stats = {
        "sample_total_points": sample_total,
        "sample_zero_hit_points": sample_zero_hit_points,
        "sample_zero_hit_pct": (100.0 * sample_zero_hit_points / sample_total) if sample_total else 0.0,
        "sample_zero_angle_points": sample_zero_angle_points,
        "sample_zero_angle_pct": (100.0 * sample_zero_angle_points / sample_total) if sample_total else 0.0,
        "has_gap": sample_zero_hit_points > 0 or sample_zero_angle_points > 0,
        "exact_available": False,
    }

    exact_gap_stats = coverage.get("exact_gap_stats") if isinstance(coverage, dict) else None
    if exact_gap_stats is not None:
        stats.update({
            "exact_available": True,
            "zero_hit_area": exact_gap_stats.get("zero_hit_area", 0.0),
            "zero_hit_pct": exact_gap_stats.get("zero_hit_pct", 0.0),
            "zero_angle_area": exact_gap_stats.get("zero_angle_area", 0.0),
            "zero_angle_pct": exact_gap_stats.get("zero_angle_pct", 0.0),
            "has_gap": bool(exact_gap_stats.get("has_gap", False)),
        })
    else:
        stats.update({
            "zero_hit_pct": stats["sample_zero_hit_pct"],
            "zero_angle_pct": stats["sample_zero_angle_pct"],
        })
    return stats




def format_gap_pct(pct: float) -> str:
    if pct <= 0:
        return "0.00%"
    if pct < 0.01:
        return "<0.01%"
    return f"{pct:.2f}%"


def coverage_sampling_label(samples_x: int, samples_y: int) -> str:
    total = int(samples_x) * int(samples_y)
    return f"{samples_x} × {samples_y} sample grid ({total:,} points)"


def classify_gap_presentation(gap_stats: dict) -> str:
    if not gap_stats.get("has_gap"):
        return "none"
    exact_available = bool(gap_stats.get("exact_available", False))
    sample_zero_hit = float(gap_stats.get("sample_zero_hit_pct", 0.0))
    sample_zero_angle = float(gap_stats.get("sample_zero_angle_pct", 0.0))
    exact_zero_hit = float(gap_stats.get("zero_hit_pct", 0.0))
    if exact_available and exact_zero_hit < 0.01 and sample_zero_hit == 0.0 and sample_zero_angle == 0.0:
        return "micro"
    return "shortfall"

def coverage_heatmap_figure(xs, ys, grid, title, colorbar_label, probe_xy=None):
    fig, ax = dark_fig(8.8, 6.2)
    extent = [xs[0], xs[-1], ys[0], ys[-1]]

    grid_int = np.asarray(grid, dtype=int)
    max_value = int(np.max(grid_int)) if grid_int.size else 0
    display_max = max(max_value, 1)

    viridis = plt.get_cmap("viridis")
    colour_steps = ["#ff3b30"]
    for i in range(1, display_max + 1):
        frac = i / max(display_max, 1)
        colour_steps.append(viridis(frac))
    cmap = mcolors.ListedColormap(colour_steps)
    boundaries = np.arange(-0.5, display_max + 1.5, 1.0)
    norm = mcolors.BoundaryNorm(boundaries, cmap.N)

    img = ax.imshow(
        grid_int,
        origin="lower",
        extent=extent,
        aspect="auto",
        interpolation="nearest",
        cmap=cmap,
        norm=norm,
    )
    cbar = fig.colorbar(img, ax=ax, shrink=0.88, ticks=np.arange(0, display_max + 1, 1))
    cbar.set_label(colorbar_label, color="#c9d1d9", fontsize=8)
    cbar.ax.yaxis.set_tick_params(color="#8b949e")
    plt.setp(cbar.ax.get_yticklabels(), color="#8b949e", fontsize=8)
    if probe_xy is not None:
        px, py = probe_xy
        ax.scatter([px], [py], s=70, marker="x", color="#f0c040", linewidths=2.0, zorder=5)
    ax.axvline(0, color="#30363d", lw=0.9, ls="--", zorder=2)
    ax.axhline(0, color="#30363d", lw=0.9, ls="--", zorder=2)
    xt = ax.get_xticks()
    ax.set_xticklabels([f"{m_to_unit(t, dist_unit):.0f}" for t in xt], color="#8b949e")
    yt = ax.get_yticks()
    ax.set_yticklabels([f"{m_to_unit(t, dist_unit):.0f}" for t in yt], color="#8b949e")
    ax.set_xlabel(f"Across-track within repeat cell ({dist_unit})", color="#8b949e")
    ax.set_ylabel(f"Along-track within repeat cell ({dist_unit})", color="#8b949e")
    ax.set_title(title, color="#c9d1d9", fontsize=11)
    fig.tight_layout()
    return fig




def _tooltip_attr_text(text):
    return html.escape(str(text), quote=True).replace("\n", "&#10;")


def metric_with_help(container, label, value, help_text):
    escaped_label = html.escape(str(label))
    escaped_value = html.escape(str(value))
    tooltip = _tooltip_attr_text(help_text)
    container.markdown(
        f"""
        <div style="border:1px solid rgba(250,250,250,0.12); border-radius:0.75rem; padding:0.8rem 0.95rem; margin-bottom:0.35rem; background:rgba(255,255,255,0.02);">
            <div style="display:flex; align-items:center; gap:0.35rem; color:#9aa4b2; font-size:0.92rem; font-weight:600; margin-bottom:0.35rem;">
                <span>{escaped_label}</span>
                <span title="{tooltip}" style="display:inline-flex; align-items:center; justify-content:center; width:1.05rem; height:1.05rem; border-radius:999px; border:1px solid #6e7681; color:#c9d1d9; font-size:0.72rem; cursor:help; line-height:1;">?</span>
            </div>
            <div style="color:#f0f6fc; font-size:1.9rem; font-weight:700; line-height:1.1;">{escaped_value}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

def help_toggle(title, text, key, level="info"):
    col_title, col_btn = st.columns([18, 1])
    with col_title:
        st.markdown(f"##### {title}")
    with col_btn:
        toggle_key = f"show_help_{key}"
        if st.button("?", key=f"btn_{key}", help=f"Show help for {title}"):
            st.session_state[toggle_key] = not st.session_state.get(toggle_key, False)
    if st.session_state.get(f"show_help_{key}", False):
        getattr(st, level)(text)



# Apply any scenario requested on the previous run before widgets are created.
pending_loaded_scenario = st.session_state.get("pending_loaded_scenario")
if pending_loaded_scenario:
    sc = pending_loaded_scenario.get("data") or {}
    scenario_label = pending_loaded_scenario.get("label", "scenario")
    if sc and "cameras" in sc:
        loaded_altitude_m = float(sc.get("altitude_m", DEFAULT_ALTITUDE_M))
        current_dist_unit = st.session_state.get("dist_unit", st.session_state.get("planner_dist_unit_last", "m"))
        st.session_state.cameras = sc["cameras"]
        bump_camera_widget_nonce()
        st.session_state.planner_altitude_m = loaded_altitude_m
        st.session_state.sidebar_alt_input = round(m_to_unit(loaded_altitude_m, current_dist_unit), 1)
        st.session_state.sidebar_speed_ms = float(sc.get("speed_ms", DEFAULT_SPEED_MS))
        st.session_state.sidebar_reciprocal = bool(sc.get("reciprocal", True))
        st.session_state.fwd_pct = int(sc.get("fwd_overlap_pct", sc.get("forward_overlap_pct", OVERLAP_PRESETS["Balanced"]["forward"])))
        st.session_state.side_pct = int(sc.get("sidelap_pct", OVERLAP_PRESETS["Balanced"]["side"]))
        st.session_state.overlap_preset = "Custom"
        st.session_state.overlap_preset_last = "Custom"
        st.session_state.selected_scenario_label = scenario_label
        st.session_state.saved_scenario_selector = scenario_label
        st.session_state.saved_scenario_selector_prev = scenario_label
        st.session_state.scenario_flash_message = {
            "level": "success",
            "text": f"Loaded {scenario_label}",
        }
    st.session_state.pending_loaded_scenario = None

# ─────────────────────────────────────────────────────────────────────────────
# Sidebar
# ─────────────────────────────────────────────────────────────────────────────

st.sidebar.image("ASL_Imagery_Icon.png", width=85)
st.sidebar.title("Oblique Survey Planner\nMatched Pair")
st.sidebar.markdown("---")
st.sidebar.subheader("Flight Parameters")

dist_unit = st.sidebar.selectbox("Distance unit", ["m", "ft"], index=0, key="dist_unit")

if st.session_state.planner_dist_unit_last != dist_unit:
    st.session_state.sidebar_alt_input = round(m_to_unit(st.session_state.planner_altitude_m, dist_unit), 1)
    st.session_state.planner_dist_unit_last = dist_unit

alt_input = st.sidebar.number_input(
    f"Altitude AGL ({dist_unit})",
    value=float(st.session_state.sidebar_alt_input),
    min_value=1.0, max_value=m_to_unit(10000.0, dist_unit),
    step=max(1.0, m_to_unit(10.0, dist_unit)),
    key="sidebar_alt_input",
)
altitude_m = unit_to_m(alt_input, dist_unit)
st.session_state.planner_altitude_m = altitude_m

speed_ms = st.sidebar.number_input(
    "Aircraft speed (m/s)",
    value=float(st.session_state.sidebar_speed_ms),
    min_value=1.0,
    max_value=300.0,
    step=0.5,
    key="sidebar_speed_ms",
)
st.sidebar.caption(f"≈ {speed_ms * 1.94384:.1f} kts  |  {speed_ms * 3.6:.1f} km/h")
reciprocal = st.sidebar.checkbox("Reciprocal (bidirectional) strips", value=bool(st.session_state.sidebar_reciprocal), key="sidebar_reciprocal")

st.sidebar.markdown("---")
st.sidebar.subheader("Overlap Targets")

if "overlap_preset" not in st.session_state:
    st.session_state.overlap_preset = "Balanced"
if "overlap_preset_last" not in st.session_state:
    st.session_state.overlap_preset_last = st.session_state.overlap_preset
if "fwd_pct" not in st.session_state:
    st.session_state.fwd_pct = OVERLAP_PRESETS["Balanced"]["forward"]
if "side_pct" not in st.session_state:
    st.session_state.side_pct = OVERLAP_PRESETS["Balanced"]["side"]

selected_overlap_preset = st.sidebar.selectbox(
    "Overlap preset",
    list(OVERLAP_PRESETS.keys()),
    key="overlap_preset",
)

if selected_overlap_preset != st.session_state.overlap_preset_last:
    preset_values = OVERLAP_PRESETS[selected_overlap_preset]
    if preset_values is not None:
        st.session_state.fwd_pct = preset_values["forward"]
        st.session_state.side_pct = preset_values["side"]
    st.session_state.overlap_preset_last = selected_overlap_preset

fwd_pct  = st.sidebar.slider("Forward overlap (%)", 0, 95, key="fwd_pct")
side_pct = st.sidebar.slider("Sidelap (%)", 0, 95, key="side_pct")
fwd_frac  = fwd_pct  / 100.0
side_frac = side_pct / 100.0

st.sidebar.markdown("---")
st.sidebar.subheader("💾 Save / Load")

if "scenario_flash_message" in st.session_state:
    flash = st.session_state.pop("scenario_flash_message")
    flash_level = flash.get("level", "success")
    getattr(st.sidebar, flash_level, st.sidebar.success)(flash.get("text", "Scenario updated."))

sc_name = st.sidebar.text_input("Scenario name", "my_survey")
if st.sidebar.button("Save scenario"):
    saved_path = save_scenario({
        "cameras": st.session_state.cameras,
        "altitude_m": altitude_m,
        "speed_ms": speed_ms,
        "fwd_overlap_pct": fwd_pct,
        "sidelap_pct": side_pct,
        "reciprocal": reciprocal,
    }, sc_name)
    st.session_state.scenario_flash_message = {
        "level": "success",
        "text": f"Saved {saved_path.name} to {saved_path.parent.as_posix()}",
    }
    st.session_state.selected_scenario_label = saved_path.name
    st.session_state.saved_scenario_selector = saved_path.name
    st.session_state.saved_scenario_selector_prev = saved_path.name
    st.rerun()

scenario_records = list_saved_scenarios()
scenario_labels = [record["label"] for record in scenario_records]
scenario_lookup = {record["label"]: record["path"] for record in scenario_records}

if scenario_records:
    default_index = 0
    selected_label = st.session_state.get("selected_scenario_label")
    if selected_label in scenario_labels:
        default_index = scenario_labels.index(selected_label)

    if st.session_state.saved_scenario_selector not in scenario_labels:
        st.session_state.saved_scenario_selector = scenario_labels[default_index]

    load_name = st.sidebar.selectbox(
        "Saved scenarios",
        scenario_labels,
        index=scenario_labels.index(st.session_state.saved_scenario_selector),
        key="saved_scenario_selector",
    )
    selected_record = next((record for record in scenario_records if record["label"] == load_name), None)
    if selected_record and selected_record["origin"] == "project_root":
        st.sidebar.caption("Legacy scenario found in the project root. Saving again will place it into saved_scenarios.")

    if st.session_state.saved_scenario_selector_prev is None:
        st.session_state.saved_scenario_selector_prev = load_name
    elif load_name != st.session_state.saved_scenario_selector_prev:
        sc = load_scenario(scenario_lookup.get(load_name, load_name))
        st.session_state.saved_scenario_selector_prev = load_name
        if sc and "cameras" in sc:
            st.session_state.pending_loaded_scenario = {
                "label": load_name,
                "data": sc,
            }
            st.rerun()
        else:
            st.sidebar.error(f"Could not load '{load_name}'")

    st.sidebar.caption("Selecting a saved scenario loads it automatically.")
else:
    load_name = None
    st.sidebar.info("No saved scenarios found in saved_scenarios yet.")

# ─────────────────────────────────────────────────────────────────────────────
# Main header
# ─────────────────────────────────────────────────────────────────────────────
col_left_logo, col_left, col_spacer, col_right = st.columns([1.2, 5.5, 1.3, 2.0])

with col_left_logo:
    st.markdown("<div style='padding-top: 22px;'></div>", unsafe_allow_html=True)
    st.image("ASL_Imagery_Icon.png", width=110)

with col_left:
    st.title("Aerial Surveys Oblique Planner")
    st.caption(
        f"Altitude: **{fmt(altitude_m, dist_unit)}**  |  "
        f"Speed: **{speed_ms:.0f} m/s ({speed_ms * 1.94384:.0f} kts)**  |  "
        f"Fwd overlap: **{fwd_pct}%**  |  Sidelap: **{side_pct}%**"
    )

with col_right:
    st.markdown("<div style='padding-top: 18px;'></div>", unsafe_allow_html=True)
    st.image("ASL_Logo_White.png", width=200)
# ─────────────────────────────────────────────────────────────────────────────
# Camera configuration table
# ─────────────────────────────────────────────────────────────────────────────

st.subheader("Camera Configuration")

with st.expander("Orientation & tilt-axis reference", expanded=False):
    st.markdown("""
| Setting | Option | Sensor across-track | Best for |
|---|---|---|---|
| **Orientation** | Portrait | **Narrow** (short) axis | L/R oblique — limits far-edge GSD stretch |
| | Landscape | **Long** axis | Nadir — maximises swath width |
| **Tilt axis** | Across (L/R) | Tilts left/right about along-track axis | Left & Right oblique |
| | Along (F/A) | Tilts fore/aft about across-track axis | Fore & Aft oblique |

**Spreadsheet match:** The Landscape sheet in the reference spreadsheet uses portrait-mounted
L/R cameras (narrow axis across-track). The default settings here reproduce those values exactly.
For the Left camera, enter the same tilt angle as Right — the mirroring is handled automatically.
    """)

cameras      = st.session_state.cameras
bodies_avail = all_body_names()

# Column headers
hdr = st.columns([0.35, 1.5, 1.7, 0.7, 0.65, 0.8, 0.85, 1.0, 0.3])
for col, lbl in zip(hdr, ["✓", "Label", "Body", "FL mm", "Tilt °", "From", "Orient.", "Tilt axis", ""]):
    col.markdown(f"<span style='color:#8b949e;font-size:0.78em;font-weight:600'>{lbl}</span>",
                 unsafe_allow_html=True)

to_delete = None
camera_widget_nonce = st.session_state.get("camera_widget_nonce", 0)
for i, cam in enumerate(cameras):
    c_en, c_lbl, c_body, c_fl, c_tilt, c_conv, c_orient, c_axis, c_del = \
        st.columns([0.35, 1.5, 1.7, 0.7, 0.65, 0.8, 0.85, 1.0, 0.3])
    colour = CAM_COLOURS[i % len(CAM_COLOURS)]

    cam["enabled"] = c_en.checkbox("##e", value=cam["enabled"], key=f"en_{camera_widget_nonce}_{i}",
                                    label_visibility="collapsed")
    c_lbl.markdown(f"<span style='color:{colour}'>●</span>", unsafe_allow_html=True)
    cam["label"] = c_lbl.text_input("##l", value=cam["label"], key=f"lbl_{camera_widget_nonce}_{i}",
                                     label_visibility="collapsed")
    body_idx = bodies_avail.index(cam["body"]) if cam["body"] in bodies_avail else 0
    cam["body"] = c_body.selectbox("##b", bodies_avail, index=body_idx, key=f"body_{camera_widget_nonce}_{i}",
                                    label_visibility="collapsed")
    cam["focal_mm"] = float(c_fl.number_input("##f", value=float(cam["focal_mm"]),
                                               min_value=1.0, max_value=2000.0, step=1.0,
                                               key=f"fl_{camera_widget_nonce}_{i}", label_visibility="collapsed"))
    cam["tilt_deg"] = float(c_tilt.number_input("##t", value=float(cam["tilt_deg"]),
                                                  min_value=0.0, max_value=85.0, step=0.5,
                                                  key=f"tilt_{camera_widget_nonce}_{i}", label_visibility="collapsed"))
    cam["tilt_conv"] = c_conv.selectbox("##c", ["horiz", "nadir"], key=f"conv_{camera_widget_nonce}_{i}",
                                         index=0 if cam["tilt_conv"] == "horiz" else 1,
                                         label_visibility="collapsed",
                                         format_func=lambda x: "Horizontal" if x == "horiz" else "Nadir")
    cam["orientation"] = c_orient.selectbox("##o", ["portrait", "landscape"], key=f"orient_{camera_widget_nonce}_{i}",
                                             index=0 if cam["orientation"] == "portrait" else 1,
                                             label_visibility="collapsed",
                                             format_func=lambda x: "Portrait" if x == "portrait" else "Landscape")
    cam["tilt_axis"] = c_axis.selectbox("##a", ["across", "along"], key=f"axis_{camera_widget_nonce}_{i}",
                                         index=0 if cam["tilt_axis"] == "across" else 1,
                                         label_visibility="collapsed",
                                         format_func=lambda x: "Across (L/R)" if x == "across" else "Along (F/A)")
    if c_del.button("✕", key=f"del_{camera_widget_nonce}_{i}", help="Remove camera"):
        to_delete = i

if to_delete is not None:
    cameras.pop(to_delete)
    bump_camera_widget_nonce()
    st.rerun()

b1, b2, b3 = st.columns([1, 1, 6])
if b1.button("➕ Add camera"):
    cameras.append({"enabled": True, "label": f"Camera {len(cameras)+1}",
                    "body": "Custom body", "focal_mm": 50.0, "tilt_deg": 50.0,
                    "tilt_conv": "horiz", "orientation": "portrait", "tilt_axis": "across"})
    bump_camera_widget_nonce()
    st.rerun()
if b2.button("↺ Reset defaults"):
    st.session_state.cameras = [dict(c) for c in DEFAULT_CAMERAS]
    bump_camera_widget_nonce()
    st.rerun()

with st.expander("Save a camera body as a custom preset"):
    pc1, pc2 = st.columns([2, 1])
    p_base = pc1.selectbox("Base body", list(BODY_PRESETS.keys()), key="p_base")
    p_name = pc2.text_input("Preset name", value=p_base, key="p_name")
    p_wmm  = pc1.number_input("Sensor long axis mm", value=BODY_PRESETS[p_base]["w_mm"], key="p_wmm")
    p_hmm  = pc1.number_input("Sensor short axis mm", value=BODY_PRESETS[p_base]["h_mm"], key="p_hmm")
    p_wpx  = pc1.number_input("Pixel count long axis", value=BODY_PRESETS[p_base]["w_px"], step=100, key="p_wpx")
    p_hpx  = pc1.number_input("Pixel count short axis", value=BODY_PRESETS[p_base]["h_px"], step=100, key="p_hpx")
    if pc2.button("Save preset"):
        save_body_preset(p_name, {"w_mm": p_wmm, "h_mm": p_hmm,
                                   "w_px": int(p_wpx), "h_px": int(p_hpx)})
        st.success(f"Saved '{p_name}'. Reload page to use it.")

st.markdown("---")

# ─────────────────────────────────────────────────────────────────────────────
# Build camera solutions
# ─────────────────────────────────────────────────────────────────────────────

active = [c for c in cameras if c["enabled"]]

solutions = []   # list of (cam_dict, CameraSolution, colour_str)
errors    = []

for i, cam in enumerate(active):
    try:
        bd     = get_body(cam["body"])
        tilt_n = normalize_tilt_angle(cam["tilt_deg"], cam["tilt_conv"])
        sol    = calculate_camera_solution(
            altitude_m          = altitude_m,
            tilt_from_nadir_deg = tilt_n,
            sensor_w_native_mm  = bd["w_mm"],
            sensor_h_native_mm  = bd["h_mm"],
            image_w_native_px   = bd["w_px"],
            image_h_native_px   = bd["h_px"],
            focal_length_mm     = cam["focal_mm"],
            orientation         = cam["orientation"],
            tilt_axis           = cam["tilt_axis"],
            label               = cam["label"],
        )
        sol = mirror_solution_for_label(sol)
        solutions.append((cam, sol, CAM_COLOURS[i % len(CAM_COLOURS)]))
    except Exception as e:
        errors.append(f"**{cam['label']}**: {e}")

for e in errors:
    st.error(e)

if not solutions:
    st.warning("No cameras enabled. Enable at least one camera in the table above.")
    st.stop()

sol_list = [s for _, s, _ in solutions]

mc = None
try:
    mc = calculate_multicamera_solution(
        camera_solutions      = sol_list,
        arrangement           = "custom",
        altitude_m            = altitude_m,
        aircraft_speed_ms     = speed_ms,
        forward_overlap_fraction = fwd_frac,
        sidelap_fraction      = side_frac,
        reciprocal_flying     = reciprocal,
    )
except Exception as e:
    st.error(f"System calculation error: {e}")

if mc and mc.warnings:
    with st.expander("⚠️ Warnings", expanded=True):
        for w in mc.warnings:
            st.warning(w)

# ─────────────────────────────────────────────────────────────────────────────
# Summary cards
# ─────────────────────────────────────────────────────────────────────────────

st.subheader("System Summary")
if mc:
    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("Combined swath",    fmt(mc.combined_swath_m,             dist_unit))
    c2.metric("Line spacing",      fmt(mc.recommended_line_spacing_m,  dist_unit))
    c3.metric("Photo spacing",     fmt(mc.recommended_photo_spacing_m, dist_unit))
    c4.metric("Exposure interval", f"{mc.photo_interval_s:.2f} s")
    c5.metric("Sidelap achieved",  f"{mc.sidelap_achieved * 100:.1f}%")

    rep = next((s for _, s, _ in solutions if abs(s.tilt_from_nadir_deg) > 1), sol_list[0])
    c6, c7, c8, c9 = st.columns(4)
    c6.metric(f"GSD inner ({rep.label})", fmt_gsd(min(rep.near_gsd_m, rep.far_gsd_m)))
    c7.metric("GSD centre",               fmt_gsd(rep.centre_gsd_m))
    c8.metric("GSD outer",                fmt_gsd(max(rep.near_gsd_m, rep.far_gsd_m)))
    c9.metric("Fwd overlap near/ctr/far",
              f"{mc.forward_overlap_near*100:.0f}% / "
              f"{mc.forward_overlap_centre*100:.0f}% / "
              f"{mc.forward_overlap_far*100:.0f}%")

st.markdown("---")

# ─────────────────────────────────────────────────────────────────────────────
# Per-camera results table  (includes obliqueness ratio)
# ─────────────────────────────────────────────────────────────────────────────

st.subheader("Per-Camera Results")

rows = []
for cam, sol, _ in solutions:
    inner_gx, outer_gx = corner_inner_outer(sol)
    inner_len, outer_len = along_lengths_for_display(sol)
    inner_gsd = min(sol.near_gsd_m, sol.far_gsd_m)
    outer_gsd = max(sol.near_gsd_m, sol.far_gsd_m)
    tilt_h = 90.0 - sol.tilt_from_nadir_deg

    rows.append({
    "Camera":                    sol.label,
    "Body / FL":                 f"{cam['body']}  {cam['focal_mm']:.0f} mm",
    "Orient.":                   sol.orientation,
    "Tilt axis":                 sol.tilt_axis,
    "Tilt nadir °":              round(sol.tilt_from_nadir_deg, 1),
    "Tilt horiz °":              round(tilt_h, 1),

    # ── Oblique angles ──
    "Near oblique angle":        format_oblique(sol.near_angle_deg),
    "Far oblique angle":         format_oblique(sol.far_angle_deg),
    # ─────────────────────

    "Pixel µm":                  round(sol.pixel_size_mm * 1000, 2),
    "FOV across °":              round(sol.full_fov_across_deg, 2),
    "FOV along °":               round(sol.full_fov_along_deg,  2),
    f"Inner edge ({dist_unit})": round(m_to_unit(abs(inner_gx), dist_unit), 1),
    f"Outer edge ({dist_unit})": round(m_to_unit(abs(outer_gx), dist_unit), 1),
    f"Inner length ({dist_unit})": round(m_to_unit(inner_len, dist_unit), 1),
    f"Outer length ({dist_unit})": round(m_to_unit(outer_len, dist_unit), 1),
    "Inner GSD cm":              round(inner_gsd * 100, 3),
    "Centre GSD cm":             round(sol.centre_gsd_m * 100, 3),
    "Outer GSD cm":              round(outer_gsd * 100, 3),
    "Inner slant m":             round(min(sol.near_slant_m, sol.far_slant_m), 1),
    "Outer slant m":             round(max(sol.near_slant_m, sol.far_slant_m), 1),
    "Diag image mm":             round(sol.diag_image_mm, 4),
})

st.dataframe(rows, width="stretch")
st.caption(
    "**Near/Far oblique angle** = angle from vertical (nadir) at the inner and outer image edges. "
    "Higher angles mean a stronger oblique view."
)
st.markdown("---")

settings_rows, system_rows, camera_rows = build_export_data(
    solutions=solutions,
    mc=mc,
    altitude_m=altitude_m,
    speed_ms=speed_ms,
    fwd_frac=fwd_frac,
    side_frac=side_frac,
    dist_unit=dist_unit,
    reciprocal=reciprocal,
)

st.markdown("---")

# then your first diagram section starts below here

# ─────────────────────────────────────────────────────────────────────────────
# DIAGRAM 1 — Footprint plan view — ALL cameras, ALL frames shown
# ─────────────────────────────────────────────────────────────────────────────

st.subheader("Footprint Plan View — All Cameras")
st.caption(
    "**Near/Far oblique angle** = angle from vertical (nadir) at the inner and outer image edges. "
    "Higher angles mean a stronger oblique view."
)

lim = axis_limits_from_solutions(solutions)

fig_fp, ax_fp = dark_fig(w=10, h=10)
ax_fp.set_aspect("equal")
ax_fp.set_xlim(-lim, lim)
ax_fp.set_ylim(-lim, lim)

# Grid / nadir
ax_fp.axhline(0, color="#21262d", lw=0.8, zorder=1)
ax_fp.axvline(0, color="#21262d", lw=0.8, zorder=1)
ax_fp.scatter([0], [0], s=200, color="#f0c040", zorder=8, marker="x", linewidths=2.5)
ax_fp.annotate("Nadir", (0, 0), xytext=(8, -14), textcoords="offset points",
               color="#f0c040", fontsize=8.5, fontweight="bold")

for cam, sol, colour in solutions:
    corners_xy = camera_polygon(sol)
    if len(corners_xy) != 4:
        continue

    poly = plt.Polygon(
        corners_xy,
        closed=True,
        facecolor=colour,
        alpha=0.18,
        edgecolor=colour,
        linewidth=2.2,
        zorder=3,
    )
    ax_fp.add_patch(poly)
    ax_fp.scatter(*zip(*corners_xy), color=colour, s=30, zorder=5, edgecolors="none")

    cx = sum(c[0] for c in corners_xy) / 4
    cy = sum(c[1] for c in corners_xy) / 4
    ax_fp.text(
        cx,
        cy,
        sol.label,
        color=colour,
        fontsize=8,
        fontweight="bold",
        ha="center",
        va="center",
        zorder=7,
        bbox=dict(facecolor="#0d1117", alpha=0.70, pad=2.5, edgecolor=colour, linewidth=0.8, boxstyle="round,pad=0.3"),
    )

    inner_gx, outer_gx = corner_inner_outer(sol)
    inner_len, outer_len = along_lengths_for_display(sol)
    (it, ib), (ot, ob) = inner_outer_corners(sol)
    label_box = dict(facecolor="#0d1117", alpha=0.82, edgecolor=colour, linewidth=0.6, boxstyle="round,pad=0.18")

    if sol.tilt_axis == "across":
        ap = dict(arrowstyle="-|>", lw=1.3, mutation_scale=11)
        ax_fp.annotate("", xy=(inner_gx, 0), xytext=(0, 0), arrowprops={**ap, "color": colour}, zorder=4)
        ax_fp.annotate("", xy=(outer_gx, 0), xytext=(inner_gx, 0), arrowprops={**ap, "color": colour}, zorder=4)

        dy = lim * 0.042
        direction = -1 if outer_gx >= 0 else 1
        ax_fp.annotate(
            f"{m_to_unit(abs(inner_gx), dist_unit):.0f} {dist_unit}",
            xy=(inner_gx / 2, 0),
            xytext=(0, 12 * direction),
            textcoords="offset points",
            color=colour,
            fontsize=7,
            ha="center",
            va="bottom" if direction > 0 else "top",
            bbox=label_box,
        )
        ax_fp.annotate(
            f"{m_to_unit(abs(outer_gx), dist_unit):.0f} {dist_unit}",
            xy=((inner_gx + outer_gx) / 2, 0),
            xytext=(0, 28 * direction),
            textcoords="offset points",
            color=colour,
            fontsize=7,
            ha="center",
            va="bottom" if direction > 0 else "top",
            bbox=label_box,
        )

        ixm = (it[0] + ib[0]) / 2
        iym = (it[1] + ib[1]) / 2
        dx_off = -11 if inner_gx >= 0 else 11
        ax_fp.annotate(
            f"{m_to_unit(inner_len, dist_unit):.0f} {dist_unit}",
            xy=(ixm, iym),
            xytext=(dx_off, 0),
            textcoords="offset points",
            color=colour,
            fontsize=6.5,
            va="center",
            ha="right" if dx_off < 0 else "left",
            bbox=label_box,
        )

        oxm = (ot[0] + ob[0]) / 2
        oym = (ot[1] + ob[1]) / 2
        dx_off2 = 11 if outer_gx >= 0 else -11
        ax_fp.annotate(
            f"{m_to_unit(outer_len, dist_unit):.0f} {dist_unit}",
            xy=(oxm, oym),
            xytext=(dx_off2, 0),
            textcoords="offset points",
            color=colour,
            fontsize=6.5,
            va="center",
            ha="left" if dx_off2 > 0 else "right",
            bbox=label_box,
        )

        inner_gsd = min(sol.near_gsd_m, sol.far_gsd_m)
        outer_gsd = max(sol.near_gsd_m, sol.far_gsd_m)
        ax_fp.annotate(
            f"{inner_gsd*100:.2f} cm/px",
            xy=(ixm, iym),
            xytext=(0, 12),
            textcoords="offset points",
            color=colour,
            fontsize=5.5,
            alpha=0.9,
            ha="center",
            va="bottom",
        )
        ax_fp.annotate(
            f"{outer_gsd*100:.2f} cm/px",
            xy=(oxm, oym),
            xytext=(0, 12),
            textcoords="offset points",
            color=colour,
            fontsize=5.5,
            alpha=0.9,
            ha="center",
            va="bottom",
        )

    else:
        ap = dict(arrowstyle="-|>", lw=1.3, mutation_scale=11)
        inner_gy = sol.near_edge_m if abs(sol.near_edge_m) < abs(sol.far_edge_m) else sol.far_edge_m
        outer_gy = sol.far_edge_m if abs(sol.far_edge_m) > abs(sol.near_edge_m) else sol.near_edge_m
        if not (np.isfinite(inner_gy) and np.isfinite(outer_gy)):
            extent = polygon_extent(sol)
            if extent is not None:
                _, _, ymin, ymax = extent
                inner_gy, outer_gy = (ymin, ymax) if abs(ymin) <= abs(ymax) else (ymax, ymin)
            else:
                inner_gy, outer_gy = 0.0, 0.0
        ax_fp.annotate("", xy=(0, inner_gy), xytext=(0, 0), arrowprops={**ap, "color": colour}, zorder=4)
        ax_fp.annotate("", xy=(0, outer_gy), xytext=(0, inner_gy), arrowprops={**ap, "color": colour}, zorder=4)
        dx = lim * 0.035
        ax_fp.annotate(
            f"{m_to_unit(abs(inner_gy), dist_unit):.0f} {dist_unit}",
            xy=(0, inner_gy / 2),
            xytext=(12, 0),
            textcoords="offset points",
            color=colour,
            fontsize=7,
            ha="left",
            va="center",
            bbox=label_box,
        )
        ax_fp.annotate(
            f"{m_to_unit(abs(outer_gy), dist_unit):.0f} {dist_unit}",
            xy=(0, (inner_gy + outer_gy) / 2),
            xytext=(12, 18 if outer_gy >= 0 else -18),
            textcoords="offset points",
            color=colour,
            fontsize=7,
            ha="left",
            va="center",
            bbox=label_box,
        )
# Tick labels in display units
xt = ax_fp.get_xticks()
ax_fp.set_xticklabels([f"{m_to_unit(t, dist_unit):.0f}" for t in xt], color="#8b949e")
yt = ax_fp.get_yticks()
ax_fp.set_yticklabels([f"{m_to_unit(t, dist_unit):.0f}" for t in yt], color="#8b949e")
ax_fp.set_xlabel(f"← Across-track ({dist_unit}) →", color="#8b949e")
ax_fp.set_ylabel(f"↑ Along-track / forward ({dist_unit})", color="#8b949e")
ax_fp.set_title("Single-Frame Footprint Plan View — All Cameras", color="#c9d1d9",
                fontsize=11, pad=10)

legend_fp = [mpatches.Patch(color=col, label=cam["label"], alpha=0.85)
             for cam, _, col in solutions]
ax_fp.legend(handles=legend_fp, loc="upper right", fontsize=8, framealpha=0.4,
             labelcolor="#c9d1d9", facecolor="#161b22")

# Forward arrow
ax_fp.annotate("", xy=(0, lim * 0.90), xytext=(0, lim * 0.74),
               arrowprops=dict(arrowstyle="->", color="#c9d1d9", lw=1.8))
ax_fp.text(lim * 0.03, lim * 0.82, "Fwd", color="#c9d1d9", fontsize=7.5, va="center")

fig_fp.tight_layout()
st.pyplot(fig_fp)
st.markdown("---")

# ─────────────────────────────────────────────────────────────────────────────
# DIAGRAM 2 — Cross-section (across-track cameras)
# ─────────────────────────────────────────────────────────────────────────────

across_sols = [(c, s, col) for c, s, col in solutions if s.tilt_axis == "across"]

st.subheader("Cross-Section View — Across-Track")
st.caption(
    "Side view looking forward. "
    "Dotted line = inner edge ray, solid = centre ray, dashed = outer edge ray. "
    "GSD values shown at ground intercepts."
)

if across_sols:
    fig_xs, ax_xs = dark_fig(12, 5)
    H = altitude_m

    ax_xs.axvline(0, color="#30363d", lw=1.0, ls="--", zorder=1)
    ax_xs.axhline(0, color="#444", lw=1.5, zorder=1)
    ax_xs.scatter([0], [H], marker="^", s=220, color="#f0c040", zorder=7)
    ax_xs.annotate(f"{fmt(H, dist_unit)} AGL", (0, H), xytext=(10, -2),
                   textcoords="offset points", color="#f0c040", fontsize=8, fontweight="bold")

    candidate_x = [0.0]

    for cam, sol, colour in across_sols:
        inner_gx, outer_gx = corner_inner_outer(sol)
        centre_gx = sol.centre_m if np.isfinite(sol.centre_m) else 0.0
        if not (np.isfinite(inner_gx) and np.isfinite(outer_gx)):
            extent = polygon_extent(sol)
            if extent is not None:
                xmin, xmax, _, _ = extent
                inner_gx, outer_gx = (xmin, xmax) if abs(xmin) <= abs(xmax) else (xmax, xmin)
                centre_gx = 0.5 * (xmin + xmax)
        inner_gsd = min(sol.near_gsd_m, sol.far_gsd_m)
        outer_gsd = max(sol.near_gsd_m, sol.far_gsd_m)
        tilt_h = 90.0 - sol.tilt_from_nadir_deg
        near_a, far_a = get_inner_outer_angles(sol)

        for gx, ls, lw in [(inner_gx, ":", 1.4), (centre_gx, "-", 2.0), (outer_gx, "--", 1.4)]:
            if np.isfinite(gx):
                ax_xs.plot([0, gx], [H, 0], color=colour, ls=ls, lw=lw, zorder=3)
                candidate_x.append(gx)

        xs = [gx for gx in [inner_gx, centre_gx, outer_gx] if np.isfinite(gx)]
        if xs:
            ax_xs.scatter(xs, [0] * len(xs), color=colour, s=50, zorder=5, edgecolors="none")

        if np.isfinite(inner_gx):
            ax_xs.text(inner_gx, -H * 0.04,
                       f"{sol.label}\n{m_to_unit(abs(inner_gx), dist_unit):.0f} {dist_unit}\nnear {near_a:.1f}°",
                       color=colour, fontsize=6, ha="center", va="top")
            ax_xs.text(inner_gx * 0.55, H * 0.18,
                       f"{inner_gsd*100:.1f} cm/px",
                       color=colour, fontsize=6, ha="center", alpha=0.9)

        if np.isfinite(outer_gx):
            ax_xs.text(outer_gx, -H * 0.09,
                       f"{m_to_unit(abs(outer_gx), dist_unit):.0f} {dist_unit}\nfar {far_a:.1f}°",
                       color=colour, fontsize=6, ha="center", va="top")
            ax_xs.text(outer_gx * 0.55, H * 0.18,
                       f"{outer_gsd*100:.1f} cm/px",
                       color=colour, fontsize=6, ha="center", alpha=0.9)

        if np.isfinite(centre_gx):
            ax_xs.text(centre_gx * 0.5, H * 0.55,
                       f"tilt {tilt_h:.1f}° from horiz",
                       color=colour, fontsize=6, ha="center", alpha=0.75, style="italic")

    valid_x = [abs(v) for v in candidate_x if np.isfinite(v)]
    max_x = max(valid_x) * 1.25 if valid_x else 1.0
    if not np.isfinite(max_x) or max_x <= 0:
        max_x = 1.0
    ax_xs.set_xlim(-max_x, max_x)
    ax_xs.set_ylim(-H * 0.18, H * 1.15)

    xt = ax_xs.get_xticks()
    ax_xs.set_xticklabels([f"{m_to_unit(t, dist_unit):.0f}" for t in xt], color="#8b949e")
    yt = ax_xs.get_yticks()
    ax_xs.set_yticklabels([f"{m_to_unit(t, dist_unit):.0f}" for t in yt], color="#8b949e")
    ax_xs.set_xlabel(f"← Across-track ({dist_unit}) →", color="#8b949e")
    ax_xs.set_ylabel(f"Altitude ({dist_unit})", color="#8b949e")
    ax_xs.set_title("Cross-Section — Across-Track Cameras", color="#c9d1d9", fontsize=11)

    legend_xs = [mpatches.Patch(color=col, label=cam["label"], alpha=0.8)
                 for cam, _, col in across_sols]
    legend_xs += [
        mpatches.Patch(color="white", alpha=0, label=""),
        plt.Line2D([0],[0], color="white", ls=":", lw=1.4, label="Inner edge ray"),
        plt.Line2D([0],[0], color="white", ls="-", lw=2.0, label="Centre ray"),
        plt.Line2D([0],[0], color="white", ls="--", lw=1.4, label="Outer edge ray"),
    ]
    ax_xs.legend(handles=legend_xs, fontsize=7, framealpha=0.35,
                 labelcolor="#c9d1d9", facecolor="#161b22", loc="upper right", ncol=2)

    fig_xs.tight_layout()
    st.pyplot(fig_xs)
else:
    st.info("No across-track cameras enabled.")
st.markdown("---")

# ─────────────────────────────────────────────────────────────────────────────
# DIAGRAM 3 — Multi-strip plan view
# ─────────────────────────────────────────────────────────────────────────────

st.subheader("Multi-Strip Plan View")
show_nadir_multistrip = st.checkbox(
    "Include nadir footprint in multi-strip view",
    value=True,
    key="show_nadir_multistrip",
)
st.caption(
    "Three adjacent flight strips offset by the recommended line spacing. "
    "Each strip shows two successive frame instances so both sidelap and forward overlap are visible. "
    "The white hatched band shows the matched-frame cross-track sidelap between the right oblique on one strip "
    "and the left oblique on the adjacent reciprocal strip."
)

multistrip_solutions = [
    (cam, sol, col)
    for cam, sol, col in solutions
    if show_nadir_multistrip or "nadir" not in sol.label.lower()
]

if mc and multistrip_solutions:
    fig_ms, ax_ms = dark_fig(14, 8)
    ax_ms.set_aspect("equal")

    line_spacing, photo_spacing, used_spacing_fallback = fallback_multistrip_spacing(
        multistrip_solutions,
        fwd_frac=fwd_frac,
        side_frac=side_frac,
        mc=mc,
    )
    if used_spacing_fallback:
        st.info("Using geometric fallback spacing in the multi-strip view because one or more system spacing values were non-finite.")

    n_strips = 3
    strip_alphas = [0.55, 0.38, 0.22]
    strip_lws = [1.6, 0.9, 0.5]
    strip_cols = ["#c9d1d9", "#6e7681", "#444"]
    strip_positions = [((i - (n_strips // 2)) * line_spacing) for i in range(n_strips)]
    frame_offsets_y = [0.0, photo_spacing]

    all_x_ms = []
    all_y_ms = []

    for si, x_off in enumerate(strip_positions):
        for frame_idx, y_off in enumerate(frame_offsets_y):
            for cam, sol, colour in multistrip_solutions:
                base_corners = camera_polygon(sol)
                if len(base_corners) != 4:
                    continue
                corners = [(x + x_off, y + y_off) for x, y in base_corners]
                alpha = strip_alphas[si] if frame_idx == 0 else strip_alphas[si] * 0.45
                poly = plt.Polygon(
                    corners,
                    closed=True,
                    facecolor=colour,
                    alpha=alpha,
                    edgecolor=colour,
                    linewidth=strip_lws[si],
                    zorder=3,
                )
                ax_ms.add_patch(poly)
                for x, y in corners:
                    all_x_ms.append(x)
                    all_y_ms.append(y)

        ax_ms.axvline(x_off, color=strip_cols[si], lw=0.9, ls="-.", alpha=0.7, zorder=2)

    if all_y_ms:
        top_y = max(all_y_ms)
        for si, x_off in enumerate(strip_positions):
            ax_ms.text(
                x_off,
                top_y * 1.02,
                f"Strip {si+1}",
                color=strip_cols[si],
                fontsize=8,
                ha="center",
                va="bottom",
            )

    sidelap_band = matched_sidelap_band(multistrip_solutions, line_spacing=line_spacing, frame_y=0.0)
    if sidelap_band is not None:
        hatch_rect = mpatches.Rectangle(
            (sidelap_band["x0"], sidelap_band["y0"]),
            sidelap_band["x1"] - sidelap_band["x0"],
            sidelap_band["y1"] - sidelap_band["y0"],
            facecolor=(1, 1, 1, 0.06),
            edgecolor="white",
            hatch="////",
            linewidth=1.1,
            zorder=4,
        )
        ax_ms.add_patch(hatch_rect)
        label_x = 0.5 * (sidelap_band["x0"] + sidelap_band["x1"])
        label_y = sidelap_band["y1"] + 0.03 * max((max(all_y_ms) - min(all_y_ms)), 1.0)
        band_pct = sidelap_band["overlap_frac"] * 100 if np.isfinite(sidelap_band["overlap_frac"]) else mc.sidelap_achieved * 100
        ax_ms.text(
            label_x,
            label_y,
            f"Matched R/L sidelap\n{band_pct:.0f}%",
            color="white",
            fontsize=7,
            ha="center",
            va="bottom",
            bbox=dict(facecolor="#21262d", alpha=0.80, pad=3, edgecolor="#c9d1d9", lw=0.6),
            zorder=6,
        )

    if all_x_ms and all_y_ms:
        y_range = max(all_y_ms) - min(all_y_ms)
        x_range = max(all_x_ms) - min(all_x_ms)
        line_arrow_y = min(all_y_ms) - 0.10 * y_range
        ax_ms.annotate(
            "",
            xy=(line_spacing, line_arrow_y),
            xytext=(0, line_arrow_y),
            arrowprops=dict(arrowstyle="<->", color="white", lw=2.0, mutation_scale=13),
            zorder=5,
        )
        ax_ms.text(
            line_spacing / 2,
            line_arrow_y - 0.03 * y_range,
            f"Line spacing\n{m_to_unit(line_spacing, dist_unit):.0f} {dist_unit}",
            color="white",
            fontsize=8,
            ha="center",
            va="top",
            bbox=dict(facecolor="#21262d", alpha=0.75, pad=3, edgecolor="#555", lw=0.5),
        )

        x_for_photo = min(all_x_ms) - 0.06 * x_range
        ax_ms.annotate(
            "",
            xy=(x_for_photo, photo_spacing),
            xytext=(x_for_photo, 0),
            arrowprops=dict(arrowstyle="<->", color="white", lw=2.0, mutation_scale=13),
            zorder=5,
        )
        ax_ms.text(
            x_for_photo + 0.03 * x_range,
            photo_spacing / 2,
            f"Photo spacing\n{m_to_unit(photo_spacing, dist_unit):.0f} {dist_unit}",
            color="white",
            fontsize=8,
            ha="left",
            va="center",
            bbox=dict(facecolor="#21262d", alpha=0.75, pad=3, edgecolor="#555", lw=0.5),
        )

        x_min = min(all_x_ms) - 0.12 * x_range
        x_max = max(all_x_ms) + 0.12 * x_range
        y_min = min(all_y_ms) - 0.16 * y_range
        y_max = max(all_y_ms) + 0.10 * y_range
    else:
        x_min, x_max, y_min, y_max = -1, 1, -1, 1

    ax_ms.set_xlim(x_min, x_max)
    ax_ms.set_ylim(y_min, y_max)

    xt = ax_ms.get_xticks()
    ax_ms.set_xticklabels([f"{m_to_unit(t, dist_unit):.0f}" for t in xt], color="#8b949e")
    yt = ax_ms.get_yticks()
    ax_ms.set_yticklabels([f"{m_to_unit(t, dist_unit):.0f}" for t in yt], color="#8b949e")
    ax_ms.set_xlabel(f"← Across-track ({dist_unit}) →", color="#8b949e")
    ax_ms.set_ylabel(f"↑ Along-track ({dist_unit})", color="#8b949e")
    ax_ms.set_title("Multi-Strip Plan View (3 strips, 2 frame instances)", color="#c9d1d9", fontsize=11)

    legend_ms = [mpatches.Patch(color=col, label=cam["label"], alpha=0.8)
                 for cam, _, col in multistrip_solutions]
    ax_ms.legend(handles=legend_ms, fontsize=8, framealpha=0.35,
                 labelcolor="#c9d1d9", facecolor="#161b22", loc="upper right")

    fig_ms.tight_layout()
    st.pyplot(fig_ms)

elif not mc:
    st.info("System solution unavailable.")
else:
    st.info("No cameras enabled for the multi-strip view.")

st.markdown("---")

# ─────────────────────────────────────────────────────────────────────────────
# DIAGRAM 4 — Point coverage / views per point
# ─────────────────────────────────────────────────────────────────────────────

coverage_fig_hits = None
coverage_fig_angles = None
coverage_result = None
coverage_probe = None

help_toggle(
    "Point Coverage / Views Per Point",
    "This section samples one repeating survey cell to estimate how many times a ground point is captured within the pattern. Image hits count every photo that sees the point. Distinct viewing angles count unique camera directions such as nadir, left oblique, right oblique, fore oblique and aft oblique. Use the min, average and max together because coverage is not uniform across the block. The min/average/max figures and heatmaps are sampled across one repeat cell, while the red gap warning below now also uses an exact geometric zero-coverage check when Shapely is available. Standard precision means a 61 × 61 sample grid, which is 3,721 tested points across one repeat cell.",
    key="point_coverage_intro",
)

coverage_precision_options = {
    "Fast preview": (31, 31),
    "Standard": (61, 61),
    "High precision": (101, 101),
}
coverage_precision = st.selectbox(
    "Coverage sampling precision",
    list(coverage_precision_options.keys()),
    index=1,
    help="Higher precision samples more points inside the repeating cell and reduces the chance of missing very narrow gaps. It is slower, but more defensible when checking low-overlap setups.",
)
samples_x, samples_y = coverage_precision_options[coverage_precision]
sample_grid_label = coverage_sampling_label(samples_x, samples_y)
st.caption(f"{coverage_precision}: {sample_grid_label} tested across one repeat cell.")

line_spacing_pc, photo_spacing_pc, used_pc_fallback = fallback_multistrip_spacing(
    solutions,
    fwd_frac=fwd_frac,
    side_frac=side_frac,
    mc=mc,
)

coverage_result = compute_point_coverage(
    solutions,
    line_spacing=line_spacing_pc,
    photo_spacing=photo_spacing_pc,
    samples_x=samples_x,
    samples_y=samples_y,
)

if coverage_result is not None:
    if used_pc_fallback:
        st.info("Using geometric fallback spacing for the point-coverage estimate because one or more system spacing values were non-finite.")

    summary = coverage_summary(coverage_result)
    gap_stats = coverage_gap_stats(coverage_result)

    help_toggle(
        "How to read the summary",
        "Minimum = the weakest sampled point inside the repeating cell. Average = a typical sampled point. Maximum = the strongest sampled overlap zone. For client conversations, the safest wording is usually to quote the range and then mention the typical value, for example: each point gets between X and Y image hits, with an average of Z. The heatmaps and min/avg/max remain sampled, but the zero-coverage warning below now also uses an exact geometric gap test when available. Standard precision uses a 61 × 61 sample grid, not a 61 m by 61 m area.",
        key="point_coverage_summary",
    )

    gap_mode = classify_gap_presentation(gap_stats)

    if gap_mode == "shortfall":
        warning_title = "⚠ COVERAGE SHORTFALL DETECTED"
        if gap_stats["exact_available"]:
            warning_summary = (
                f"Exact zero-coverage gap present. {format_gap_pct(gap_stats['zero_hit_pct'])} of the repeat-cell area has no image coverage."
            )
            warning_detail = (
                f"The heatmaps and min/avg/max values are still sampled at {coverage_precision.lower()} "
                f"using a {sample_grid_label}, but the gap warning itself is based on an exact geometric check. "
                f"For reference, {gap_stats['sample_zero_hit_pct']:.1f}% of sampled points had zero image hits and "
                f"{gap_stats['sample_zero_angle_pct']:.1f}% had zero viewing angles."
            )
        else:
            warning_summary = (
                f"Sampled coverage gaps detected at {coverage_precision.lower()} using a {sample_grid_label}."
            )
            warning_detail = (
                f"{gap_stats['sample_zero_hit_pct']:.1f}% of sampled points had zero image hits, and "
                f"{gap_stats['sample_zero_angle_pct']:.1f}% had zero viewing angles. "
                "This warning is currently based on the sampled repeat-cell check."
            )
        st.markdown(
            f"""
            <div style="background:rgba(248,81,73,0.18); border:2px solid rgba(248,81,73,0.95); border-left:10px solid #f85149; border-radius:0.95rem; padding:1rem 1.15rem; margin:0.45rem 0 1rem 0; box-shadow:0 0 0 1px rgba(248,81,73,0.18) inset;">
                <div style="font-size:1.42rem; font-weight:900; letter-spacing:0.02em; color:#ffb3ad; margin-bottom:0.28rem;">{warning_title}</div>
                <div style="font-size:1.06rem; font-weight:800; color:#f0f6fc; margin-bottom:0.28rem;">{warning_summary}</div>
                <div style="font-size:0.98rem; font-weight:600; color:#f0f6fc; line-height:1.45;">{warning_detail}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    elif gap_mode == "micro":
        st.warning(
            f"Possible microscopic exact-gap sliver detected: {format_gap_pct(gap_stats['zero_hit_pct'])} of the repeat-cell area. "
            f"The exact geometric check found a tiny uncovered sliver, but the {sample_grid_label} found 0.0% zero-hit and 0.0% zero-angle sampled points. "
            "Treat this as a vanishingly small seam or numerical precision edge unless it persists at higher precision."
        )
    else:
        if gap_stats["exact_available"]:
            st.success(
                f"No exact zero-coverage gaps detected inside the repeat cell. "
                f"The heatmaps and min/avg/max below are still sampled at {coverage_precision.lower()} using a {sample_grid_label}."
            )
        else:
            st.success(
                f"No sampled gaps detected at {coverage_precision.lower()} using a {sample_grid_label}. "
                "This is a sampled check of the repeat cell, not a formal geometric proof."
            )

    m1, m2, m3 = st.columns(3)
    metric_with_help(
        m1,
        "Min image hits",
        f"{summary['hits_min']}",
        "The weakest sampled point inside the repeating survey cell. If this is zero, the sampled check found a coverage gap somewhere under the current settings. For client conversations, this is the safest number to describe as the minimum image count per point.",
    )
    metric_with_help(
        m2,
        "Average image hits",
        f"{summary['hits_avg']:.1f}",
        "A typical sampled point across the repeating survey cell. This is useful for explaining the general level of redundancy, but it is not a guarantee because some points will sit in weaker or stronger overlap zones.",
    )
    metric_with_help(
        m3,
        "Max image hits",
        f"{summary['hits_max']}",
        "The strongest sampled overlap zone inside the repeating survey cell. This often occurs where strip overlap and forward overlap stack together. It is helpful for understanding peak redundancy, but it should not be presented as typical coverage.",
    )

    m4, m5, m6 = st.columns(3)
    metric_with_help(
        m4,
        "Min viewing angles",
        f"{summary['angles_min']}",
        "The weakest directional diversity at any sampled point. A value of zero means there is no sampled coverage at that location. A value of one means the point is seen from only one camera family or direction. This is the conservative number to quote when clients ask about minimum angular diversity.",
    )
    metric_with_help(
        m5,
        "Average viewing angles",
        f"{summary['angles_avg']:.1f}",
        "A typical number of distinct camera directions seeing a sampled point across the repeat cell. This is often the easiest client-facing indicator of how much 3D viewing diversity the survey provides overall.",
    )
    metric_with_help(
        m6,
        "Max viewing angles",
        f"{summary['angles_max']}",
        "The strongest directional diversity inside the sampled repeat cell. These are usually the best-covered overlap zones where several camera families see the same point. Use this as an upper bound, not the expected value everywhere.",
    )

    st.caption(
        "Image hits = total photos containing a point. Viewing angles = unique camera families seeing that point, not repeated photos from the same direction."
    )

    help_toggle(
        "Probe a single point",
        "Use this to answer a very specific question such as: how many photos and how many distinct angles would a point near the strip centre, overlap zone or edge receive? \
The probe point sits inside one repeating line-spacing by photo-spacing cell.",
        key="point_coverage_probe",
    )

    probe_col1, probe_col2 = st.columns(2)
    with probe_col1:
        probe_x_display = st.slider(
            f"Probe across-track position ({dist_unit})",
            min_value=float(m_to_unit(-0.5 * line_spacing_pc, dist_unit)),
            max_value=float(m_to_unit(0.5 * line_spacing_pc, dist_unit)),
            value=0.0,
            step=max(0.5, round(m_to_unit(line_spacing_pc / 40.0, dist_unit), 1)),
        )
    with probe_col2:
        probe_y_display = st.slider(
            f"Probe along-track position ({dist_unit})",
            min_value=float(m_to_unit(-0.5 * photo_spacing_pc, dist_unit)),
            max_value=float(m_to_unit(0.5 * photo_spacing_pc, dist_unit)),
            value=0.0,
            step=max(0.5, round(m_to_unit(photo_spacing_pc / 40.0, dist_unit), 1)),
        )

    probe_x = unit_to_m(probe_x_display, dist_unit)
    probe_y = unit_to_m(probe_y_display, dist_unit)
    coverage_probe = point_coverage_at(probe_x, probe_y, coverage_result["sources"])

    p1, p2 = st.columns(2)
    with p1:
        metric_with_help(
            p1,
            "Probe image hits",
            str(coverage_probe["hits"]),
            "This is the exact number of photos covering the chosen probe point. Use it when a client asks about a specific position, such as the strip centre, overlap zone, or edge, rather than the whole coverage pattern.",
        )
        metric_with_help(
            p1,
            "Probe viewing angles",
            str(coverage_probe["unique_angles"]),
            "This is the exact number of distinct camera directions covering the chosen probe point. It helps explain whether that location is being seen from just one direction, or from several different look angles that improve 3D interpretation.",
        )
    with p2:
        st.markdown("**Viewing angle families at probe point**")
        if coverage_probe["families"]:
            st.write(", ".join(coverage_probe["families"]))
        else:
            st.write("No enabled camera footprints intersect this point.")

    help_toggle(
        "Heatmaps",
        "The heatmaps show how sampled coverage changes across one repeating cell of the survey pattern. Red cells mark sampled zero-coverage holes and the colour scale starts at zero, so it is easier to spot likely gaps. The separate warning above now also uses an exact geometric zero-coverage check when available. The current precision setting controls how many test points are laid across that repeat cell, for example standard is a 61 × 61 sample grid. Central overlap zones will often have more hits and sometimes more unique viewing angles than edge zones.",
        key="point_coverage_heatmaps",
    )

    coverage_fig_hits = coverage_heatmap_figure(
        coverage_result["xs"],
        coverage_result["ys"],
        coverage_result["hits"],
        title="Point Coverage Heatmap — Total Image Hits",
        colorbar_label="Image hits per point",
        probe_xy=(probe_x, probe_y),
    )
    coverage_fig_angles = coverage_heatmap_figure(
        coverage_result["xs"],
        coverage_result["ys"],
        coverage_result["angle_counts"],
        title="Point Coverage Heatmap — Distinct Viewing Angles",
        colorbar_label="Distinct viewing angles",
        probe_xy=(probe_x, probe_y),
    )

    h1, h2 = st.columns(2)
    with h1:
        st.pyplot(coverage_fig_hits)
    with h2:
        st.pyplot(coverage_fig_angles)

    st.caption(
        f"Repeat cell used for the estimate: {m_to_unit(line_spacing_pc, dist_unit):.0f} {dist_unit} line spacing by {m_to_unit(photo_spacing_pc, dist_unit):.0f} {dist_unit} photo spacing."
    )
else:
    st.info("Point coverage could not be estimated from the current camera footprints.")

st.markdown("---")


st.subheader("Export Results")
report_figures = [("Footprint plan view", fig_to_png_bytes(fig_fp))]
if across_sols:
    report_figures.append(("Cross-section view", fig_to_png_bytes(fig_xs)))
if mc and multistrip_solutions:
    report_figures.append(("Multi-strip plan view", fig_to_png_bytes(fig_ms)))
if coverage_fig_hits is not None:
    report_figures.append(("Point coverage heatmap - image hits", fig_to_png_bytes(coverage_fig_hits)))
if coverage_fig_angles is not None:
    report_figures.append(("Point coverage heatmap - viewing angles", fig_to_png_bytes(coverage_fig_angles)))

col_exp1, col_exp2 = st.columns(2)
with col_exp1:
    excel_bytes = make_excel_export(settings_rows, system_rows, camera_rows)
    st.download_button(
        label="Download Excel data report",
        data=excel_bytes,
        file_name="oblique_planner_report.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
with col_exp2:
    word_bytes = make_word_export(settings_rows, system_rows, camera_rows, report_figures=report_figures)
    st.download_button(
        label="Download Word client report",
        data=word_bytes,
        file_name="oblique_planner_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

st.markdown("---")
# ─────────────────────────────────────────────────────────────────────────────
# Formula trace
# ─────────────────────────────────────────────────────────────────────────────

st.subheader("Formula Trace")
with st.expander("Show intermediate calculations per camera", expanded=False):
    for cam, sol, colour in solutions:
        bd = get_body(cam["body"])
        inner_gx, outer_gx = corner_inner_outer(sol)
        inner_len, outer_len = along_lengths_for_display(sol)
        inner_gsd = min(sol.near_gsd_m, sol.far_gsd_m)
        outer_gsd = max(sol.near_gsd_m, sol.far_gsd_m)
        tilt_h = 90.0 - sol.tilt_from_nadir_deg
        ob = obliqueness_ratio(sol)

        st.markdown(
            f"<span style='color:{colour}'>●</span> **{sol.label}**",
            unsafe_allow_html=True
        )
        st.markdown(f"""
| Quantity | Formula | Value |
|---|---|---|
| Sensor native | — | {bd['w_mm']} × {bd['h_mm']} mm, {bd['w_px']} × {bd['h_px']} px |
| Orientation | **{sol.orientation}** | across: **{sol.sensor_across_mm:.4f} mm**, along: **{sol.sensor_along_mm:.4f} mm** |
| Focal length | — | **{cam['focal_mm']} mm** |
| Pixel size | `sensor_across / image_across_px` | **{sol.pixel_size_mm*1000:.3f} µm** |
| Tilt | — | **{sol.tilt_from_nadir_deg:.2f}° from nadir** / **{tilt_h:.2f}° from horizontal** |
| Tilt axis | — | **{sol.tilt_axis}** |
| Half FOV across | `atan(sensor_across / (2×fl))` | {sol.half_fov_across_deg:.4f}° → full {sol.full_fov_across_deg:.4f}° |
| Half FOV along | `atan(sensor_along / (2×fl))` | {sol.half_fov_along_deg:.4f}° → full {sol.full_fov_along_deg:.4f}° |
| Diag PP→edge | `sqrt((sensor_across/2)² + fl²)` | **{sol.diag_image_mm:.4f} mm** |
| Inner edge | `H × tan(θ − φ_w)` | **{m_to_unit(abs(inner_gx), dist_unit):.2f} {dist_unit}** from nadir |
| Outer edge | `H × tan(θ + φ_w)` | **{m_to_unit(abs(outer_gx), dist_unit):.2f} {dist_unit}** from nadir |
| Inner slant | `sqrt(H² + Gx²)` | **{min(sol.near_slant_m, sol.far_slant_m):.2f} m** |
| Outer slant | `sqrt(H² + Gx²)` | **{max(sol.near_slant_m, sol.far_slant_m):.2f} m** |
| Inner length | exact 4-corner | **{m_to_unit(inner_len, dist_unit):.2f} {dist_unit}** |
| Outer length | exact 4-corner | **{m_to_unit(outer_len, dist_unit):.2f} {dist_unit}** |
| Inner GSD | `px_mm × slant_mm / diag_mm` | **{fmt_gsd(inner_gsd)}** |
| Centre GSD | — | **{fmt_gsd(sol.centre_gsd_m)}** |
| Outer GSD | `px_mm × slant_mm / diag_mm` | **{fmt_gsd(outer_gsd)}** |
| **Obliqueness ratio** | `outer GSD / inner GSD` | **{ob:.3f}×** |
        """)

st.markdown("---")

# ─────────────────────────────────────────────────────────────────────────────
# Assumptions
# ─────────────────────────────────────────────────────────────────────────────

with st.expander("ℹ️ Assumptions, conventions and spreadsheet verification"):
    st.markdown("""
### Spreadsheet verification (Oblique_setup9_working_2.xls)

**Landscape sheet** — Sony A7R V, nadir fl=21 mm, oblique fl=50 mm, 50° from horizontal, GSD 8.5 cm:

| Value | Our app | Spreadsheet |
|---|---|---|
| Flying height | 469.737 m | 469.737 m |
| Oblique inner edge | 233.820 m | 233.820 m |
| Oblique outer edge | 635.679 m | 635.679 m |
| Inner length | 368.473 m | 368.473 m |
| Outer length | 555.051 m | 555.051 m |
| Inner GSD | 3.877 cm/px | 3.877 cm/px |
| Outer GSD | 5.840 cm/px | 5.840 cm/px |
| Slope to inner | 524.714 m | 524.714 m |
| Slope to outer | 790.405 m | 790.405 m |

All values match to 3+ decimal places.

**Portrait sheet** — uses landscape-mounted oblique cameras (long axis across-track) at a different
flying height (H=631.58 m, nadir fl=48 mm, GSD 5 cm). To reproduce that sheet, set the L/R oblique
cameras to **Landscape** orientation. Both sheets match our model when the correct orientation is used.

### Obliqueness ratio
Defined as `outer GSD / inner GSD`. Values:
- **1.0** = nadir camera (uniform GSD across the image)
- **1.5** = mild oblique (50% more GSD at the far edge than inner edge)
- **>3** = highly oblique (consider whether far-edge GSD meets mission requirements)

### Sensor orientation convention

| Setting | Sensor dimension across-track | Along-track | Typical use |
|---|---|---|---|
| **Portrait** | Short (narrow) axis | Long axis | L/R oblique |
| **Landscape** | Long axis | Short axis | Nadir camera |

### GSD formula (slant-plane, matching spreadsheet)
```
GSD = pixel_size_mm × slant_2d_mm / diag_image_mm
diag_image_mm = sqrt((sensor_across/2)² + focal_length²)
```

### Photo spacing
Uses the **inner (minimum) footprint length** so the target forward overlap is
met at the most constrained position. The far edge will have higher overlap than requested.

### Assumptions (v1)
- Flat terrain
- Pinhole camera, square pixels
- Pure single-axis tilt per camera
- No lens distortion, wind drift, or lever-arm effects
    """)

st.markdown("---")
st.caption(
    "Oblique Survey Planner v3  ·  Flat terrain  ·  Pinhole model  ·  "
    "Verified against Oblique_setup9_working_2.xls"
)
